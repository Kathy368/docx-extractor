# -*- coding: utf-8 -*-
"""
检测要求和检测结论字段专属提取方案2
功能：仅重新提取 Word 检测报告中的"检测要求"和"检测结论"两个字段，
      输出新 Excel，通过"文件路径"列与原汇总结果一一对应，便于复制回填。

版本：2.0
更新说明：补充多列检测要求表提取逻辑，覆盖以下特殊格式：
  1. 合并表头的多列检测要求表（北方现代、澳美、爱博升等）
  2. 2列检测要求表无结论列（北汽等）
  3. "结果"列作为结论类列的识别与排除
"""

# ============================================================
# 【配置区 - 改这里就行】
# ============================================================

# 输入：现有的汇总结果 Excel（必须包含"文件路径"列）
INPUT_EXCEL = r"D:\Kathy\PDF提取工具\中英文汇总提取结果_20260722_003948全提取2筛选版.xlsx"

# 输出：新的专属提取结果 Excel
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\检测要求和检测结论字段专属提取结果5.xlsx"

# 测试限制：0=跑全部；N>0=只跑前 N 条（用于快速调测）
TEST_LIMIT = 0

# 读取的 sheet 名
SHEET_NAME = "提取结果汇总"

# 是否显示进度条
SHOW_PROGRESS = True

# ============================================================
# 【依赖导入】
# ============================================================
import os
import re
import html
import pandas as pd
from docx import Document
from tqdm import tqdm


# ============================================================
# 【工具函数】
# ============================================================
def get_cell_text(cell):
    """读取单元格完整文本（含 SDT 内容控件）"""
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    return html.unescape(result).strip()


def get_row_cells_text_fast(row):
    """读取一行所有单元格文本（基于 XML 正则，快速但合并单元格可能少报列数）"""
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cells_text.append(html.unescape(cell_text).strip())
    return cells_text


def normalize_text(text):
    """文本清洗"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def contains_pass(text):
    """是否包含符合语义"""
    if not text:
        return False
    t = text.lower()
    return any(k in t for k in ['符合', '合格', '通过', 'pass', 'comply', 'complies',
                                 'conform', 'conforms', '阴性', '未检出'])


def contains_fail(text):
    """是否包含不符合语义"""
    if not text:
        return False
    t = text.lower()
    return any(k in t for k in ['不符合', '不合格', '不通过', 'fail', 'failed', 'failing'])


def is_nd(text):
    """是否为 N.D.（未检出）"""
    if not text:
        return False
    return text.strip().lower() in ['nd', 'n.d.', 'n.d', '未检出']


def is_numeric_result(val):
    """是否为具体数值结果（按用户要求视为不符合）"""
    if not val or val in ['/', '-', '—', 'P', 'X', '']:
        return False
    if is_nd(val) or contains_pass(val) or contains_fail(val):
        return False
    return bool(re.search(r'\d', val))


def is_conclusion_like(text):
    """判断文本是否像结论列内容（短结论词）"""
    if not text:
        return False
    t = text.strip().lower()
    if len(t) > 15:
        return False
    conclusion_keywords = [
        '结论', 'conclusion',
        '符合', '不符合', '合格', '不合格', '通过', '不通过',
        'pass', 'fail', 'failed', 'passing', 'failing',
        'conform', 'conforms', 'conformed', 'non-conform', 'nonconform',
        'nd', 'n.d.', 'n.d', '未检出', '阳性', '阴性'
    ]
    return any(k in t for k in conclusion_keywords)


def classify_text(text):
    """对整段文本判断：pass / fail / partial / unknown"""
    if not text:
        return 'unknown'
    has_fail = contains_fail(text)
    has_pass = contains_pass(text)
    if has_fail and not has_pass:
        return 'fail'
    if has_fail and has_pass:
        return 'partial'
    if has_pass and not has_fail:
        return 'pass'
    return 'unknown'


def format_fail_details(items):
    """格式化不符合项列表"""
    seen = set()
    details = []
    for item, val in items:
        key = f"{item}|{val}"
        if key in seen:
            continue
        seen.add(key)
        if item:
            if val in ['不符合', 'Fail', 'failed', '不合格', '不通过']:
                details.append(item)
            else:
                details.append(f"{item}({val})")
        else:
            details.append(val)
    return details


def summarize_conclusion_items(items, all_fail_short=False):
    """
    综合多项结论：
    - items: [(项目, 值), ...]
    - all_fail_short=True 时，全部不符合只输出"不符合"
    """
    if not items:
        return ""

    pass_items = [(i, v) for i, v in items
                  if (contains_pass(v) and not contains_fail(v)) or is_nd(v)]
    fail_items = [(i, v) for i, v in items
                  if contains_fail(v) or is_numeric_result(v)]

    if fail_items and not pass_items:
        if all_fail_short:
            return "不符合"
        details = format_fail_details(fail_items)
        if details:
            return "不符合，不符合的内容有：" + "、".join(details)
        return "不符合"

    if fail_items:
        details = format_fail_details(fail_items)
        if details:
            return "部分项不符合，不符合的内容有：" + "、".join(details)
        return "部分项不符合"

    if pass_items:
        return "符合"

    # 只有未知值，保留原文
    return " / ".join(set(v for _, v in items))


# ============================================================
# 【检测要求提取】（v2.0 新增合并表头+多列补充逻辑）
# ============================================================
def extract_requirement_from_paragraphs(doc):
    """
    从段落文本中提取检测要求（兜底逻辑）

    适用场景：报告没有独立的检测要求表格，检测要求直接写在正文段落中。
    典型文件：奥托立夫-SAB导流片、爱博升-双组份多元结构胶等。
    """
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paras:
        return ""

    stop_patterns = [
        '检测结果', 'Test Result',
        '检测仪器', 'Test Instrument',
        '检测流程', 'Test Flow',
        '样品描述', 'Sample Description',
        '样品照片', 'Sample photo',
        '判定标准', 'Decision criteria',
        '****报告结束', 'End of Report'
    ]

    method_stop_patterns = [
        '检测结果', 'Test Result',
        '检测仪器', 'Test Instrument',
        '样品描述', 'Sample Description',
        '样品照片', 'Sample photo',
        '****报告结束', 'End of Report'
    ]

    # ============================================================
    # 策略1：找"检测要求"标题后的内容
    # ============================================================
    req_lines = []
    capture = False
    for text in paras:
        if any(k in text for k in ['检测要求', 'Test Requirement']):
            content = text
            for k in ['检测要求', 'Test Requirement', '：', ':', '；', ';']:
                content = content.replace(k, '')
            content = content.strip()
            if content and len(content) > 5 and '请参考下页' not in content and 'Please refer' not in content:
                req_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if len(text) < 3:
                continue
            req_lines.append(text)

    if req_lines:
        return '\n'.join(req_lines)

    # ============================================================
    # 策略2：找"检测方法"/"检测流程"标题后的内容（奥托立夫、爱博升等）
    # ============================================================
    method_like_patterns = ['检测方法', 'Test Method', '检测流程', 'Test Flow']
    method_lines = []
    capture = False
    for text in paras:
        if any(p in text for p in method_like_patterns):
            content = text
            for k in ['检测方法', 'Test Method', '检测流程', 'Test Flow', '：', ':', '；', ';']:
                content = content.replace(k, '')
            content = content.strip().lower()
            # 过滤掉只剩 method/test/flow 等无意义词的标题残留
            meaningless = {'method', 'methods', 'test', 'tests', 'test method', 'flow', 'flows'}
            if content and len(content) > 5 and content not in meaningless \
                    and '请参考下页' not in text and 'Please refer' not in text:
                method_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in method_stop_patterns):
                capture = False
                continue
            if len(text) < 3:
                continue
            method_lines.append(text)

    if method_lines:
        return '\n'.join(method_lines)

    # ============================================================
    # 策略3：找"检测结果"标题下的编号列表
    # ============================================================
    result_lines = []
    capture = False
    for text in paras:
        if '检测结果' in text or 'Test Result' in text:
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if re.match(r'^[\d一二三四五六七八九十]+[\.\、\．]', text) or text.startswith(('•', '-', '*')):
                result_lines.append(text)

    if result_lines:
        return '\n'.join(result_lines)

    return ""


# ============================================================
# 【检测要求提取】（v2.0 新增合并表头+多列补充逻辑）
# ============================================================
def extract_requirement(doc):
    """
    提取检测要求（v2.0 版）

    覆盖规则（按顺序执行）：
    1. 遍历所有 Word 表格，提取所有匹配"检测要求"的非空表格内容；
    2. 若表格中未提取到，再从段落文本中兜底提取。

    表格提取子规则：
    A. 合并表头预检：处理因表头单元格合并导致 get_row_cells_text_fast 少报列数的表格
       （北方现代、澳美POPs、爱博升POPs/REACH等）
    B. 单列表格：表头仅1列且含"检测要求"
    C. 2列混合表：表头同时含"检测要求"和"结论"，取检测要求列
    D. 通用多列表格：表头含"检测要求"的多列表格，排除"结论"/"结果"列后提取
    """
    all_lines = []

    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        # 通过 fast 获取表头文本（用于关键词检测，列数可能因合并不准）
        first_row_fast = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row_fast)
        fast_col_count = len(first_row_fast)

        # 判断是否含有"检测要求"关键词
        has_req_keyword = any('检测要求' in cell or '测试要求' in cell or 'Test Requirement' in cell
                             for cell in first_row_fast)

        # ================================================================
        # A. 合并表头预检：表头合并导致 fast 列数少于真实列数
        # ================================================================
        true_col_count = len(table.rows[0].cells)
        if fast_col_count < true_col_count and has_req_keyword and len(table.rows) >= 2:
            max_cols = max(len(row.cells) for row in table.rows[1:])

            col_values = [[] for _ in range(max_cols)]
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                for i, val in enumerate(row_cells):
                    if i < max_cols and val:
                        col_values[i].append(val)

            skip_cols = set()
            for i, vals in enumerate(col_values):
                if not vals:
                    continue
                con_count = sum(1 for v in vals if is_conclusion_like(v))
                if con_count / len(vals) >= 0.6:
                    skip_cols.add(i)

            lines = []
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                parts = []
                for i, val in enumerate(row_cells):
                    if i < max_cols and i not in skip_cols and val:
                        parts.append(val)
                if parts:
                    lines.append(' '.join(parts))
            if lines:
                all_lines.extend(lines)
                continue

        # ================================================================
        # B. 单列表格
        # ================================================================
        if fast_col_count == 1:
            header = first_row_fast[0]
            if '检测要求' in header or '测试要求' in header or 'Test Requirement' in header:
                lines = []
                for row in table.rows[1:]:
                    cell_text = get_cell_text(row.cells[0]).strip()
                    if cell_text:
                        lines.append(cell_text)
                if lines:
                    all_lines.extend(lines)
                    continue

        # ================================================================
        # C. 2行2列混合表（检测要求 + 结论）
        # ================================================================
        if fast_col_count == 2:
            has_req = ('检测要求' in first_row_text or '测试要求' in first_row_text or 'Test Requirement' in first_row_text)
            has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text)
            if has_req and has_con and len(table.rows) >= 2:
                if '检测要求' in first_row_fast[1] or '测试要求' in first_row_fast[1] or 'Test Requirement' in first_row_fast[1]:
                    req_col = 1
                else:
                    req_col = 0

                lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    if req_col < len(cells):
                        val = cells[req_col].strip()
                        if val:
                            lines.append(val)
                if lines:
                    all_lines.extend(lines)
                    continue

        # ================================================================
        # D. 通用多列检测要求表处理（非合并表头，2+列）
        # ================================================================
        if has_req_keyword and len(table.rows) >= 2:
            skip_cols = set()
            for i, cell in enumerate(first_row_fast):
                cell_s = cell.strip()
                if '结论' in cell or 'Conclusion' in cell:
                    skip_cols.add(i)
                elif cell_s in ('结果', 'Result'):
                    skip_cols.add(i)

            extract_cols = [i for i, cell in enumerate(first_row_fast)
                           if ('检测要求' in cell or '测试要求' in cell or 'Test Requirement' in cell)
                           and i not in skip_cols]

            if extract_cols:
                lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    row_parts = []
                    for col in extract_cols:
                        if col < len(cells):
                            val = cells[col].strip()
                            if val:
                                row_parts.append(val)
                    if row_parts:
                        lines.append(' '.join(row_parts))
                if lines:
                    all_lines.extend(lines)
                    continue

    # 表格中提取到内容，直接返回
    if all_lines:
        # 去重同时保留顺序
        seen = set()
        unique_lines = []
        for line in all_lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)
        return '\n'.join(unique_lines)

    # ============================================================
    # 兜底：从段落文本提取
    # ============================================================
    return extract_requirement_from_paragraphs(doc)


# ============================================================
# 【检测结论提取】（未变更）
# ============================================================
def extract_conclusion_from_filename(filename):
    """从文件名判断结论"""
    name = os.path.splitext(filename)[0]
    if '不合格' in name or '不符合' in name:
        return '不符合'
    if '合格' in name or '符合' in name:
        return '符合'
    return ""


def extract_conclusion_from_independent_table(doc):
    """从独立检测结论表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 1:
            continue

        header = first_row[0]
        if '检测结论' not in header and '测试结论' not in header and 'Test Conclusion' not in header:
            continue

        lines = []
        for row in table.rows[1:]:
            val = get_cell_text(row.cells[0]).strip()
            if val:
                lines.append(val)

        full_text = '\n'.join(lines)
        if not full_text:
            return ""

        cat = classify_text(full_text)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)
        if cat == 'partial':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)

        # 无法判断，保留原文
        return full_text

    return ""


def extract_conclusion_from_mixed_table(doc):
    """从2行2列检测要求/结论混合表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 2:
            continue

        first_row_text = ' | '.join(first_row)
        has_req = ('检测要求' in first_row_text or '测试要求' in first_row_text or 'Test Requirement' in first_row_text)
        has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text)
        if not (has_req and has_con):
            continue

        con_col = 1
        if '结论' in first_row[0] or 'Conclusion' in first_row[0]:
            con_col = 0

        items = []
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if con_col < len(cells):
                val = cells[con_col].strip()
                if val:
                    items.append(("", val))

        if items:
            summary = summarize_conclusion_items(items, all_fail_short=True)
            if summary:
                return summary
            return "\n".join([v for _, v in items])

    return ""


def extract_conclusion_from_result_tables(doc):
    """从检测结果表结论列/行综合判断"""
    all_items = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row)
        if len(first_row) == 1:
            continue

        is_result = any(k in first_row_text for k in ['结果', 'Result', '结论',
                                                       'Conclusion', '限值', 'Limit'])
        if not is_result:
            continue

        # 结论列
        con_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['结论', 'Conclusion']):
                con_col = idx
                break

        # 检测项目列
        item_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['检测项目', 'Test Item', 'Item', '测试项目']):
                item_col = idx
                break

        current_item = ""
        table_rows = list(table.rows)

        for row_idx, row in enumerate(table_rows[1:], start=1):
            cells = get_row_cells_text_fast(row)
            if not cells:
                continue

            first_cell = cells[0].strip() if len(cells) > 0 else ""
            is_conclusion_row = first_cell in ['结论', 'Conclusion']

            # 更新当前项目（处理合并单元格），结论行不更新
            if item_col is not None and item_col < len(cells) and not is_conclusion_row:
                item_val = cells[item_col].strip()
                if item_val and item_val not in ['检测项目', 'Test Item', 'Item',
                                                  '测试项目', '结论', 'Conclusion']:
                    current_item = item_val

            # 结论列
            if con_col is not None and con_col < len(cells):
                val = cells[con_col].strip()
                if val and val not in ['/', '-', '—']:
                    item_for = current_item if current_item not in ['结论', 'Conclusion'] else ""
                    all_items.append((item_for, val))

            # 结论行
            if is_conclusion_row:
                for idx, val in enumerate(cells[1:], start=1):
                    val = val.strip()
                    if not val or val in ['/', '-', '—']:
                        continue

                    target_item = ""
                    # 向上扫描该列，找结果列有具体数字的项目
                    for prev_row in reversed(table_rows[1:row_idx]):
                        prev_cells = get_row_cells_text_fast(prev_row)
                        if idx < len(prev_cells):
                            prev_res = prev_cells[idx].strip()
                            prev_item = (prev_cells[item_col].strip()
                                         if item_col is not None and item_col < len(prev_cells)
                                         else "")
                            if is_numeric_result(prev_res) and prev_item \
                                    and prev_item not in ['结论', 'Conclusion']:
                                target_item = prev_item
                                break

                    if not target_item:
                        target_item = current_item if current_item not in ['结论', 'Conclusion'] else ""

                    all_items.append((target_item, val))

    if all_items:
        return summarize_conclusion_items(all_items, all_fail_short=True)
    return ""


def extract_conclusion(doc, filename):
    """按优先级提取检测结论"""
    # 1. 文件名
    con = extract_conclusion_from_filename(filename)
    if con:
        return con

    # 2. 独立检测结论表格
    con = extract_conclusion_from_independent_table(doc)
    if con:
        return con

    # 3. 2行2列混合表格
    con = extract_conclusion_from_mixed_table(doc)
    if con:
        return con

    # 4. 从检测结果表综合判断
    con = extract_conclusion_from_result_tables(doc)
    if con:
        return con

    return ""


# ============================================================
# 【主程序】
# ============================================================
def main():
    print("=" * 70)
    print("检测要求和检测结论字段专属提取方案2")
    print(f"输入：{INPUT_EXCEL}")
    print(f"输出：{OUTPUT_EXCEL}")
    if TEST_LIMIT > 0:
        print(f"[测试模式] 只处理前 {TEST_LIMIT} 条")
    print("=" * 70)

    # 1. 读取汇总 Excel
    print("\n[1/4] 读取汇总 Excel...")
    df = pd.read_excel(INPUT_EXCEL, sheet_name=SHEET_NAME)
    total = len(df)
    print(f"      共 {total} 行")

    if '文件路径' not in df.columns:
        print("[错误] Excel 中不存在'文件路径'列")
        return

    file_paths = df['文件路径'].dropna().tolist()
    if TEST_LIMIT > 0:
        file_paths = file_paths[:TEST_LIMIT]

    # 2. 批量处理
    print("\n[2/4] 开始提取检测要求和检测结论...")
    results = []

    iterator = tqdm(file_paths, desc="提取进度", unit="份") if SHOW_PROGRESS else file_paths
    for fp in iterator:
        try:
            doc = Document(fp)
            req = extract_requirement(doc)
            con = extract_conclusion(doc, os.path.basename(fp))
            results.append({
                "文件路径": fp,
                "检测要求（新）": req,
                "检测结论（新）": con,
            })
        except Exception as e:
            results.append({
                "文件路径": fp,
                "检测要求（新）": "",
                "检测结论（新）": f"提取失败: {str(e)[:80]}",
            })

    # 3. 保存结果
    print("\n[3/4] 保存结果 Excel...")
    df_out = pd.DataFrame(results)
    df_out = df_out[["文件路径", "检测要求（新）", "检测结论（新）"]]

    # 尝试保存 Excel，若文件被占用则自动换名
    saved_path = None
    try:
        df_out.to_excel(OUTPUT_EXCEL, index=False)
        saved_path = OUTPUT_EXCEL
        print(f"      已保存至：{saved_path}")
    except Exception as e:
        import datetime
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        alt_path = OUTPUT_EXCEL.replace('.xlsx', f'_备份{ts}.xlsx')
        try:
            df_out.to_excel(alt_path, index=False)
            saved_path = alt_path
            print(f"      [注意] 原文件被占用，已另存为：{alt_path}")
            print(f"      占用原因：{e}")
        except Exception as e2:
            print(f"      [错误] Excel 保存失败（{e}），尝试保存 CSV...")
            csv_path = OUTPUT_EXCEL.replace('.xlsx', f'_备份{ts}.csv')
            try:
                df_out.to_csv(csv_path, index=False, encoding='utf_8_sig')
                saved_path = csv_path
                print(f"      已保存 CSV 至：{csv_path}")
            except Exception as e3:
                print(f"      [严重错误] 所有保存方式均失败：{e3}")
                # 至少保留数据在内存中，不崩溃

    # 4. 统计
    print("\n[4/4] 统计信息：")
    print(f"      处理文件数：{len(results)}")
    print(f"      检测要求非空：{df_out['检测要求（新）'].astype(str).str.strip().ne('').sum()}")
    print(f"      检测结论非空：{df_out['检测结论（新）'].astype(str).str.strip().ne('').sum()}")
    if saved_path:
        print(f"      输出文件：{saved_path}")
    print("\n[完成] 可直接按'文件路径'列 VLOOKUP 或复制填回原汇总表。")


if __name__ == "__main__":
    main()
