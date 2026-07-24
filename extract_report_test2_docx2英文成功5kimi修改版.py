# -*- coding: utf-8 -*-
"""
英文检测报告信息提取工具 5
功能：从英文 RoHS/ELV/REACH 检测报告 Word 文档中提取关键信息，生成 Excel 表格

改进点：
1. 支持自动修复部分损坏的 docx 文件（NULL 关系、Bad CRC-32 等）
2. 新增无序号标注时的样品部件总数统计（按 Test item(s) 行数）
3. 新增图片型英文报告的 OCR 识别兜底方案
4. 集成《检测要求和检测结论字段专属提取方案2》的英文检测要求/检测结论提取逻辑，
   覆盖合并表头、多列表格、段落兜底等更多格式，并保留原有逻辑作为兜底。

OCR 依赖（可选）：
- 首选：Tesseract OCR + pytesseract + Pillow
  安装：pip install pytesseract Pillow
  并下载安装 Tesseract：https://github.com/UB-Mannheim/tesseract/wiki
- 备选：easyocr 或 paddleocr

基础依赖库：python-docx, pandas, openpyxl
安装命令：pip install python-docx pandas openpyxl
"""
import os
import re
import sys
import zipfile
import shutil
import tempfile
import html
import pandas as pd

# ==============================================
# 【配置区】请根据实际情况修改以下参数
# ==============================================
# 待处理的 Word 文档路径（支持 Windows 路径，如 r"D:\Kathy\PDF提取工具\test.docx"）
INPUT_FILE = r"G:\中心实验室\中心实验室-新(2023.07.31)\002 报告\报告组\常规组报告 2017.11月份开始\待审核报告 新\非汽车事业部\A 傲川\A 傲川 2023.4.4\S23040400301C-硅元素.docx"
# 输出 Excel 文件路径
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\英测2提取结果.xlsx"
# 最大检测方法数量（预留列数）
MAX_TEST_METHODS = 100
# 是否显示详细处理日志
VERBOSE = True


# ==============================================
# 【英文字段映射表】英文表头 → 中文含义
# ==============================================
EN_FIELD_MAPPING = {
    # 申请商相关
    "Applicant": "申请商",
    "Client": "申请商",
    "Company": "申请商",
    "Address": "申请商地址",

    # 样品相关
    "Product name": "样品名称",
    "Product Name": "样品名称",
    "Sample name": "样品名称",
    "Sample Name": "样品名称",
    "Sample Description": "样品名称",
    "Part Name": "样品名称",
    "Model": "零件号",
    "Part No.": "零件号",
    "Part Number": "零件号",
    "Model No.": "零件号",
    "Test model": "零件号",
    "Client Ref. Info.": "客户参考信息",
    "Client Reference": "客户参考信息",
    "Material": "材质",
    "Production Date": "生产日期",
    "Manufacturer": "制造商",
    "Manufacturer & Factory": "制造商",
    "Factory": "制造商",
    "OEM": "制造商",
    "Trade mark": "商标",
    "Trade Mark": "商标",

    # 汽车部新增关键字段
    "Vehicle type": "车型",
    "Vehicle model": "车型",
    "Automobile company": "主机厂",
    "Buyer": "主机厂",
    "Supplier": "供应商代码",

    # 样品编号相关
    "Sample No.": "样品编号",
    "Sample Number": "样品编号",
    "Sample number": "样品编号",
    "Sample ID": "样品编号",

    # 日期相关
    "Sample Received Date": "样品接收日期",
    "Date of Receipt": "样品接收日期",
    "Date Received": "样品接收日期",
    "Received Date": "样品接收日期",
    "Receipt Date": "样品接收日期",
    "Testing Period": "检测期间",
    "Test Period": "检测期间",
    "Test Date": "检测期间",
    "Testing Date": "检测期间",
    "Date of Test": "检测期间",
    "Date": "报告日期",
    "Report Date": "报告日期",
    "Issue Date": "报告日期",

    # 人员相关
    "Compiled by": "编制",
    "Prepared by": "编制",
    "Reviewed by": "审核",
    "Checked by": "审核",
    "Approved by": "批准",

    # 其他
    "Note": "说明",
    "Notes": "说明",
    "Remark": "备注",
    "Remarks": "备注",
}


# ==============================================
# 工具函数：路径处理（解决 Windows 长路径问题）
# ==============================================
def normalize_path(file_path):
    """为 Windows 长路径添加 \\\\?\\\\ 前缀"""
    prefix = r"\\\\?\\"
    if sys.platform == 'win32' and not file_path.startswith(prefix):
        abs_path = os.path.abspath(file_path)
        if len(abs_path) >= 260:  # 超过 Windows 传统 MAX_PATH 再加前缀
            return prefix + abs_path
    return file_path


# ==============================================
# 工具函数：修复有问题的 docx 文件
# ==============================================
def fix_docx_if_needed(file_path):
    """
    检查并修复有问题的 docx 文件
    支持：NULL 关系问题、Bad CRC-32 等 zip 损坏
    返回修复后的文件路径，如果不需要修复则返回原路径
    """
    file_path = normalize_path(file_path)

    def _try_open(path):
        from docx import Document
        try:
            Document(path)
            return True
        except Exception:
            return False

    if _try_open(file_path):
        return file_path

    if VERBOSE:
        print(f"  ⚠️  检测到文档损坏，正在自动修复...")

    temp_dir = tempfile.gettempdir()
    prefix = r"\\\\?\\"
    base_name = os.path.basename(file_path.replace(prefix, ''))
    fixed_path = os.path.join(temp_dir, f"fixed_{base_name}")

    try:
        with zipfile.ZipFile(file_path, 'r') as zin:
            # 第一轮：找出所有损坏的媒体文件
            skipped_media = set()
            for item in zin.infolist():
                if not item.filename.startswith('word/media/'):
                    continue
                try:
                    zin.read(item.filename)
                except zipfile.BadZipFile:
                    skipped_media.add(item.filename)
                    if VERBOSE:
                        print(f"    跳过损坏的图片: {item.filename}")

            with zipfile.ZipFile(fixed_path, 'w') as zout:
                for item in zin.infolist():
                    # 跳过已损坏的媒体文件
                    if item.filename in skipped_media:
                        continue

                    # 修复 NULL 关系，并删除指向损坏媒体的关系
                    if item.filename == 'word/_rels/document.xml.rels':
                        try:
                            content = zin.read(item.filename).decode('utf-8')
                            content = re.sub(r'<Relationship[^>]*Target="[^"]*NULL[^"]*"[^>]*/>', '', content)
                            for media in skipped_media:
                                # document.xml.rels 中 Target 为相对 word/ 目录，如 media/image5.png
                                target = media.replace('word/', '')
                                content = re.sub(r'<Relationship[^>]*Target="[^"]*' + re.escape(target) + r'"[^>]*/>', '', content)
                            zout.writestr(item, content)
                        except zipfile.BadZipFile:
                            continue
                        continue

                    # 尝试复制其他文件
                    try:
                        data = zin.read(item.filename)
                        zout.writestr(item, data)
                    except zipfile.BadZipFile:
                        raise

        if _try_open(fixed_path):
            if VERBOSE:
                print(f"  ✅ 文档修复成功")
            return fixed_path
        else:
            raise Exception("修复后仍无法打开文档")
    except Exception as fix_error:
        if VERBOSE:
            print(f"  ❌ 文档修复失败: {fix_error}")
        raise


# ==============================================
# 工具函数：获取单元格完整文本（包括 SDT 内容控件）
# ==============================================
def get_cell_text(cell):
    """
    获取单元格的完整文本，包括 SDT（结构化文档标签/内容控件）中的内容
    """
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    result = html.unescape(result)
    return result.strip()


def get_row_cells_text_fast(row):
    """
    获取一行中所有单元格的文本，包括被 SDT（结构化文档标签）包裹的单元格
    处理 w:gridSpan 水平合并单元格，用空字符串填充被合并的列，
    确保返回列表长度与表格实际网格列数对齐，避免合并表头列索引偏移。
    """
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cell_text = html.unescape(cell_text)
        cells_text.append(cell_text.strip())
        # 处理 gridSpan：提取合并列数，超出1的部分用空字符串填充以对齐列位置
        gs_match = re.search(r'<w:gridSpan[^>]*w:val\s*=\s*"(\d+)"', tc_xml)
        if gs_match:
            span = int(gs_match.group(1))
            for _ in range(span - 1):
                cells_text.append('')
    return cells_text


# 兼容总提取代码的别名
get_row_cells_text = get_row_cells_text_fast


def normalize_text(text):
    """清理文本：去首尾空、统一换行/空格"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def normalize_key(key):
    """清理字段名：去冒号、去首尾空、统一空格"""
    if not key:
        return ""
    key = normalize_text(key)
    key = key.rstrip('：:').strip()
    return key


def normalize_key_bilingual(key):
    """
    针对报告中双语字段名做归一化。
    同时覆盖以下常见格式：
      - 中文(English):     Applicant(申请商)
      - 中文English:       样品名称Sample name
      - English中文:       Sample name 样品名称
      - English : 中文:    Sample name : 样品名称
    去掉非目标语言部分、尾部中英文冒号/空格，保留目标键名。
    若键名纯为英文，则保留原值。
    """
    if not key:
        return ""
    key = normalize_text(key)
    # 去掉英文括号及其中的内容
    key = re.sub(r'\s*\([A-Za-z0-9\s&/\.\-_,#()]+\)\s*', '', key)
    # 去掉尾部中英文冒号
    key = key.rstrip('：:').strip()
    # 去掉末尾连续的非中文字段（英文 + 数字 + 常见标点和空格）
    suffix_match = re.search(r'[A-Za-z0-9\s&/\.\-_,#]+$', key)
    if suffix_match and len(suffix_match.group().strip()) >= 2:
        candidate = key[:suffix_match.start()].strip()
        if candidate:
            key = candidate
    # 去掉开头连续的非中文字段
    prefix_match = re.match(r'^[A-Za-z0-9\s&/\.\-_,#]+', key)
    if prefix_match and len(prefix_match.group().strip()) >= 2:
        candidate = key[prefix_match.end():].strip()
        if candidate:
            key = candidate
    return key.rstrip('：:').strip()


# ==============================================
# 【专题方案2】检测要求/结论辅助函数
# ==============================================
def contains_pass(text):
    """是否包含符合语义"""
    if not text:
        return False
    t = text.lower()
    # 先排除明确的不符合语义（避免 "不符合" 被 "符合" 子串误匹配）
    if any(k in t for k in ['不符合', '不合格', '不通过', '不达标', '不满足']):
        return False
    return any(k in t for k in ['符合', '合格', '通过', '达标', '满足',
                                 'pass', 'comply', 'complies', 'compliant',
                                 'conform', 'conforms', 'conformed',
                                 '阴性', '未检出', '不得检出', 'not detected', 'ok'])


def contains_fail(text):
    """是否包含不符合语义"""
    if not text:
        return False
    t = text.lower()
    # 简单子串匹配
    if any(k in t for k in ['不符合', '不合格', '不通过', '不达标', '不满足', '超标', '超出',
                             'fail', 'failed', 'failing', 'non-conform', 'nonconform',
                             '阳性', 'ol', 'f']):
        return True
    # '检出' 用正则排除 '未检出'（(?<!未)检出 = 前面不是"未"的"检出"）
    # 并排除 '不得检出'（限值要求，非实际检出）、'检出限'（方法学术语）
    if re.search(r'(?<!未)检出', t) and '检出限' not in t and '不得检出' not in t:
        return True
    return False


def is_nd(text):
    """是否为 N.D.（未检出）"""
    if not text:
        return False
    return text.strip().lower() in ['nd', 'n.d.', 'n.d', '未检出', 'not detected']


def is_numeric_result(val):
    """是否为具体数值结果（按用户要求视为不符合）"""
    if not val or val in ['/', '-', '—', 'P', 'X', '']:
        return False
    if is_nd(val) or contains_pass(val) or contains_fail(val):
        return False
    return bool(re.search(r'\d', val))


def is_conclusion_like(text):
    """判断文本是否像结论列内容（短结论词）"""
    if not text:
        return False
    t = text.strip().lower()
    if len(t) > 15:
        return False
    conclusion_keywords = [
        '结论', 'conclusion',
        '符合', '不符合', '合格', '不合格', '通过', '不通过',
        'pass', 'fail', 'failed', 'passing', 'failing',
        'conform', 'conforms', 'conformed', 'non-conform', 'nonconform',
        'nd', 'n.d.', 'n.d', '未检出', '阳性', '阴性'
    ]
    return any(k in t for k in conclusion_keywords)


def classify_text(text):
    """对整段文本判断：pass / fail / partial / unknown"""
    if not text:
        return 'unknown'
    has_fail = contains_fail(text)
    has_pass = contains_pass(text)
    if has_fail and not has_pass:
        return 'fail'
    if has_fail and has_pass:
        return 'partial'
    if has_pass and not has_pass:
        return 'pass'
    return 'unknown'


def format_fail_details(items):
    """格式化不符合项列表"""
    seen = set()
    details = []
    for item, val in items:
        key = f"{item}|{val}"
        if key in seen:
            continue
        seen.add(key)
        if item:
            if val in ['不符合', 'Fail', 'failed', '不合格', '不通过']:
                details.append(item)
            else:
                details.append(f"{item}({val})")
        else:
            details.append(val)
    return details


def summarize_conclusion_items(items, all_fail_short=False):
    """
    综合多项结论：
    - items: [(项目, 值), ...]
    - all_fail_short=True 时，全部不符合只输出"不符合"
    """
    if not items:
        return ""

    pass_items = [(i, v) for i, v in items
                  if (contains_pass(v) and not contains_fail(v)) or is_nd(v)]
    fail_items = [(i, v) for i, v in items
                  if contains_fail(v)]

    if fail_items and not pass_items:
        if all_fail_short:
            return "不符合"
        details = format_fail_details(fail_items)
        if details:
            return "不符合，不符合的内容有：" + "、".join(details)
        return "不符合"

    if fail_items:
        details = format_fail_details(fail_items)
        if details:
            return "部分项不符合，不符合的内容有：" + "、".join(details)
        return "部分项不符合"

    if pass_items:
        return "符合"

    # 只有未知值，保留原文
    return " / ".join(set(v for _, v in items))


# ==============================================
# 函数1：提取基本信息
# ==============================================
def find_column_index(first_row, keywords):
    """根据关键词列表找到列索引"""
    for idx, cell in enumerate(first_row):
        for kw in keywords:
            if kw in cell:
                return idx
    return None


def extract_basic_info(doc):
    """从两列表格中提取英文报告基础字段"""
    fields = {}
    manufacturer_address_found = False

    for table_idx, table in enumerate(doc.tables[:20]):
        if len(table.rows) < 1:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' '.join(first_row)
        first_cell_count = len(first_row)

        # 跳过检测结果表格
        if 'Part No.' in first_row_text or 'Test Items' in first_row_text or 'Test item(s)' in first_row_text:
            continue

        # 跳过检测方法表格
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue

        # 跳过 XRF 限值表格
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue

        # 跳过检测要求/结论表格（单独处理）
        if 'Test Requirement' in first_row_text or 'Test Conclusion' in first_row_text:
            continue

        # 处理两列表格（键值对）
        if first_cell_count == 2:
            for row_idx, row in enumerate(table.rows):
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]

                # 处理只有 1 个单元格的行
                if len(cells) == 1:
                    key = normalize_key_bilingual(cells[0])
                    if key and len(key) < 60:
                        fields[key] = ""
                    continue

                if len(cells) == 2:
                    key = normalize_key_bilingual(cells[0])
                    val = cells[1].strip()
                    key = re.sub(r'\s+', ' ', key).strip()

                    if key and len(key) < 60:
                        # 跳过说明性文字
                        if 'The following sample' in key or 'was/were submitted' in key:
                            continue

                        # 特殊处理：遇到 Manufacturer/Factory 后，下一个 Address 是制造商地址
                        if key in ['Manufacturer', 'Manufacturer & Factory', 'Factory']:
                            manufacturer_address_found = True
                            if val:
                                fields[key] = val
                            continue

                        # 特殊处理 Address
                        if key == 'Address':
                            if not manufacturer_address_found:
                                if val:
                                    fields[key] = val
                            else:
                                if val and '制造商地址' not in fields:
                                    fields['制造商地址'] = val
                            continue

                        if val:
                            fields[key] = val

        # 处理四列表格（编制/审核/批准/日期等）
        elif first_cell_count == 4:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) == 4:
                    for i in range(0, 4, 2):
                        key = normalize_key_bilingual(cells[i])
                        val = cells[i+1].strip()
                        if key and len(key) < 40:
                            if val:
                                fields[key] = val

    return fields


# ==============================================
# 函数2：提取检测要求和检测结论（完整原文）
# ==============================================
def extract_test_requirement_and_conclusion(doc):
    """提取 Test Requirement 和 Test Conclusion 的完整原文（优先从首页提取）"""
    test_requirement = ""
    test_conclusion = ""

    for table_idx, table in enumerate(doc.tables):
        if test_requirement and test_conclusion:
            break

        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c) for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)

        # 情况1：2行2列表格（Test Requirement 和 Conclusion 在同一行）
        if (len(table.rows[0].cells) == 2 and
            'Test Requirement' in first_row_text and
            'Conclusion' in first_row_text):
            if VERBOSE:
                print(f"  找到检测要求/结论表格（2行2列格式，首页）")
            if len(table.rows) >= 2:
                test_requirement = get_cell_text(table.rows[1].cells[0])
                test_conclusion = get_cell_text(table.rows[1].cells[1])
            break

        # 情况2：2行1列表格（Test Requirement 单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Requirement' in first_row_text:
            if not test_requirement:
                if VERBOSE:
                    print(f"  找到检测要求表格（独立表格）")
                if len(table.rows) >= 2:
                    test_requirement = get_cell_text(table.rows[1].cells[0])

        # 情况3：2行1列表格（Test Conclusion 单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Conclusion' in first_row_text:
            if not test_conclusion:
                if VERBOSE:
                    print(f"  找到检测结论表格（独立表格）")
                if len(table.rows) >= 2:
                    test_conclusion = get_cell_text(table.rows[1].cells[0])

    return test_requirement, test_conclusion


# ==============================================
# 【专题方案2】英文检测要求提取
# ==============================================
def extract_requirement_from_paragraphs_en(doc):
    """
    从段落文本中提取 Test Requirement（兜底逻辑）
    """
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paras:
        return ""

    stop_patterns = [
        'Test Result', '检测结果',
        'Test Instrument', '检测仪器',
        'Test Flow', '检测流程',
        'Sample Description', '样品描述',
        'Sample photo', '样品照片',
        'Decision criteria', '判定标准',
        'End of Report', '****报告结束'
    ]

    method_stop_patterns = [
        'Test Result', '检测结果',
        'Test Instrument', '检测仪器',
        'Sample Description', '样品描述',
        'Sample photo', '样品照片',
        'End of Report', '****报告结束'
    ]

    # 策略1：找 "Test Requirement" 标题后的内容
    req_lines = []
    capture = False
    for text in paras:
        if any(k in text for k in ['Test Requirement', 'Test Requirements', 'Requirement', 'Requirements', 'Test Basis', 'Test Standard', 'Standards', '标准']):
            content = text
            for k in ['Test Requirement', 'Test Requirements', 'Requirement', 'Requirements', 'Test Basis', 'Test Standard', 'Standards', '标准', ':', '：', ';', '；']:
                content = content.replace(k, '')
            content = content.strip()
            if content and len(content) > 5 and 'Please refer' not in content and '请参考下页' not in content:
                req_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if len(text) < 3:
                continue
            req_lines.append(text)

    if req_lines:
        return '\n'.join(req_lines)

    # 策略2：找 "Test Method" / "Test Flow" / "Reference" 标题后的内容
    method_like_patterns = ['Test Method', 'Method', 'Test Flow', 'Flow', 'Reference', 'References', 'With reference to']
    method_lines = []
    capture = False
    for text in paras:
        if any(p in text for p in method_like_patterns):
            content = text
            for k in ['Test Method', 'Method', 'Test Flow', 'Flow', 'Reference', 'References', 'With reference to', ':', '：', ';', '；']:
                content = content.replace(k, '')
            content = content.strip().lower()
            meaningless = {'method', 'methods', 'test', 'tests', 'test method', 'flow', 'flows', 'reference', 'references'}
            if content and len(content) > 5 and content not in meaningless \
                    and 'Please refer' not in text and '请参考下页' not in text:
                method_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in method_stop_patterns):
                capture = False
                continue
            if len(text) < 3:
                continue
            method_lines.append(text)

    if method_lines:
        return '\n'.join(method_lines)

    # 策略3：找 "Test Result" 标题下的编号列表
    result_lines = []
    capture = False
    for text in paras:
        if 'Test Result' in text or '检测结果' in text:
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if re.match(r'^[\d一二三四五六七八九十]+[\.\、\．]', text) or text.startswith(('•', '-', '*')):
                result_lines.append(text)

    if result_lines:
        return '\n'.join(result_lines)

    return ""


def _is_requirement_text_en(text):
    """判断文本是否像英文检测要求描述（而非标准编号或检测项目名）"""
    if not text:
        return False
    t = text.lower()
    keywords = ['as specified by', 'according to', 'in accordance with',
                'with reference to', 'based on', 'as required by',
                'to determine', 'to test', 'to screen']
    if any(k in t for k in keywords):
        return True
    if len(text) > 30 and re.search(r'(IEC|ISO|ASTM|EPA|US\s*EPA|GB/T|GB|QC/T)\s*\d', text):
        return True
    return False


def _is_serial_column_en(table, col_idx, sample_rows=5):
    """判断指定列是否主要为数字/序号列"""
    serial_count = 0
    total = 0
    for row in table.rows[1:sample_rows + 1]:
        cells = get_row_cells_text_fast(row)
        if col_idx < len(cells):
            val = cells[col_idx].strip()
            if val:
                total += 1
                if re.match(r'^[\d一二三四五六七八九十]+[\.\、\．)\s]*$', val):
                    serial_count += 1
    return total > 0 and serial_count / total >= 0.5


def _classify_requirement_table_columns_en(first_row):
    """根据表头给每列打角色标签"""
    roles = []
    for cell in first_row:
        cs = cell.strip()
        if any(k in cs for k in ['Test Requirement', 'Requirement', '测试要求']):
            roles.append('req')
        elif any(k in cs for k in ['Conclusion', '结论']):
            roles.append('con')
        elif any(k in cs for k in ['Result', '结果']):
            roles.append('res')
        else:
            roles.append('other')
    return roles


def _determine_requirement_cols_en(table, first_row, col_roles):
    """
    在英文多列检测要求表中，确定真正的要求文本列。
    """
    candidate_cols = [i for i, r in enumerate(col_roles) if r in ('req', 'other')]
    candidate_cols = [c for c in candidate_cols if not _is_serial_column_en(table, c)]

    if not candidate_cols:
        return []

    best_col = None
    best_score = -1
    for col in candidate_cols:
        req_text_count = 0
        total_len = 0
        std_only_count = 0
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if col >= len(cells):
                continue
            val = cells[col].strip()
            if not val:
                continue
            total_len += len(val)
            if _is_requirement_text_en(val):
                req_text_count += 1
            if re.search(r'^(IEC|ISO|ASTM|EPA|US\s*EPA|GB/T|GB|QC/T)\s*\d', val) and len(val) < 60:
                std_only_count += 1

        score = req_text_count * 100 + total_len - std_only_count * 50
        if score > best_score:
            best_score = score
            best_col = col

    return [best_col] if best_col is not None else []


def extract_requirement_en(doc):
    """
    提取英文检测要求 Test Requirement（基于专题方案2 v2.0）
    """
    all_lines = []

    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        first_row_fast = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row_fast)
        fast_col_count = len(first_row_fast)

        has_req_keyword = any('Test Requirement' in cell or 'Requirement' in cell or '测试要求' in cell
                             for cell in first_row_fast)

        # A. 合并表头预检
        true_col_count = len(table.rows[0].cells)
        if fast_col_count < true_col_count and has_req_keyword and len(table.rows) >= 2:
            max_cols = max(len(row.cells) for row in table.rows[1:])

            col_values = [[] for _ in range(max_cols)]
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                for i, val in enumerate(row_cells):
                    if i < max_cols and val:
                        col_values[i].append(val)

            skip_cols = set()
            for i, vals in enumerate(col_values):
                if not vals:
                    continue
                con_count = sum(1 for v in vals if is_conclusion_like(v))
                if con_count / len(vals) >= 0.6:
                    skip_cols.add(i)

            lines = []
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                parts = []
                for i, val in enumerate(row_cells):
                    if i < max_cols and i not in skip_cols and val:
                        parts.append(val)
                if parts:
                    lines.append(' '.join(parts))
            if lines:
                all_lines.extend(lines)
                continue

        # B. 单列表格
        if fast_col_count == 1:
            header = first_row_fast[0]
            if 'Test Requirement' in header or 'Requirement' in header or '测试要求' in header:
                lines = []
                for row in table.rows[1:]:
                    cell_text = get_cell_text(row.cells[0]).strip()
                    if cell_text:
                        lines.append(cell_text)
                if lines:
                    all_lines.extend(lines)
                    continue

        # C. 2行2列混合表（Requirement + Conclusion）
        if fast_col_count == 2:
            has_req = ('Test Requirement' in first_row_text or 'Requirement' in first_row_text or '测试要求' in first_row_text)
            has_con = ('Conclusion' in first_row_text or '结论' in first_row_text)
            if has_req and has_con and len(table.rows) >= 2:
                if 'Test Requirement' in first_row_fast[1] or 'Requirement' in first_row_fast[1] or '测试要求' in first_row_fast[1]:
                    req_col = 1
                else:
                    req_col = 0

                lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    if req_col < len(cells):
                        val = cells[req_col].strip()
                        if val:
                            lines.append(val)
                if lines:
                    all_lines.extend(lines)
                    continue

        # D. 通用多列表格（列角色识别 + 语义过滤）
        if has_req_keyword and len(table.rows) >= 2:
            col_roles = _classify_requirement_table_columns_en(first_row_fast)
            extract_cols = _determine_requirement_cols_en(table, first_row_fast, col_roles)

            if extract_cols:
                # 优先提取明确要求语义的长文本
                req_lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    for col in extract_cols:
                        if col < len(cells):
                            val = cells[col].strip()
                            if val and _is_requirement_text_en(val):
                                req_lines.append(val)

                # 若未命中要求语义，兜底提取候选列全部非空文本
                if not req_lines:
                    for row in table.rows[1:]:
                        cells = get_row_cells_text_fast(row)
                        row_parts = []
                        for col in extract_cols:
                            if col < len(cells):
                                val = cells[col].strip()
                                if val:
                                    row_parts.append(val)
                        if row_parts:
                            req_lines.append(' '.join(row_parts))

                if req_lines:
                    all_lines.extend(req_lines)
                    continue

    if all_lines:
        seen = set()
        unique_lines = []
        for line in all_lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)
        return '\n'.join(unique_lines)

    return extract_requirement_from_paragraphs_en(doc)


# ==============================================
# 【专题方案2】英文检测结论提取
# ==============================================
def _resolve_result_value(val, cells, result_col):
    """
    判断结果列中单个单元格值的结论语义（用于无"Conclusion"列时的兜底判断）。
    返回: ('pass',) / ('fail', detail) / None（跳过）/ ('unknown',)
    """
    val = val.strip()
    if not val or val in ['/', '-', '—', 'P', 'X']:
        return None

    # 1. N.D. → pass
    if is_nd(val):
        return ('pass',)

    # 2. 明确 Fail 语义（且不含 Pass 词） → fail
    if contains_fail(val) and not contains_pass(val):
        return ('fail', val)

    # 3. 明确 Pass 语义（且不含 Fail 词） → pass
    if contains_pass(val) and not contains_fail(val):
        return ('pass',)

    # 4. 同时含 Pass 和 Fail 语义 → fail（偏保守）
    if contains_pass(val) and contains_fail(val):
        return ('fail', val)

    # 5. 含 "/" 的值 → 向左读同行单元格获取上下文
    if '/' in val:
        for i in range(result_col - 1, -1, -1):
            if i >= len(cells):
                continue
            left_val = cells[i].strip()
            if not left_val or left_val in ['/', '-', '—']:
                continue
            if contains_fail(left_val):
                return ('fail', val)
            if contains_pass(left_val) or is_nd(left_val):
                return ('pass',)
        # 左边无任何语义信息 → 保守判为无法判断
        return ('unknown',)

    # 6. 兜底：无法判断 → 需人工查看
    return ('unknown',)


def extract_conclusion_from_filename_en(filename):
    """从文件名判断结论（文件名多为中文命名习惯）"""
    name = os.path.splitext(filename)[0]
    if '不合格' in name or '不符合' in name:
        return '不符合'
    if '合格' in name or '符合' in name:
        return '符合'
    return ""


def extract_conclusion_from_independent_table_en(doc):
    """从独立 Test Conclusion 表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 1:
            continue

        header = first_row[0]
        if 'Test Conclusion' not in header and 'Conclusion' not in header and '测试结论' not in header:
            continue

        lines = []
        for row in table.rows[1:]:
            val = get_cell_text(row.cells[0]).strip()
            if val:
                lines.append(val)

        full_text = '\n'.join(lines)
        if not full_text:
            return ""

        cat = classify_text(full_text)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)
        if cat == 'partial':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)

        return full_text

    return ""


def extract_conclusion_from_mixed_table_en(doc):
    """从2行2列 Requirement/Conclusion 混合表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 2:
            continue

        first_row_text = ' | '.join(first_row)
        has_req = ('Test Requirement' in first_row_text or 'Requirement' in first_row_text or '测试要求' in first_row_text)
        has_con = ('Conclusion' in first_row_text or '结论' in first_row_text or
                   'Result' in first_row_text or '结果' in first_row_text)
        if not (has_req and has_con):
            continue

        con_col = 1
        if 'Conclusion' in first_row[0] or '结论' in first_row[0]:
            con_col = 0

        items = []
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if con_col < len(cells):
                val = cells[con_col].strip()
                if val:
                    items.append(("", val))

        if items:
            summary = summarize_conclusion_items(items, all_fail_short=True)
            if summary:
                return summary
            return "\n".join([v for _, v in items])

    return ""


def _is_result_table_en(table):
    """
    判断表格是否为英文检测结果表。
    不仅看首行表头，还扫描前 10 行，支持跨行合并表头。
    """
    result_keywords = ['Result', 'Conclusion', 'Limit', 'Pass', 'Fail']
    for r_idx in range(min(10, len(table.rows))):
        row_text = ' | '.join(get_row_cells_text_fast(table.rows[r_idx]))
        if any(k in row_text for k in result_keywords):
            return True
    return False


def _locate_conclusion_column_en(table, first_row, max_scan_rows=10):
    """
    在英文结果表中定位结论列。
    1. 先在首行表头中找明确含 Conclusion/结论 的列；
    2. 若未找到，扫描前 max_scan_rows 行，按每列出现结论关键词的频率定位；
    3. 兜底：找 Result/结果 列。
    """
    for idx, cell in enumerate(first_row):
        if any(k in cell for k in ['Conclusion', '结论']):
            return idx, False

    col_scores = [0] * len(first_row)
    conclusion_keywords = ['Conclusion', '结论', 'Pass', 'Fail', 'Compliant', 'Non-conform', 'ND', 'N.D.']
    for r_idx in range(1, min(max_scan_rows + 1, len(table.rows))):
        cells = get_row_cells_text_fast(table.rows[r_idx])
        if len(cells) != len(first_row):
            continue
        for c_idx, val in enumerate(cells):
            v = val.strip()
            if any(k in v for k in conclusion_keywords):
                weight = 2 if len(v) <= 12 else 1
                col_scores[c_idx] += weight

    if max(col_scores) > 0:
        return col_scores.index(max(col_scores)), True

    first_row_text = ' | '.join(first_row)
    for idx, cell in enumerate(first_row):
        if any(k in cell for k in ['Result', '结果']) and \
           not any(k in first_row_text for k in ['Test Result', '检测结果', 'Requirement', '要求']):
            return idx, False

    return None, False


def _read_conclusion_value_en(row, con_col, search_range=2):
    """
    读取一行中的英文结论值。
    若定位列无效，则在右侧相邻列及行尾搜索结论关键词。
    """
    cells = get_row_cells_text_fast(row)
    if not cells:
        return ""

    conclusion_keywords = ['Pass', 'Fail', 'Compliant', 'Non-conform', 'Conforms', 'ND', 'N.D.', 'Not Detected']

    if 0 <= con_col < len(cells):
        val = cells[con_col].strip()
        if val and val not in ['/', '-', '—', 'Conclusion', 'Result']:
            if any(k in val for k in conclusion_keywords) or len(val) <= 15:
                return val

    start_col = min(con_col + 1, len(cells) - 1)
    for c in range(start_col, min(len(cells), start_col + search_range)):
        val = cells[c].strip()
        if val and val not in ['/', '-', '—']:
            if any(k in val for k in conclusion_keywords):
                return val

    for c in range(max(0, len(cells) - 2), len(cells)):
        val = cells[c].strip()
        if val and val not in ['/', '-', '—', 'Conclusion', 'Result']:
            if any(k in val for k in conclusion_keywords):
                return val

    return ""


def extract_conclusion_from_result_tables_en(doc):
    """从英文 Test Result 表格的 Conclusion 列/行综合判断；支持合并单元格、跨行表头。"""
    all_items = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row)
        if len(first_row) == 1:
            continue

        # 判断是否为结果表（支持跨行合并表头）
        if not _is_result_table_en(table):
            continue

        # 跳过说明/备注/注释表
        if any(k in first_row_text for k in ['说明', '备注', '注释', 'Note', 'Remark', 'Annotation']):
            has_item_col = any(k in first_row_text for k in ['Test Item', 'Item',
                                                              '检测项目', '测试项目'])
            if not has_item_col:
                continue

        # 定位结论列
        con_col, _ = _locate_conclusion_column_en(table, first_row)

        is_result_fallback = False
        if con_col is None:
            for idx, cell in enumerate(first_row):
                if any(k in cell for k in ['Result', '结果']) and \
                   not any(k in first_row_text for k in ['Test Result', '检测结果',
                                                          'Requirement', '要求']):
                    con_col = idx
                    is_result_fallback = True
                    break
        if con_col is None:
            continue

        item_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['Test Item', 'Item', '测试项目', 'Test Items']):
                item_col = idx
                break

        current_item = ""
        table_rows = list(table.rows)

        for row_idx, row in enumerate(table_rows[1:], start=1):
            cells = get_row_cells_text_fast(row)
            if not cells:
                continue

            first_cell = cells[0].strip() if len(cells) > 0 else ""
            is_conclusion_row = first_cell in ['Conclusion', '结论']

            if item_col is not None and item_col < len(cells) and not is_conclusion_row:
                item_val = cells[item_col].strip()
                if item_val and item_val not in ['Test Item', 'Item', '测试项目', 'Test Items', 'Conclusion', '结论']:
                    current_item = item_val

            # 读取结论值（支持列偏移兜底）
            val = _read_conclusion_value_en(row, con_col)
            if val:
                item_for = current_item if current_item not in ['Conclusion', '结论'] else ""

                if is_result_fallback:
                    resolved = _resolve_result_value(val, cells, con_col)
                    if resolved is None:
                        continue
                    if resolved[0] == 'pass':
                        all_items.append((item_for, '符合'))
                    elif resolved[0] == 'fail':
                        all_items.append((item_for, resolved[1]))
                    else:
                        all_items.append((item_for, '需人工查看'))
                else:
                    all_items.append((item_for, val))

            if is_conclusion_row:
                for idx, val in enumerate(cells[1:], start=1):
                    val = val.strip()
                    if not val or val in ['/', '-', '—']:
                        continue

                    target_item = ""
                    for prev_row in reversed(table_rows[1:row_idx]):
                        prev_cells = get_row_cells_text_fast(prev_row)
                        if idx < len(prev_cells):
                            prev_res = prev_cells[idx].strip()
                            prev_item = (prev_cells[item_col].strip()
                                         if item_col is not None and item_col < len(prev_cells)
                                         else "")
                            if prev_res and prev_item \
                                    and prev_item not in ['Conclusion', '结论']:
                                target_item = prev_item
                                break

                    if not target_item:
                        target_item = current_item if current_item not in ['Conclusion', '结论'] else ""

                    all_items.append((target_item, val))

    if all_items:
        return summarize_conclusion_items(all_items, all_fail_short=True)
    return ""


def extract_conclusion_from_paragraphs_en(doc):
    """
    兜底：从段落文本中抓取英文检测结论。
    寻找包含 Conclusion / Pass / Fail / Comply 等关键词的句子。
    """
    paras = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    if not paras:
        return ""

    # 直接包含 "Conclusion" / "Test Conclusion" 标题的段落
    capture = False
    captured_lines = []
    for text in paras:
        if re.search(r'Test Conclusion|Conclusion[：:\s]', text, re.IGNORECASE):
            content = re.sub(r'.*?\b(Test Conclusion|Conclusion)[：:\s]*', '', text, count=1, flags=re.IGNORECASE).strip()
            if content and len(content) > 3:
                captured_lines.append(content)
            capture = True
            continue
        if capture:
            if any(k in text for k in ['Test Result', 'Test Instrument', 'Sample Description', 'Sample photo', 'Prepared by', 'End of Report']):
                break
            if len(text) < 3:
                continue
            captured_lines.append(text)

    if captured_lines:
        full = '\n'.join(captured_lines)
        cat = classify_text(full)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            return summarize_conclusion_items([("", v) for v in captured_lines], all_fail_short=True)
        return full

    # 没有明确结论标题时，找含结论语义的单句
    conclusion_keywords = ['pass', 'passed', 'comply', 'complies', 'compliant', 'conform', 'conforms',
                           'fail', 'failed', 'failing', 'not detected', 'nd']
    for text in paras:
        if len(text) < 5 or len(text) > 200:
            continue
        t = text.lower()
        if any(k in t for k in conclusion_keywords):
            # 过滤掉明显不是结论的句子
            if any(k in t for k in ['test item', 'test method', 'instrument', 'sample description']):
                continue
            cat = classify_text(text)
            if cat == 'pass':
                return "符合"
            if cat == 'fail':
                return text
            if '需人工查看' not in text:
                return text

    return ""


def extract_conclusion_en(doc, filename):
    """按优先级提取英文检测结论"""
    con = extract_conclusion_from_filename_en(filename)
    if con:
        return con

    con = extract_conclusion_from_independent_table_en(doc)
    if con:
        return con

    con = extract_conclusion_from_mixed_table_en(doc)
    if con:
        return con

    con = extract_conclusion_from_result_tables_en(doc)
    if con:
        return con

    con = extract_conclusion_from_paragraphs_en(doc)
    if con:
        return con

    return ""


# ==============================================
# 函数3：提取样品部件总数
# ==============================================
def extract_sample_count(doc):
    """
    提取样品部件总数（ds修改版）：
    0. 新增：优先扫描"样品描述"特征表（Sample Description / 产品编号 / 样品序号等），
       找到直接返回；有表但无部件拆分则返回1。
    1. 其次在 Test Result(s) 模块内寻找：
       - 有 Part No / Sample No. 列时，取最大基号；同时用 Part Description / Sample Description 去重行数做二次校验。
       - 无序号列时，只按 Part Description / Sample Description 去重行数统计。
    2. 识别无表头部件表（首行首列为数字编号，第二列为描述）。
    3. 识别邻苯类组合编号表（如 1+9+19+25+26+33）。
    4. 排除 CAS No. 被误当成序号列。
    5. Test Result(s) 模块没有部件信息时，退到独立的 Sample Description 模块，取最大序号。
    """
    max_serial = 0
    desc_count = 0
    found_result_part_info = False
    has_result_table_no_parts = False

    result_section_markers = ['Test Result(s)', 'Test Result', 'Test Results', 'Result(s),mg/kg', '结果,mg/kg']
    result_markers = ['Result', 'Limit', 'Conclusion', 'mg/kg']
    test_item_keywords = {'pb', 'cd', 'hg', 'cr', 'br', 'lead', 'cadmium', 'mercury', 'chromium', 'brominated'}
    serial_keywords = ['Part No', 'Sample No', 'No.',
                       '序号', '部件编号', '部件号', '产品编号', '样品编号',
                       '样品序号', '产品序号', '部件序号',
                       'No', '部件号', 'Serial No', 'Item No']
    desc_keywords = ['Part Description', 'Sample Description', 'Description',
                     '样品描述', '部件描述', '描述', '部件名称', '样品名称']

    # 样品描述表特征关键词（含任一即优先识别为样品描述表）
    sample_desc_markers = ['Sample Description', 'Sample composition', 'Part Description',
                           '样品描述', '样品组成', '样品序号', '产品序号', '部件序号', '产品编号', '部件名称']
    # 样品描述表互斥关键词：含任一结果表特征的表格不按样品描述表处理
    sample_desc_exclude = ['Test item', 'Test method', 'Result', 'Limit', 'mg/kg',
                           'CAS No', 'CAS No.', 'XRF', 'Method',
                           '检测项目', '限值', '检测方法', '检测仪器', '方法检出限']

    part_id_pattern = re.compile(r'^(\d+)(-\d+)?[A-Z]?\)?$')

    def _base_num(val):
        m = re.match(r'^(\d+)', str(val).strip())
        return int(m.group(1)) if m else None

    def _update_max(val):
        nonlocal max_serial
        n = _base_num(val)
        if n and n > max_serial:
            max_serial = n

    def _is_header_like_part_table(first_row):
        if len(first_row) < 2:
            return False
        h0 = first_row[0]
        h1 = first_row[1]
        has_no = any(k in h0 for k in ['No.', 'No', '序号', 'Part No', 'Sample No',
                                        'Serial No', 'Item No',
                                        '产品编号', '样品编号', '样品序号', '产品序号', '部件序号', '部件编号'])
        has_desc = any(k in h1 for k in ['Description', '描述', 'Part', 'Sample'])
        return has_no and has_desc

    # ================================================================
    # 阶段0（新增）：样品描述表优先提取
    # 含 Sample Description / 产品编号 / 样品序号等特征、且不含结果表特征的表格，
    # 优先按样品描述表提取序号/描述列，提取逻辑与结果表趋同。
    # 找到 → 直接返回；有表但无拆分 → 返回 1；未找到 → 继续后续逻辑。
    # ================================================================
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)
        first_row_lower = first_row_text.lower()

        # 跳过方法/限值表
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue

        has_sample_desc = any(k in first_row_text for k in sample_desc_markers)
        has_result_feature = any(k in first_row_text for k in sample_desc_exclude)

        if has_sample_desc and not has_result_feature:
            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                sample_max_serial = 0
                sample_desc_count = 0

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        m = re.match(r'^(\d+)', str(val).strip())
                        if m:
                            n = int(m.group(1))
                            if n > sample_max_serial:
                                sample_max_serial = n

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    sample_desc_count = len(descs)

                count = sample_max_serial if sample_max_serial > 0 else sample_desc_count
                if count > 2000:
                    return "样品部件总数异常请人工检查"
                if count == 0:
                    return 1
                return count

    # ================================================================
    # 阶段1（原有）：检测结果表/部件表提取
    # ================================================================
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)
        first_row_lower = first_row_text.lower()

        # 排除方法/限值表
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue

        # ===== A. 带表头的 Test Result(s) 表 =====
        is_result_table = (any(k in first_row_text for k in result_section_markers) or
                           any(k in first_row_text for k in result_markers) or
                           any(kw in first_row_lower for kw in test_item_keywords))
        if is_result_table:
            # 排除 SVHC/REACH 化学物质清单表：表头含 CAS No. 且含 EC No. 或物质名称
            # 这类表的“序号”是化学物质编号，不是样品部件序号
            is_chemical_list = (
                'CAS' in first_row_text.upper() and
                ('EC' in first_row_text.upper() or
                 'Substance' in first_row_text or
                 '物质名称' in first_row_text)
            )
            if is_chemical_list:
                continue

            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                found_result_part_info = True

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        _update_max(val)

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    if len(descs) > desc_count:
                        desc_count = len(descs)
            else:
                # 有结果表但没有序号/描述列，通常是单一样品的结果表
                # 对于含 CAS 的短汇总表（如 SVHC 汇总结果），同样视为单一样品
                has_result_table_no_parts = True

        # ===== B. 无表头部件表 或 No./序号 头部件表 =====
        is_part_table = False
        start_row_idx = 0
        if len(first_row) >= 2:
            c0 = first_row[0].strip()
            c1 = first_row[1].strip() if len(first_row) > 1 else ""
            if (part_id_pattern.match(c0) and c1 and not c1.isdigit() and
                    not any(k in c1.lower() for k in ['conclusion', 'result', 'method', 'limit', 'note'])):
                is_part_table = True
            elif _is_header_like_part_table(first_row):
                is_part_table = True
                start_row_idx = 1

        if is_part_table:
            numeric_rows = 0
            for row in table.rows[start_row_idx:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                if part_id_pattern.match(val):
                    _update_max(val)
                    numeric_rows += 1
            if numeric_rows > 0:
                found_result_part_info = True

        # ===== C. 邻苯类组合编号表 =====
        if (len(table.rows) >= 3 and
                any(k in first_row_text for k in ['Test Items', '测试项目']) and
                'Result' in first_row_text):
            sub_header = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[1])]
            if sub_header and ('Test Items' in sub_header[0] or '测试项目' in sub_header[0]):
                found_nums = False
                for cell in sub_header:
                    nums = re.findall(r'\b\d+\b', cell)
                    if nums:
                        found_nums = True
                        for n in nums:
                            _update_max(n)
                if found_nums:
                    found_result_part_info = True

    if found_result_part_info:
        count = max_serial if max_serial > 0 else desc_count
        if count > 2000:
            return "样品部件总数异常请人工检查"
        # 有结果表但序号/描述列均无有效值，视为单一样品的汇总结果
        if count == 0:
            return 1
        return count

    # 退到独立的 Sample Description 模块
    count = extract_sample_count_from_sample_description_section(doc)
    if count > 0:
        if count > 2000:
            return "样品部件总数异常请人工检查"
        return count

    # 存在结果表但无部件拆分，视为 1 个样品
    if has_result_table_no_parts:
        return 1

    return "未检测到"


def extract_sample_count_from_sample_description_section(doc):
    """
    从独立的 Sample Description 模块（段落标题）中提取最大序号。
    排除包含 Photo / 照片的段落。
    """
    body = doc.element.body
    in_section = False
    target_table = None
    section_paragraphs = []

    for child in body:
        tag = child.tag.split('}')[-1]
        text = ''.join(child.itertext()).strip()

        if tag == 'p':
            lower = text.lower()
            if 'sample description' in lower:
                if 'photo' in lower or '照片' in text:
                    continue
                in_section = True
                continue
            if in_section and any(k in lower for k in ['test method', 'test result', 'test item', 'conclusion', 'requirement']):
                break

        if in_section:
            if tag == 'p':
                section_paragraphs.append(text)
            elif tag == 'tbl' and target_table is None:
                target_table = child
                break

    max_serial = 0
    found = False

    def _try_match(s):
        nonlocal max_serial, found
        s = s.strip()
        for pattern in [r'^(\d+)\s*[:：.\)、]\s*', r'^(\d+)$']:
            m = re.match(pattern, s)
            if m:
                found = True
                seq = int(m.group(1))
                if seq > max_serial:
                    max_serial = seq
                return True
        return False

    if target_table is not None:
        for tc in target_table.iter():
            if tc.tag.endswith('tc'):
                cell_text = ''.join(tc.itertext()).strip()
                _try_match(cell_text)
    else:
        for text in section_paragraphs:
            for line in re.split(r'[\r\n]+', text):
                _try_match(line.strip())

    # 若段落标题方式没找到，再扫描全文档中表头为 Sample Description 的表格
    if max_serial == 0:
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
            first_row_text = ' '.join(first_row)
            if 'Sample Description' not in first_row_text and 'sample description' not in first_row_text.lower():
                continue
            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                m = re.match(r'^(\d+)', val)
                if m:
                    seq = int(m.group(1))
                    if seq > max_serial:
                        max_serial = seq

    return max_serial


def extract_sample_count_from_test_items(doc):
    """
    兜底：当样品没有按 Part No. 拆分时，根据结果表的 Result(s) 列数推断样品数。
    例如 Test item(s) | Limit | Result(s) 只有一列结果 → 1 个样品。
    """
    max_count = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 只关心结果表
        if not ('Result' in first_row_text or 'result' in first_row_text.lower()):
            continue
        # 已带 Part/Sample 编号的表由 extract_sample_count 处理，这里跳过
        if any(k in first_row_text for k in ['Part No', 'Sample No', 'No.']):
            continue

        # 若表头是数字编号（如 1 | 2 | 3），取最大编号
        nums = []
        for cell in first_row:
            cell = cell.strip()
            if re.match(r'^\d+$', cell):
                nums.append(int(cell))
        if nums:
            max_count = max(max_count, max(nums))
            continue

        # 若有 Test item(s) + Result(s) 列，则视为 1 个样品
        has_test_item = any('Test item' in c or 'item' in c.lower() for c in first_row)
        has_result = any('Result' in c for c in first_row)
        if has_test_item and has_result:
            max_count = max(max_count, 1)

    return max_count


def extract_test_methods(doc):
    """从 Chemical Test 章节的检测方法表格中提取"""
    methods = []
    seen_items = set()

    for table_idx, table in enumerate(doc.tables[:30]):
        if len(table.rows) < 3:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 检测方法表格特征：有 Test item / Testing item / Test Items 列和 Test method 列
        first_row_lower = first_row_text.lower()
        has_item_col = any(k in first_row_text for k in ['Test item', 'Testing item', 'Test Items', 'Testing Items', 'Items'])
        has_method_col = 'test method' in first_row_lower or 'method' in first_row_lower
        if not (has_item_col and has_method_col):
            continue

        if VERBOSE:
            print(f"  找到检测方法表格（{len(table.rows)}行）")

        item_col = find_column_index(first_row, ['Test item', 'Testing item', 'Test Items', 'Testing Items', 'Items', 'Item'])
        method_col = find_column_index(first_row, ['Test method', 'Testing method', 'Method'])
        instrument_col = find_column_index(first_row, ['Test instrument', 'Instrument', 'Equipment', 'Apparatus', '测试仪器', '检测仪器'])

        for row_idx, row in enumerate(table.rows[1:]):
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if len(cells) < 2:
                continue

            item = cells[item_col].strip() if item_col is not None and item_col < len(cells) else cells[0].strip()
            method = cells[method_col].strip() if method_col is not None and method_col < len(cells) else ""
            instrument = cells[instrument_col].strip() if instrument_col is not None and instrument_col < len(cells) else ""

            # 跳过说明行和空行
            if not item or not method:
                continue
            if 'Limit' in item or '△' in item or 'Note' in item:
                continue

            # 按检测项目名称去重
            if item in seen_items:
                continue
            seen_items.add(item)

            methods.append({
                "检测项目": item,
                "检测方法": method,
                "检测仪器": instrument
            })

    return methods


# ==============================================
# 函数5：判断是否为图片型文档
# ==============================================
def is_image_based_document(doc):
    """
    判断文档是否为图片型（正文几乎没有可提取文本，如 CTT...EN_sign.docx）
    """
    paras = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    total_text = ' '.join(paras)

    # 非空段落少于 3 个，且总文本长度小于 50，认为是图片型
    if len(paras) < 3 and len(total_text) < 50:
        return True

    # 检查前 5 个表格是否都几乎无文本
    empty_tables = 0
    for table in doc.tables[:5]:
        has_text = False
        for row in table.rows[:3]:
            for cell in row.cells:
                if get_cell_text(cell):
                    has_text = True
                    break
            if has_text:
                break
        if not has_text:
            empty_tables += 1

    if empty_tables >= 3 and len(total_text) < 100:
        return True

    return False


# ==============================================
# 函数6：图片型文档 OCR 识别
# ==============================================
def extract_images_from_docx(docx_path):
    """
    从 docx 文件中提取所有内嵌图片到临时目录
    返回图片文件路径列表
    """
    image_paths = []
    temp_dir = tempfile.mkdtemp(prefix="docx_ocr_")

    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            for item in zin.infolist():
                if item.filename.startswith('word/media/'):
                    try:
                        data = zin.read(item.filename)
                        ext = os.path.splitext(item.filename)[1].lower()
                        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                            img_name = f"img_{len(image_paths):04d}{ext}"
                            img_path = os.path.join(temp_dir, img_name)
                            with open(img_path, 'wb') as f:
                                f.write(data)
                            image_paths.append(img_path)
                    except Exception:
                        continue
    except Exception:
        pass

    return image_paths, temp_dir


def ocr_image(image_path):
    """
    对单张图片进行 OCR 识别
    优先使用 pytesseract；未安装则尝试 easyocr/paddleocr
    对低分辨率/小字号图片会自动放大，提升识别率
    """
    from PIL import Image, ImageEnhance

    try:
        img = Image.open(image_path)
    except Exception:
        return ""

    # 跳过极小图片（通常是图标、装饰线）
    if img.width < 80 or img.height < 20:
        return ""

    # 转换并增强：统一转灰度，低分辨率时放大
    img = img.convert('L')
    scale = 1.0
    if img.height < 80:
        scale = max(scale, 200.0 / img.height)
    if img.width < 300:
        scale = max(scale, 600.0 / img.width)
    if scale > 1.0:
        new_size = (int(img.width * scale), int(img.height * scale))
        img = img.resize(new_size, Image.LANCZOS)

    # 适度增强对比度
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.5)

    # 尝试 pytesseract
    try:
        import pytesseract
        # 兼容 Windows 上未将 Tesseract 加入 PATH 的情况
        _tess_path = r'C:\Users\szbc124\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
        if os.path.exists(_tess_path):
            pytesseract.pytesseract.tesseract_cmd = _tess_path

        # 优先只用英文包；若未安装中文包，eng+chi_sim 会报错
        for lang in ('eng', 'eng+chi_sim'):
            try:
                text = pytesseract.image_to_string(img, lang=lang)
                if text and text.strip():
                    return text
            except Exception:
                continue
    except ImportError:
        pass
    except Exception:
        pass

    # 尝试 easyocr
    try:
        import easyocr
        reader = easyocr.Reader(['en', 'ch_sim'], gpu=False)
        results = reader.readtext(image_path)
        text = '\n'.join([r[1] for r in results])
        return text
    except ImportError:
        pass
    except Exception:
        pass

    # 尝试 paddleocr
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        result = ocr.ocr(image_path, cls=True)
        texts = []
        if result and result[0]:
            for line in result[0]:
                if line:
                    texts.append(line[1][0])
        return '\n'.join(texts)
    except ImportError:
        pass
    except Exception:
        pass

    return ""


def extract_text_via_ocr(docx_path):
    """
    对图片型 docx 的所有内嵌图片进行 OCR，合并识别结果
    按文档顺序处理，主要保留可能是文字行的图片
    """
    from PIL import Image

    image_paths, temp_dir = extract_images_from_docx(docx_path)
    if not image_paths:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return ""

    # 图片已按文档出现顺序命名（img_0000, img_0001...），保持该顺序
    all_texts = []
    total_len = 0
    max_images = 400       # 文字行型 docx 可能拆成很多小图，放宽数量
    text_threshold = 8000  # 累计识别到约 8000 字符后停止，避免过长

    for idx, img_path in enumerate(image_paths[:max_images]):
        # 优先处理整行文字/段落/表格图片；跳过图标、小碎片、超大签名页
        size = os.path.getsize(img_path)
        if size < 512 or size > 2 * 1024 * 1024:
            continue
        try:
            with Image.open(img_path) as img:
                # 文字行/段落/表格宽度通常 ≥200；过低则多为图标
                if img.width < 200 or img.height < 15 or img.height > 2000:
                    continue
        except Exception:
            continue

        try:
            text = ocr_image(img_path)
            text = text.strip() if text else ""
            if text:
                all_texts.append(text)
                total_len += len(text)
                if total_len >= text_threshold:
                    break
        except Exception:
            continue

    shutil.rmtree(temp_dir, ignore_errors=True)
    return '\n'.join(all_texts)


def extract_fields_from_ocr_text(text):
    """
    从 OCR 识别出的英文文本中用正则提取关键字段
    返回 {英文字段名: 值} 的字典
    """
    fields = {}
    if not text:
        return fields

    # 标准化：合并多余空格、统一冒号
    text = re.sub(r'\s+', ' ', text)

    # 标签映射：英文标签 → 字段名
    ocr_patterns = [
        (r'Report\s*No\.?\s*[:#]\s*([A-Z]{1,5}\d{6,15}[A-Z0-9]*)', 'Report No.'),
        (r'Applicant\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Client|Company|Product|Sample|Date|Test)|$)', 'Applicant'),
        (r'Client\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Company|Product|Sample|Date|Test)|$)', 'Client'),
        (r'Company\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Product|Sample|Date|Test)|$)', 'Company'),
        (r'Address\s*[:：]\s*([^\n]+?)(?=\s+(?:Product|Sample|Model|Date|Test|Manufacturer)|$)', 'Address'),
        (r'Product\s*Name\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Sample|Date|Test)|$)', 'Product name'),
        (r'Sample\s*Name\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Date|Test)|$)', 'Sample name'),
        (r'Sample\s*Description\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Date|Test)|$)', 'Sample Description'),
        (r'Model\s*[:：]\s*([^\n]+?)(?=\s+(?:Part|Sample|Date|Test)|$)', 'Model'),
        (r'Part\s*No\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Sample|Date|Test|Model)|$)', 'Part No.'),
        (r'Part\s*Number\s*[:：]\s*([^\n]+?)(?=\s+(?:Sample|Date|Test|Model)|$)', 'Part Number'),
        (r'Sample\s*No\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Received)|$)', 'Sample No.'),
        (r'Sample\s*Received\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Period)|$)', 'Sample Received Date'),
        (r'Date\s*of\s*Receipt\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Period)|$)', 'Date of Receipt'),
        (r'Testing\s*Period\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Report)|$)', 'Testing Period'),
        (r'Test\s*Period\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Report)|$)', 'Test Period'),
        (r'Report\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Issue)|$)', 'Report Date'),
        (r'Issue\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Report)|$)', 'Issue Date'),
        (r'Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Test|Report|No)|$)', 'Date'),
        (r'Manufacturer\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Factory|Date|Test)|$)', 'Manufacturer'),
        (r'Factory\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Date|Test)|$)', 'Factory'),
        (r'Material\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Production)|$)', 'Material'),
        (r'Production\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Production Date'),
        (r'Trade\s*Mark\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Trade Mark'),
        (r'Client\s*Ref\.?\s*Info\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Client Ref. Info.'),
        (r'Compiled\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Reviewed|Checked|Approved)|$)', 'Compiled by'),
        (r'Prepared\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Reviewed|Checked|Approved)|$)', 'Prepared by'),
        (r'Reviewed\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Checked|Approved|Date)|$)', 'Reviewed by'),
        (r'Checked\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Approved|Date)|$)', 'Checked by'),
        (r'Approved\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Report)|$)', 'Approved by'),
    ]

    for pattern, field_name in ocr_patterns:
        if field_name in fields and fields[field_name]:
            continue
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            # 清理末尾可能截断的短语
            value = re.sub(r'\s+(?:Test|Date|Report|No|Sample)\s*$', '', value, flags=re.IGNORECASE)
            if value:
                fields[field_name] = value

    # 检测要求和结论：尝试从 Test Requirement / Test Conclusion 段落提取
    req_match = re.search(
        r'Test\s*Requirement\s*[:：]?\s*([^\n]*(?:\n(?![A-Z][a-zA-Z\s]{2,20}[:：]).*)*)',
        text, re.IGNORECASE
    )
    if req_match:
        fields['Test Requirement'] = req_match.group(1).strip()

    con_match = re.search(
        r'Test\s*Conclusion\s*[:：]?\s*([^\n]*(?:\n(?![A-Z][a-zA-Z\s]{2,20}[:：]).*)*)',
        text, re.IGNORECASE
    )
    if con_match:
        fields['Test Conclusion'] = con_match.group(1).strip()

    return fields


def extract_dates_from_paragraphs_en(doc):
    """
    兜底：从段落文本中提取英文日期（样品接收日期、检测期间）。
    """
    recv_date = ""
    period = ""

    # 常见英文日期格式：2024-01-01 / Jan 01, 2024 / 01-Jan-2024
    date_pattern = r'(\d{1,2}\s*[-/]\s*[A-Za-z]{3,9}\s*[-/]\s*\d{2,4}|\d{4}\s*[-/]\s*\d{1,2}\s*[-/]\s*\d{1,2}|[A-Za-z]{3,9}\s+\d{1,2}[,.]?\s+\d{4})'
    period_pattern = r'(\d{1,2}\s*[-/]\s*[A-Za-z]{3,9}\s*[-/]\s*\d{2,4}|\d{4}\s*[-/]\s*\d{1,2}\s*[-/]\s*\d{1,2}|[A-Za-z]{3,9}\s+\d{1,2}[,.]?\s+\d{4})\s*(?:to|-)\s*(\d{1,2}\s*[-/]\s*[A-Za-z]{3,9}\s*[-/]\s*\d{2,4}|\d{4}\s*[-/]\s*\d{1,2}\s*[-/]\s*\d{1,2}|[A-Za-z]{3,9}\s+\d{1,2}[,.]?\s+\d{4})'

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        if not period and any(k in text for k in ['Testing Period', 'Test Period', 'Period', 'Date of Test']):
            m = re.search(period_pattern, text, re.IGNORECASE)
            if m:
                period = f"{m.group(1).replace(' ', '')} to {m.group(2).replace(' ', '')}"
            else:
                m = re.search(date_pattern, text, re.IGNORECASE)
                if m:
                    period = m.group(1).replace(' ', '')

        if not recv_date and any(k in text for k in ['Date of Receipt', 'Received Date', 'Date Received', 'Receipt Date']):
            m = re.search(date_pattern, text, re.IGNORECASE)
            if m:
                recv_date = m.group(1).replace(' ', '')

        if recv_date and period:
            break

    return recv_date, period


def extract_test_methods_from_paragraphs_en(doc):
    """
    兜底：从英文段落中提取检测方法（标准方法号如 IEC 62321, GB/T 30512 等）。
    """
    methods = []
    seen = set()

    method_patterns = [
        r'\b(?:IEC|ISO|ASTM|GB/T|GB|EPA|US EPA|QC/T)\s*\d[\d\-/:]*\s*(?:[-:]\s*\d{4})?',
        r'With\s+reference\s+to\s+([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'Reference\s*[:：]?\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
    ]

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        for pattern in method_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                method = match.strip()
                if not method or len(method) < 4:
                    continue
                # 简单过滤：方法号应包含标准组织前缀
                if not re.search(r'\b(IEC|ISO|ASTM|GB|EPA|QC)', method, re.IGNORECASE):
                    continue
                if method not in seen:
                    seen.add(method)
                    methods.append({
                        "检测项目": "",
                        "检测方法": method,
                        "检测仪器": ""
                    })

    return methods


# ==============================================
# 函数7：完整提取英文报告
# ==============================================
def extract_english_report(doc, filename, file_path=None):
    """英文报告完整提取，返回字段字典和检测方法列表"""

    # 图片型文档：使用 OCR 识别
    if is_image_based_document(doc):
        if VERBOSE:
            print(f"  ⚠️  检测到图片型文档，尝试 OCR 识别...")

        result = {}
        ocr_text = ""

        if file_path:
            try:
                ocr_text = extract_text_via_ocr(file_path)
                if ocr_text:
                    ocr_fields = extract_fields_from_ocr_text(ocr_text)
                    for raw_key, raw_val in ocr_fields.items():
                        std_key = EN_FIELD_MAPPING.get(raw_key)
                        if std_key and raw_val:
                            result[std_key] = raw_val
                    # 保留 OCR 原文作为备注，便于人工核对
                    if '备注' not in result or not result['备注']:
                        result['备注'] = f"OCR识别内容：{ocr_text[:500]}"
                else:
                    result['备注'] = "图片型文档，OCR 未识别到有效文本（请检查是否已安装 Tesseract/pytesseract）"
            except Exception as e:
                result['备注'] = f"图片型文档，OCR 识别失败: {str(e)[:120]}"
        else:
            result['备注'] = "图片型文档，缺少文件路径无法 OCR"

        # 若 OCR 未识别出报告编号，再尝试从文件名补
        if not result.get("报告编号"):
            m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z0-9]*)', filename, re.IGNORECASE)
            if m:
                result["报告编号"] = m.group(1)

        return result, []

    # 1. 基础字段
    raw_fields = extract_basic_info(doc)

    result = {}
    for raw_key, raw_val in raw_fields.items():
        std_key = EN_FIELD_MAPPING.get(raw_key)
        if std_key and std_key in result and raw_val:
            # 已存在则不覆盖（取首次出现的值）
            pass
        elif std_key and raw_val:
            result[std_key] = raw_val

    # 单独处理制造商地址
    if '制造商地址' in raw_fields and raw_fields['制造商地址']:
        result['制造商地址'] = raw_fields['制造商地址']

    # 2. 检测要求/结论（优先使用专题方案2逻辑，空时回退旧逻辑）
    req = extract_requirement_en(doc)
    con = extract_conclusion_en(doc, filename)

    # 旧逻辑兜底
    if not req or not con:
        old_req, old_con = extract_test_requirement_and_conclusion(doc)
        if not req and old_req:
            req = old_req
        if not con and old_con:
            con = old_con

    if req:
        result["检测要求"] = req
    if con:
        result["检测结论"] = con

    # 3. 报告编号从文件名补
    m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z]?\d*)', filename, re.IGNORECASE)
    if m and not result.get("报告编号"):
        result["报告编号"] = m.group(1)

    # 3.5 日期兜底：若基础字段未取到样品接收日期/检测期间，从段落补
    if not result.get("样品接收日期") or not result.get("检测期间"):
        recv_date, period = extract_dates_from_paragraphs_en(doc)
        if recv_date and not result.get("样品接收日期"):
            result["样品接收日期"] = recv_date
        if period and not result.get("检测期间"):
            result["检测期间"] = period

    # 4. 样品部件总数
    sample_count = extract_sample_count(doc)
    if sample_count == 0:
        sample_count = extract_sample_count_from_test_items(doc)
    result["样品部件总数"] = str(sample_count)

    # 5. 检测方法
    test_methods = extract_test_methods(doc)
    if not test_methods:
        test_methods = extract_test_methods_from_paragraphs_en(doc)
    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]

    return result, test_methods


# ==============================================
# 主函数：处理单个文档
# ==============================================
def process_word_file(input_path, output_excel):
    """
    处理单个英文 Word 检测报告，提取信息并生成 Excel

    参数:
        input_path: Word 文档路径
        output_excel: 输出 Excel 路径
    """
    from docx import Document

    print("=" * 60)
    print(f"开始处理: {os.path.basename(input_path)}")
    print("=" * 60)

    # 1. 检查文件是否存在
    if not os.path.exists(normalize_path(input_path)):
        print(f"❌ 错误：文件不存在 - {input_path}")
        return False

    # 2. 修复文档（如果需要）
    try:
        actual_file = fix_docx_if_needed(input_path)
    except Exception as e:
        print(f"❌ 错误：无法打开文档 - {e}")
        print("   请检查文件是否为有效的 Word 文档(.docx格式)")
        return False

    # 3. 打开文档
    try:
        doc = Document(actual_file)
    except ImportError:
        print("❌ 错误：未安装 python-docx 库")
        print("   请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 错误：打开文档失败 - {e}")
        return False

    if VERBOSE:
        print(f"\n📄 文档信息:")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")

    # 4. 初始化结果字典
    result = {
        "文件名": os.path.basename(input_path),
        "报告编号": "",
        "申请商": "",
        "申请商地址": "",
        "制造商地址": "",
        "样品名称": "",
        "零件号": "",
        "商标": "",
        "客户参考信息": "",
        "材质": "",
        "生产日期": "",
        "制造商": "",
        "样品编号": "",
        "样品接收日期": "",
        "检测期间": "",
        "检测要求": "",
        "检测结论": "",
        "报告日期": "",
        "样品部件总数": "",
        "编制": "",
        "审核": "",
        "批准": "",
        "备注": "",
    }

    # 5. 提取
    extracted, test_methods = extract_english_report(doc, os.path.basename(input_path), input_path)
    result.update(extracted)

    # 6. 生成 Excel
    print(f"\n💾 生成 Excel 文件...")

    col_order = [
        "报告编号", "申请商", "申请商地址", "制造商地址",
        "样品名称", "零件号", "商标",
        "客户参考信息", "材质", "生产日期", "制造商", "样品编号",
        "样品接收日期", "检测期间", "检测要求", "检测结论",
        "报告日期", "样品部件总数",
        "编制", "审核", "批准"
    ]

    actual_method_count = min(len(test_methods), MAX_TEST_METHODS)
    for i in range(1, actual_method_count + 1):
        col_order.append(f"检测项目{i}")
        col_order.append(f"检测方法{i}")
        col_order.append(f"检测仪器{i}")

    col_order.extend(["备注", "文件名"])

    df = pd.DataFrame([result])
    for col in col_order:
        if col not in df.columns:
            df[col] = ""
    df = df[col_order]

    try:
        df.to_excel(output_excel, index=False)
        print(f"✅ 成功保存: {output_excel}")
    except ImportError:
        print("❌ 错误：未安装 openpyxl 库")
        print("   请运行: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 错误：保存 Excel 失败 - {e}")
        return False

    # 统计信息
    non_empty = sum(1 for val in result.values() if val)
    print(f"\n📊 提取完成！共提取 {non_empty} 个非空字段")
    print("=" * 60)

    return True


# ==============================================
# 程序入口
# ==============================================
def main():
    print("\n" + "=" * 60)
    print("  英文检测报告信息提取工具 4")
    print("=" * 60)

    # 检查依赖
    try:
        import docx
        import pandas
    except ImportError as e:
        print(f"\n❌ 缺少依赖库: {e}")
        print("请运行以下命令安装依赖:")
        print("  pip install python-docx pandas openpyxl")
        sys.exit(1)

    # 处理文件
    success = process_word_file(INPUT_FILE, OUTPUT_EXCEL)

    if success:
        print("\n🎉 处理完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)


if __name__ == "__main__":
    main()
