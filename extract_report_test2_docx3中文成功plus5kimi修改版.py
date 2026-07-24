# -*- coding: utf-8 -*-
"""
中文检测报告信息提取工具 Plus 5
功能：从中文 RoHS/ELV/REACH 检测报告 Word 文档中提取关键信息，生成 Excel 表格
使用说明：
1. 修改下方【配置区】的文件路径和输出路径
2. 运行脚本即可生成 Excel 文件
3. 支持自动修复部分损坏的 docx 文件（NULL 关系、Bad CRC-32 等）

更新说明（Plus 5）：
- 集成《检测要求和检测结论字段专属提取方案2》的检测要求/检测结论提取逻辑，
  覆盖合并表头、多列检测要求表、2列无结论表、段落兜底等更多格式。
- 保留原有检测要求/结论逻辑作为兜底，确保不降低已有提取效果。

依赖库：python-docx, pandas, openpyxl
安装命令：pip install python-docx pandas openpyxl
"""
import os
import re
import sys
import zipfile
import tempfile
import html
import pandas as pd

# ==============================================
# 【配置区】请根据实际情况修改以下参数
# ==============================================
# 待处理的 Word 文档路径（支持 Windows 路径，如 r"D:\Kathy\PDF提取工具\test.docx"）
INPUT_FILE = r"D:\Kathy\PDF提取工具\ROHS-S19062103203001.docx"
# 输出 Excel 文件路径
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\检测报告提取信息中文docx4.xlsx"
# 最大检测方法数量（预留列数）
MAX_TEST_METHODS = 100
# 是否显示详细处理日志
VERBOSE = True


# ==============================================
# 【中文字段映射表】中文表头 → 标准含义
# ==============================================
ZH_FIELD_ALIASES = {
    # 申请商相关
    "申请商": ["申请商", "委托单位", "委托方", "客户", "客户名称"],
    "申请商地址": ["地址", "申请商地址", "委托单位地址", "客户地址"],

    # 样品相关
    "样品名称": ["样品名称", "产品名称", "样品种类", "样品描述", "产品描述", "零部件名称", "Part Name"],
    "零件号": ["零件号", "型号", "产品型号", "料号", "Part No", "Part number", "零件型号"],
    "参考零件号": ["参考零件号", "参考型号"],
    "材质": ["材质", "材料", "样品材质"],
    "车型": ["车型", "车辆型号", "车型/项目号", "项目号", "Vehicle model"],
    "生产日期": ["生产日期", "生产批号", "Production Date"],
    "供应商代码": ["供应商代码", "供应商", "Supplier"],
    "测试类型": ["测试类型", "检测类型", "测试项目"],

    # 汽车部新增关键字段
    "主机厂": ["主机厂", "汽车公司", "Automobile company", "Buyer", "买家"],
    "制造商": ["制造商", "OEM", "生产厂家", "工厂", "Manufacturer", "Factory"],

    # 样品编号相关
    "样品编号": ["样品编号", "样品号", "样本编号", "Sample No", "Sample Number", "Sample number"],

    # 日期相关
    "样品接收日期": ["样品接收日期", "接收日期", "样品收到日期", "收样日期", "收到日期", "到样日期", "收样时间"],
    "检测期间": ["检测期间", "测试周期", "检测周期", "测试期间", "测试日期", "检测日期"],
    "报告日期": ["报告日期", "日期", "签发日期", "发布日期"],

    # 人员相关
    "编制": ["编制", "编写", "编制人"],
    "审核": ["审核", "校核", "审核人"],
    "批准": ["批准", "签发", "批准人"],

    # 其他
    "备注": ["备注", "说明", "附注"],
}


def build_zh_field_mapping():
    """把别名表展开为 {别名: 标准字段名}"""
    mapping = {}
    for std_key, aliases in ZH_FIELD_ALIASES.items():
        for alias in aliases:
            mapping[alias] = std_key
    return mapping


ZH_FIELD_MAPPING = build_zh_field_mapping()


# ==============================================
# 工具函数：路径处理（解决 Windows 长路径问题）
# ==============================================
def normalize_path(file_path):
    """为 Windows 长路径添加 \\\\?\\\\ 前缀"""
    prefix = r"\\\\?\\"
    if sys.platform == 'win32' and not file_path.startswith(prefix):
        abs_path = os.path.abspath(file_path)
        if len(abs_path) >= 260:  # 超过 Windows 传统 MAX_PATH 再加前缀
            return prefix + abs_path
    return file_path


# ==============================================
# 工具函数：修复有问题的 docx 文件
# ==============================================
def fix_docx_if_needed(file_path):
    """
    检查并修复有问题的 docx 文件
    支持：NULL 关系问题、Bad CRC-32 等 zip 损坏
    返回修复后的文件路径，如果不需要修复则返回原路径
    """
    file_path = normalize_path(file_path)

    def _try_open(path):
        from docx import Document
        try:
            Document(path)
            return True
        except Exception:
            return False

    if _try_open(file_path):
        return file_path

    # 需要修复
    if VERBOSE:
        print(f"  ⚠️  检测到文档损坏，正在自动修复...")

    temp_dir = tempfile.gettempdir()
    prefix = r"\\\\?\\"
    base_name = os.path.basename(file_path.replace(prefix, ''))
    fixed_path = os.path.join(temp_dir, f"fixed_{base_name}")

    try:
        with zipfile.ZipFile(file_path, 'r') as zin:
            # 第一轮：找出所有损坏的媒体文件
            skipped_media = set()
            for item in zin.infolist():
                if not item.filename.startswith('word/media/'):
                    continue
                try:
                    zin.read(item.filename)
                except zipfile.BadZipFile:
                    skipped_media.add(item.filename)
                    if VERBOSE:
                        print(f"    跳过损坏的图片: {item.filename}")

            with zipfile.ZipFile(fixed_path, 'w') as zout:
                for item in zin.infolist():
                    # 跳过已损坏的媒体文件
                    if item.filename in skipped_media:
                        continue

                    # 修复 NULL 关系，并删除指向损坏媒体的关系
                    if item.filename == 'word/_rels/document.xml.rels':
                        try:
                            content = zin.read(item.filename).decode('utf-8')
                            content = re.sub(r'<Relationship[^>]*Target="[^"]*NULL[^"]*"[^>]*/>', '', content)
                            for media in skipped_media:
                                # document.xml.rels 中 Target 为相对 word/ 目录，如 media/image5.png
                                target = media.replace('word/', '')
                                content = re.sub(r'<Relationship[^>]*Target="[^"]*' + re.escape(target) + r'"[^>]*/>', '', content)
                            zout.writestr(item, content)
                        except zipfile.BadZipFile:
                            # 如果连关系文件都损坏，跳过（极少见）
                            continue
                        continue

                    # 尝试复制其他文件
                    try:
                        data = zin.read(item.filename)
                        zout.writestr(item, data)
                    except zipfile.BadZipFile:
                        raise

        if _try_open(fixed_path):
            if VERBOSE:
                print(f"  ✅ 文档修复成功")
            return fixed_path
        else:
            raise Exception("修复后仍无法打开文档")
    except Exception as fix_error:
        if VERBOSE:
            print(f"  ❌ 文档修复失败: {fix_error}")
        raise


# ==============================================
# 工具函数：获取单元格完整文本（包括 SDT 内容控件）
# ==============================================
def get_cell_text(cell):
    """
    获取单元格的完整文本，包括 SDT（结构化文档标签/内容控件）中的内容
    解决 python-docx 无法读取内容控件文本的问题
    """
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    result = html.unescape(result)
    return result.strip()


def get_row_cells_text_fast(row):
    """
    获取一行中所有单元格的文本，包括被 SDT（结构化文档标签）包裹的单元格
    解决 python-docx 无法识别 SDT 包裹的 tc 单元格的问题
    处理 w:gridSpan 水平合并单元格，用空字符串填充被合并的列，
    确保返回列表长度与表格实际网格列数对齐，避免合并表头列索引偏移。
    """
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cell_text = html.unescape(cell_text)
        cells_text.append(cell_text.strip())
        # 处理 gridSpan：提取合并列数，超出1的部分用空字符串填充以对齐列位置
        gs_match = re.search(r'<w:gridSpan[^>]*w:val\s*=\s*"(\d+)"', tc_xml)
        if gs_match:
            span = int(gs_match.group(1))
            for _ in range(span - 1):
                cells_text.append('')
    return cells_text


# 兼容总提取代码的别名
get_row_cells_text = get_row_cells_text_fast


# ==============================================
# 工具函数：文本预处理
# ==============================================
def normalize_text(text):
    """清理文本：去首尾空、统一换行/空格"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def normalize_key(key):
    """清理字段名：去冒号、去首尾空、统一空格"""
    if not key:
        return ""
    key = normalize_text(key)
    key = key.rstrip('：:').strip()
    return key


def normalize_key_bilingual(key):
    """
    针对汽车部报告中双语字段名做归一化。
    同时覆盖以下常见格式：
      - 中文(English):     申请商(Applicant) :
      - 中文English:       样品接收日期Sample Received Date：
      - English中文:       Sample Received Date 样品接收日期：
      - 中文 : English:    样品名称 : Sample name
    去掉英文部分、尾部中英文冒号/空格，保留中文键名。
    若键名纯为英文，则保留原值。
    """
    if not key:
        return ""
    # 基础清洗
    key = normalize_text(key)
    # 去掉英文括号及其中的内容（支持嵌套一层括号）
    key = re.sub(r'\s*\([A-Za-z0-9\s&/\.\-_,#()]+\)\s*', '', key)
    # 去掉尾部中英文冒号
    key = key.rstrip('：:').strip()
    # 去掉末尾连续的非中文字段（英文 + 数字 + 常见标点和空格）
    suffix_match = re.search(r'[A-Za-z0-9\s&/\.\-_,#]+$', key)
    if suffix_match and len(suffix_match.group().strip()) >= 2:
        candidate = key[:suffix_match.start()].strip()
        if candidate:
            key = candidate
    # 去掉开头连续的非中文字段
    prefix_match = re.match(r'^[A-Za-z0-9\s&/\.\-_,#]+', key)
    if prefix_match and len(prefix_match.group().strip()) >= 2:
        candidate = key[prefix_match.end():].strip()
        if candidate:
            key = candidate
    return key.rstrip('：:').strip()


# ==============================================
# 【专题方案2】检测要求/结论辅助函数
# ==============================================
def contains_pass(text):
    """是否包含符合语义"""
    if not text:
        return False
    t = text.lower()
    # 先排除明确的不符合语义（避免 "不符合" 被 "符合" 子串误匹配）
    if any(k in t for k in ['不符合', '不合格', '不通过', '不达标', '不满足']):
        return False
    return any(k in t for k in ['符合', '合格', '通过', '达标', '满足',
                                 'pass', 'comply', 'complies', 'compliant',
                                 'conform', 'conforms', 'conformed',
                                 '阴性', '未检出', '不得检出', 'not detected', 'ok'])


def contains_fail(text):
    """是否包含不符合语义"""
    if not text:
        return False
    t = text.lower()
    # 简单子串匹配
    if any(k in t for k in ['不符合', '不合格', '不通过', '不达标', '不满足', '超标', '超出',
                             'fail', 'failed', 'failing', 'non-conform', 'nonconform',
                             '阳性', 'ol', 'f']):
        return True
    # '检出' 用正则排除 '未检出'（(?<!未)检出 = 前面不是"未"的"检出"）
    # 并排除 '不得检出'（限值要求，非实际检出）、'检出限'（方法学术语）
    if re.search(r'(?<!未)检出', t) and '检出限' not in t and '不得检出' not in t:
        return True
    return False


def is_nd(text):
    """是否为 N.D.（未检出）"""
    if not text:
        return False
    return text.strip().lower() in ['nd', 'n.d.', 'n.d', '未检出', 'not detected']


def is_numeric_result(val):
    """是否为具体数值结果（按用户要求视为不符合）"""
    if not val or val in ['/', '-', '—', 'P', 'X', '']:
        return False
    if is_nd(val) or contains_pass(val) or contains_fail(val):
        return False
    return bool(re.search(r'\d', val))


def is_conclusion_like(text):
    """判断文本是否像结论列内容（短结论词）"""
    if not text:
        return False
    t = text.strip().lower()
    if len(t) > 15:
        return False
    conclusion_keywords = [
        '结论', 'conclusion',
        '符合', '不符合', '合格', '不合格', '通过', '不通过',
        'pass', 'fail', 'failed', 'passing', 'failing',
        'conform', 'conforms', 'conformed', 'non-conform', 'nonconform',
        'nd', 'n.d.', 'n.d', '未检出', '阳性', '阴性'
    ]
    return any(k in t for k in conclusion_keywords)


def classify_text(text):
    """对整段文本判断：pass / fail / partial / unknown"""
    if not text:
        return 'unknown'
    has_fail = contains_fail(text)
    has_pass = contains_pass(text)
    if has_fail and not has_pass:
        return 'fail'
    if has_fail and has_pass:
        return 'partial'
    if has_pass and not has_fail:
        return 'pass'
    return 'unknown'


def format_fail_details(items):
    """格式化不符合项列表"""
    seen = set()
    details = []
    for item, val in items:
        key = f"{item}|{val}"
        if key in seen:
            continue
        seen.add(key)
        if item:
            if val in ['不符合', 'Fail', 'failed', '不合格', '不通过']:
                details.append(item)
            else:
                details.append(f"{item}({val})")
        else:
            details.append(val)
    return details


def summarize_conclusion_items(items, all_fail_short=False):
    """
    综合多项结论：
    - items: [(项目, 值), ...]
    - all_fail_short=True 时，全部不符合只输出"不符合"
    """
    if not items:
        return ""

    pass_items = [(i, v) for i, v in items
                  if (contains_pass(v) and not contains_fail(v)) or is_nd(v)]
    fail_items = [(i, v) for i, v in items
                  if contains_fail(v)]

    if fail_items and not pass_items:
        if all_fail_short:
            return "不符合"
        details = format_fail_details(fail_items)
        if details:
            return "不符合，不符合的内容有：" + "、".join(details)
        return "不符合"

    if fail_items:
        details = format_fail_details(fail_items)
        if details:
            return "部分项不符合，不符合的内容有：" + "、".join(details)
        return "部分项不符合"

    if pass_items:
        return "符合"

    # 只有未知值，保留原文
    return " / ".join(set(v for _, v in items))


# ==============================================
# 函数1：判断文档类型（长文档/短文档）
# ==============================================
def detect_doc_type(doc):
    """
    判断中文文档类型：
    - 长文档：首页有检测方法总结表格（两列，方法编号+名称，如 QC/T、GB/T、IEC 开头）
    - 短文档：首页没有检测方法总结，检测方法在后面的详细表格中
    """
    # 方法1：通过表格数量判断（长文档表格很多）
    if len(doc.tables) > 30:
        return "long"

    # 方法2：检查前 5 个表格中是否有检测方法总结
    for table in doc.tables[:5]:
        if len(table.rows) >= 2 and len(table.rows[0].cells) == 2:
            first_cell = get_cell_text(table.rows[0].cells[0])
            if re.match(r'^(QC/T|GB/T|IEC|GB|ISO|ASTM|EPA|US EPA)\s*\d', first_cell, re.IGNORECASE):
                return "long"

    return "short"


# ==============================================
# 表格类型判断辅助函数
# ==============================================
def is_result_table(first_row):
    """判断是否为检测结果表格（应跳过）"""
    text = ' '.join(first_row)
    has_item = any(kw in text for kw in ["检测项目", "Test Item", "测试项目", "Test item(s)"])
    has_result = any(kw in text for kw in ["限值", "结果", "方法检出限", "Limit", "Result", "MDL"])
    return has_item and has_result


def is_sample_composition_table(table):
    """判断是否为样品组成/描述表格（第一列是数字序号）"""
    if len(table.rows) < 3:
        return False

    checked = 0
    digit_count = 0
    for row in table.rows[:5]:
        cells = get_row_cells_text_fast(row)
        if not cells:
            continue
        checked += 1
        first_cell = cells[0].strip()
        if re.match(r'^(\d+)(-\d+)?$', first_cell):
            digit_count += 1

    return checked >= 3 and digit_count >= 3


def is_instrument_table(first_row):
    """判断是否为仪器/设备表格"""
    text = ' '.join(first_row)
    instrument_keywords = ["仪器名称", "仪器", "设备名称", "设备", "Instrument", "Equipment"]
    return any(kw in text for kw in instrument_keywords)


def find_column_index(first_row, keywords):
    """根据关键词列表找到列索引"""
    for idx, cell in enumerate(first_row):
        for kw in keywords:
            if kw in cell:
                return idx
    return None


# ==============================================
# 函数2：提取基本信息
# ==============================================
def extract_basic_info(doc):
    """从两列、四列表格中提取中文报告基础字段"""
    fields = {}
    manufacturer_found = False

    for table_idx, table in enumerate(doc.tables[:20], 1):
        if len(table.rows) < 1:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' '.join(first_row)
        cell_count = len(first_row)

        # 跳过检测结果表格
        if is_result_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 检测结果表格，跳过")
            continue

        # 跳过样品组成/样品描述表格
        if is_sample_composition_table(table):
            if VERBOSE:
                print(f"  表格 {table_idx}: 样品组成表格，跳过")
            continue

        # 跳过仪器明细表（单独处理）
        if is_instrument_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 仪器表格，单独处理")
            continue

        # 处理两列表格（键值对）
        if cell_count == 2:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]

                # 处理只有 1 个单元格的行（可能是合并单元格的标签行）
                if len(cells) == 1:
                    key = normalize_key_bilingual(cells[0])
                    if key and len(key) < 60:
                        fields[key] = ""
                    continue

                if len(cells) == 2:
                    key = normalize_key_bilingual(cells[0])
                    val = cells[1].strip()

                    if not key or len(key) >= 60:
                        continue

                    # 跳过说明性文字
                    if '以下的检测样品' in key or '由客户提供' in key:
                        continue

                    # 特殊处理：遇到制造商/工厂后，下一个地址是制造商地址
                    if key in ["制造商", "工厂", "生产厂家"]:
                        manufacturer_found = True
                        if val and val != '/':
                            fields[key] = val
                        continue

                    # 特殊处理地址：区分申请商地址和制造商地址
                    if key == "地址":
                        if not manufacturer_found:
                            if val and val != '/':
                                fields["申请商地址"] = val
                        else:
                            if val and val != '/' and '制造商地址' not in fields:
                                fields["制造商地址"] = val
                        continue

                    # 忽略空值占位符
                    if val == '/':
                        val = ""

                    if val:
                        fields[key] = val

        # 处理四列表格（编制/审核/批准/日期等）
        elif cell_count == 4:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) == 4:
                    for i in range(0, 4, 2):
                        key = normalize_key_bilingual(cells[i])
                        val = cells[i + 1].strip()
                        if key and len(key) < 40:
                            if val and val != '/':
                                fields[key] = val

    return fields


# ==============================================
# 函数3：提取检测要求和检测结论（完整原文）
# ==============================================
def _split_requirement_conclusion(text):
    """
    对“检测要求及结论”混合文本进行拆分。
    把包含结论语义（符合/通过/合格/满足/达标）的最后一个分句作为检测结论，
    其余部分作为检测要求。
    """
    text = normalize_text(text)
    if not text:
        return "", ""

    # 按中文/英文句读拆分，保留分隔符
    parts = re.split(r'([。；，,!])', text)
    sentences = []
    cur = ""
    for p in parts:
        cur += p
        if p in "。；，,!":
            s = cur.strip()
            if s:
                sentences.append(s)
            cur = ""
    if cur.strip():
        sentences.append(cur.strip())

    conclusion_keywords = ['符合', '通过', '合格', '满足', '达标',
                           'comply', 'complies', 'pass', 'passed', 'conform']
    con_idx = None
    for i in range(len(sentences) - 1, -1, -1):
        low = sentences[i].lower()
        if any(k in low or k in sentences[i] for k in conclusion_keywords):
            con_idx = i
            break

    if con_idx is not None:
        requirement = "".join(sentences[:con_idx]).strip()
        conclusion = sentences[con_idx].strip()
        if not requirement:
            requirement = text
        return requirement, conclusion

    return text, ""


def extract_test_requirement_and_conclusion(doc):
    """提取检测要求和检测结论的完整原文"""
    test_requirement = ""
    test_conclusion = ""

    for table_idx, table in enumerate(doc.tables, 1):
        if test_requirement and test_conclusion:
            break

        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c) for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)

        # 情况1：2行2列表格（检测要求和结论在同一行）
        if (len(table.rows[0].cells) == 2 and
                '检测要求' in first_row_text and
                '结论' in first_row_text):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到检测要求/结论表格（2行2列）")
            if len(table.rows) >= 2:
                test_requirement = get_cell_text(table.rows[1].cells[0])
                test_conclusion = get_cell_text(table.rows[1].cells[1])
            break

        # 情况2/3：1列表格
        elif len(table.rows[0].cells) == 1:
            header = normalize_text(first_row[0])

            # 2a. 表头同时含“检测要求”和“结论”（如“检测要求及结论”）
            if '检测要求' in header and '结论' in header:
                if VERBOSE:
                    print(f"  表格 {table_idx}: 找到检测要求及结论混合表格")
                if len(table.rows) >= 2:
                    full = get_cell_text(table.rows[1].cells[0])
                    req, con = _split_requirement_conclusion(full)
                    if req and not test_requirement:
                        test_requirement = req
                    if con and not test_conclusion:
                        test_conclusion = con

            # 2b. 表头仅含“检测要求”，可能在后续行出现“检测结论:”标签
            elif '检测要求' in header:
                if not test_requirement and len(table.rows) >= 2:
                    test_requirement = get_cell_text(table.rows[1].cells[0])
                # 扫描后续标签行
                if not test_conclusion:
                    for r in range(1, len(table.rows)):
                        txt = normalize_text(get_cell_text(table.rows[r].cells[0]))
                        if '检测结论' in txt and r + 1 < len(table.rows):
                            test_conclusion = get_cell_text(table.rows[r + 1].cells[0])
                            if VERBOSE:
                                print(f"  表格 {table_idx}: 在单列表格后续行找到检测结论")
                            break

            # 2c. 表头含“检测结论”
            elif '结论' in header and '检测' in header:
                if not test_conclusion:
                    if VERBOSE:
                        print(f"  表格 {table_idx}: 找到检测结论表格（独立表格）")
                    if len(table.rows) >= 2:
                        test_conclusion = get_cell_text(table.rows[1].cells[0])

    return test_requirement, test_conclusion


def extract_conclusion_from_result_column_zh(doc):
    """
    兜底：从检测结果表格的“结论”列提取结论（如“符合”）。
    对多行结论去重后用“ / ”拼接。
    """
    conclusions = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        # 跳过“检测要求/结论”标签表
        if len(first_row) == 1 and ('检测要求' in first_row[0] or '检测结论' in first_row[0]):
            continue

        first_row_text = ' '.join(first_row)
        # 只考虑结果/结论相关表格
        if not any(k in first_row_text for k in ['结论', 'Conclusion', '符合', 'Pass', '结果', 'Result']):
            continue

        # 跳过说明/备注/注释表：表头含 说明/备注/注释/Note/Remark/Annotation 且不含检测项目列
        if any(k in first_row_text for k in ['说明', '备注', '注释', 'Note', 'Remark', 'Annotation']):
            has_item_col = any(k in first_row_text for k in ['检测项目', 'Test Item',
                                                              'Item', '测试项目'])
            if not has_item_col:
                continue

        con_col = find_column_index(first_row, ['结论', 'Conclusion', '符合', 'Pass',
                                                 '结果', 'Result'])
        if con_col is None:
            continue

        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if con_col >= len(cells):
                continue
            val = cells[con_col].strip()
            if not val:
                continue
            # 跳过表头残留、检测项目名等
            if val in first_row:
                continue
            if val in ['结论', 'Conclusion', '结果', 'Result']:
                continue
            if val not in seen:
                seen.add(val)
                conclusions.append(val)

    # 兜底 2：结果表中直接有一行首列为”结论”，其右侧单元格为”符合/Pass”等
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text2 = ' '.join(first_row)
        if not any(k in first_row_text2 for k in ['结论', 'Conclusion', '符合', 'Pass', '结果', 'Result']):
            continue
        # 跳过说明/备注/注释表
        if any(k in first_row_text2 for k in ['说明', '备注', '注释', 'Note', 'Remark', 'Annotation']):
            has_item_col = any(k in first_row_text2 for k in ['检测项目', 'Test Item',
                                                               'Item', '测试项目'])
            if not has_item_col:
                continue
        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if not cells:
                continue
            if cells[0].strip() in ['结论', 'Conclusion']:
                for cell in cells[1:]:
                    val = cell.strip()
                    if val and val not in ['结论', 'Conclusion', '/', '-', '—', '结果', 'Result']:
                        if val not in seen:
                            seen.add(val)
                            conclusions.append(val)

    return ' / '.join(conclusions) if conclusions else ""


# ==============================================
# 【专题方案2】检测要求提取
# ==============================================
def extract_requirement_from_paragraphs_zh(doc):
    """
    从段落文本中提取检测要求（兜底逻辑）
    适用场景：报告没有独立的检测要求表格，检测要求直接写在正文段落中。
    """
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paras:
        return ""

    stop_patterns = [
        '检测结果', 'Test Result',
        '检测仪器', 'Test Instrument',
        '检测流程', 'Test Flow',
        '样品描述', 'Sample Description',
        '样品照片', 'Sample photo',
        '判定标准', 'Decision criteria',
        '****报告结束', 'End of Report'
    ]

    method_stop_patterns = [
        '检测结果', 'Test Result',
        '检测仪器', 'Test Instrument',
        '样品描述', 'Sample Description',
        '样品照片', 'Sample photo',
        '****报告结束', 'End of Report'
    ]

    # 策略1：找"检测要求"标题后的内容
    req_lines = []
    capture = False
    for text in paras:
        if any(k in text for k in ['检测要求', 'Test Requirement', '检测依据', '测试依据', '测试要求', '标准要求']):
            content = text
            for k in ['检测要求', 'Test Requirement', '检测依据', '测试依据', '测试要求', '标准要求', '：', ':', '；', ';']:
                content = content.replace(k, '')
            content = content.strip()
            if content and len(content) > 5 and '请参考下页' not in content and 'Please refer' not in content:
                req_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if len(text) < 3:
                continue
            req_lines.append(text)

    if req_lines:
        return '\n'.join(req_lines)

    # 策略2：找"检测方法"/"检测流程"标题后的内容
    method_like_patterns = ['检测方法', 'Test Method', '检测流程', 'Test Flow']
    method_lines = []
    capture = False
    for text in paras:
        if any(p in text for p in method_like_patterns):
            content = text
            for k in ['检测方法', 'Test Method', '检测流程', 'Test Flow', '：', ':', '；', ';']:
                content = content.replace(k, '')
            content = content.strip().lower()
            meaningless = {'method', 'methods', 'test', 'tests', 'test method', 'flow', 'flows'}
            if content and len(content) > 5 and content not in meaningless \
                    and '请参考下页' not in text and 'Please refer' not in text:
                method_lines.append(content)
            capture = True
            continue

        if capture:
            if any(stop in text for stop in method_stop_patterns):
                capture = False
                continue
            if len(text) < 3:
                continue
            method_lines.append(text)

    if method_lines:
        return '\n'.join(method_lines)

    # 策略3：找"检测结果"标题下的编号列表
    result_lines = []
    capture = False
    for text in paras:
        if '检测结果' in text or 'Test Result' in text:
            capture = True
            continue

        if capture:
            if any(stop in text for stop in stop_patterns):
                break
            if re.match(r'^[\d一二三四五六七八九十]+[\.\、\．]', text) or text.startswith(('•', '-', '*')):
                result_lines.append(text)

    if result_lines:
        return '\n'.join(result_lines)

    return ""


def _is_requirement_text(text):
    """判断文本是否像检测要求描述（而非标准编号或检测项目名）"""
    if not text:
        return False
    t = text.lower()
    keywords = ['根据客户要求', '依据', '根据', '参考', '按照',
                '对所提交样品', '对所送样品', '进行以下项目',
                '对所送样', '按客户要求', 'to determine',
                'in accordance with', 'as specified by']
    if any(k in t for k in keywords):
        return True
    # 长文本且包含标准方法号
    if len(text) > 30 and re.search(r'(GB/T|GB|IEC|ISO|ASTM|EPA|US\s*EPA|QC/T|QJ/GAC|Q/ALKS)\s*\d', text):
        return True
    return False


def _is_serial_column(table, col_idx, sample_rows=5):
    """判断指定列是否主要为数字/中文序号列"""
    serial_count = 0
    total = 0
    for row in table.rows[1:sample_rows + 1]:
        cells = get_row_cells_text_fast(row)
        if col_idx < len(cells):
            val = cells[col_idx].strip()
            if val:
                total += 1
                if re.match(r'^[\d一二三四五六七八九十]+[\.\、\．)\s]*$', val):
                    serial_count += 1
    return total > 0 and serial_count / total >= 0.5


def _classify_requirement_table_columns(first_row):
    """根据表头给每列打角色标签"""
    roles = []
    for cell in first_row:
        cs = cell.strip()
        if any(k in cs for k in ['检测要求', '测试要求', 'Test Requirement', 'Requirement']):
            roles.append('req')
        elif any(k in cs for k in ['结论', 'Conclusion']):
            roles.append('con')
        elif any(k in cs for k in ['结果', 'Result']):
            roles.append('res')
        else:
            roles.append('other')
    return roles


def _determine_requirement_cols(table, first_row, col_roles):
    """
    在通用多列检测要求表中，确定真正的要求文本列。
    会排除结论列、结果列和序号列，并在剩余列中选择要求语义最强的列。
    """
    candidate_cols = [i for i, r in enumerate(col_roles) if r in ('req', 'other')]

    # 排除明显是序号列的
    candidate_cols = [c for c in candidate_cols if not _is_serial_column(table, c)]

    if not candidate_cols:
        return []

    best_col = None
    best_score = -1
    for col in candidate_cols:
        req_text_count = 0
        total_len = 0
        std_only_count = 0
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if col >= len(cells):
                continue
            val = cells[col].strip()
            if not val:
                continue
            total_len += len(val)
            if _is_requirement_text(val):
                req_text_count += 1
            # 纯标准编号行（短且匹配标准号开头）
            if re.search(r'^(GB/T|GB|IEC|ISO|ASTM|EPA|US\s*EPA|QC/T|QJ/GAC|Q/ALKS)\s*\d', val) and len(val) < 60:
                std_only_count += 1

        # 评分：要求文本命中加分，总长度加分，纯标准编号减分
        score = req_text_count * 100 + total_len - std_only_count * 50
        if score > best_score:
            best_score = score
            best_col = col

    return [best_col] if best_col is not None else []


def extract_requirement_zh(doc):
    """
    提取检测要求（基于专题方案2 v2.0）
    覆盖合并表头、单列表格、2列混合表、通用多列表格、段落兜底。
    """
    all_lines = []

    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        first_row_fast = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row_fast)
        fast_col_count = len(first_row_fast)

        has_req_keyword = any('检测要求' in cell or '测试要求' in cell or 'Test Requirement' in cell
                             for cell in first_row_fast)

        # A. 合并表头预检
        true_col_count = len(table.rows[0].cells)
        if fast_col_count < true_col_count and has_req_keyword and len(table.rows) >= 2:
            max_cols = max(len(row.cells) for row in table.rows[1:])

            col_values = [[] for _ in range(max_cols)]
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                for i, val in enumerate(row_cells):
                    if i < max_cols and val:
                        col_values[i].append(val)

            skip_cols = set()
            for i, vals in enumerate(col_values):
                if not vals:
                    continue
                con_count = sum(1 for v in vals if is_conclusion_like(v))
                if con_count / len(vals) >= 0.6:
                    skip_cols.add(i)

            lines = []
            for row in table.rows[1:]:
                row_cells = [get_cell_text(c).strip() for c in row.cells]
                parts = []
                for i, val in enumerate(row_cells):
                    if i < max_cols and i not in skip_cols and val:
                        parts.append(val)
                if parts:
                    lines.append(' '.join(parts))
            if lines:
                all_lines.extend(lines)
                continue

        # B. 单列表格
        if fast_col_count == 1:
            header = first_row_fast[0]
            if '检测要求' in header or '测试要求' in header or 'Test Requirement' in header:
                lines = []
                for row in table.rows[1:]:
                    cell_text = get_cell_text(row.cells[0]).strip()
                    if cell_text:
                        lines.append(cell_text)
                if lines:
                    all_lines.extend(lines)
                    continue

        # C. 2行2列混合表（检测要求 + 结论）
        if fast_col_count == 2:
            has_req = ('检测要求' in first_row_text or '测试要求' in first_row_text or 'Test Requirement' in first_row_text)
            has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text)
            if has_req and has_con and len(table.rows) >= 2:
                if '检测要求' in first_row_fast[1] or '测试要求' in first_row_fast[1] or 'Test Requirement' in first_row_fast[1]:
                    req_col = 1
                else:
                    req_col = 0

                lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    if req_col < len(cells):
                        val = cells[req_col].strip()
                        if val:
                            lines.append(val)
                if lines:
                    all_lines.extend(lines)
                    continue

        # D. 通用多列检测要求表处理（列角色识别 + 语义过滤）
        if has_req_keyword and len(table.rows) >= 2:
            col_roles = _classify_requirement_table_columns(first_row_fast)
            extract_cols = _determine_requirement_cols(table, first_row_fast, col_roles)

            if extract_cols:
                # 优先提取明确要求语义的长文本
                req_lines = []
                for row in table.rows[1:]:
                    cells = get_row_cells_text_fast(row)
                    for col in extract_cols:
                        if col < len(cells):
                            val = cells[col].strip()
                            if val and _is_requirement_text(val):
                                req_lines.append(val)

                # 若未命中要求语义，兜底提取候选列全部非空文本
                if not req_lines:
                    for row in table.rows[1:]:
                        cells = get_row_cells_text_fast(row)
                        row_parts = []
                        for col in extract_cols:
                            if col < len(cells):
                                val = cells[col].strip()
                                if val:
                                    row_parts.append(val)
                        if row_parts:
                            req_lines.append(' '.join(row_parts))

                if req_lines:
                    all_lines.extend(req_lines)
                    continue

    if all_lines:
        seen = set()
        unique_lines = []
        for line in all_lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)
        return '\n'.join(unique_lines)

    return extract_requirement_from_paragraphs_zh(doc)


# ==============================================
# 【专题方案2】检测结论提取
# ==============================================
def _resolve_result_value(val, cells, result_col):
    """
    判断结果列中单个单元格值的结论语义（用于无"结论"列时的兜底判断）。
    返回: ('pass',) / ('fail', detail) / None（跳过）
    """
    val = val.strip()
    if not val or val in ['/', '-', '—', 'P', 'X']:
        return None

    # 1. N.D. → pass
    if is_nd(val):
        return ('pass',)

    # 2. 明确 Fail 语义（且不含 Pass 词） → fail
    if contains_fail(val) and not contains_pass(val):
        return ('fail', val)

    # 3. 明确 Pass 语义（且不含 Fail 词） → pass
    if contains_pass(val) and not contains_fail(val):
        return ('pass',)

    # 4. 同时含 Pass 和 Fail 语义 → fail（偏保守）
    if contains_pass(val) and contains_fail(val):
        return ('fail', val)

    # 5. 含 "/" 的值 → 向左读同行单元格获取上下文
    if '/' in val:
        for i in range(result_col - 1, -1, -1):
            if i >= len(cells):
                continue
            left_val = cells[i].strip()
            if not left_val or left_val in ['/', '-', '—']:
                continue
            if contains_fail(left_val):
                return ('fail', val)
            if contains_pass(left_val) or is_nd(left_val):
                return ('pass',)
        # 左边无任何语义信息 → 保守判为无法判断
        return ('unknown',)

    # 6. 兜底：无法判断 → 需人工查看
    return ('unknown',)


def extract_conclusion_from_filename_zh(filename):
    """从文件名判断结论"""
    name = os.path.splitext(filename)[0]
    if '不合格' in name or '不符合' in name:
        return '不符合'
    if '合格' in name or '符合' in name:
        return '符合'
    return ""


def extract_conclusion_from_independent_table_zh(doc):
    """从独立检测结论表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 1:
            continue

        header = first_row[0]
        if '检测结论' not in header and '测试结论' not in header and 'Test Conclusion' not in header:
            continue

        lines = []
        for row in table.rows[1:]:
            val = get_cell_text(row.cells[0]).strip()
            if val:
                lines.append(val)

        full_text = '\n'.join(lines)
        if not full_text:
            return ""

        cat = classify_text(full_text)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)
        if cat == 'partial':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)

        return full_text

    return ""


def extract_conclusion_from_mixed_table_zh(doc):
    """从2行2列检测要求/结论混合表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 2:
            continue

        first_row_text = ' | '.join(first_row)
        has_req = ('检测要求' in first_row_text or '测试要求' in first_row_text or 'Test Requirement' in first_row_text)
        has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text or
                   '结果' in first_row_text or 'Result' in first_row_text)
        if not (has_req and has_con):
            continue

        con_col = 1
        if '结论' in first_row[0] or 'Conclusion' in first_row[0]:
            con_col = 0

        items = []
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if con_col < len(cells):
                val = cells[con_col].strip()
                if val:
                    items.append(("", val))

        if items:
            summary = summarize_conclusion_items(items, all_fail_short=True)
            if summary:
                return summary
            return "\n".join([v for _, v in items])

    return ""


def _is_result_table(table):
    """
    判断表格是否为检测结果表。
    不仅看首行表头，还扫描前 10 行，支持跨行合并表头。
    """
    result_keywords = ['结果', 'Result', '结论', 'Conclusion', '限值', 'Limit']
    for r_idx in range(min(10, len(table.rows))):
        row_text = ' | '.join(get_row_cells_text_fast(table.rows[r_idx]))
        if any(k in row_text for k in result_keywords):
            return True
    return False


def _locate_conclusion_column(table, first_row, max_scan_rows=10):
    """
    在结果表中定位结论列。
    1. 先在首行表头中找明确含“结论/Conclusion”的列；
    2. 若未找到，扫描前 max_scan_rows 行，按每列出现结论关键词的频率定位；
    3. 兜底：找“结果/Result”列。
    返回: (列索引, 是否由内容定位)
    """
    # 策略1：表头明确含结论
    for idx, cell in enumerate(first_row):
        if any(k in cell for k in ['结论', 'Conclusion']):
            return idx, False

    # 策略2：扫描前N行，统计结论关键词出现频率
    col_scores = [0] * len(first_row)
    conclusion_keywords = ['结论', 'Conclusion', '符合', 'Pass', '合格', '不合格', 'Fail', '不符合', '未检出']
    for r_idx in range(1, min(max_scan_rows + 1, len(table.rows))):
        cells = get_row_cells_text_fast(table.rows[r_idx])
        if len(cells) != len(first_row):
            continue
        for c_idx, val in enumerate(cells):
            v = val.strip()
            if any(k in v for k in conclusion_keywords):
                # 短结论词权重更高
                weight = 2 if len(v) <= 12 else 1
                col_scores[c_idx] += weight

    if max(col_scores) > 0:
        return col_scores.index(max(col_scores)), True

    # 策略3：兜底找“结果/Result”列
    first_row_text = ' | '.join(first_row)
    for idx, cell in enumerate(first_row):
        if any(k in cell for k in ['结果', 'Result']) and \
           not any(k in first_row_text for k in ['检测结果', 'Test Result', '要求', 'Requirement']):
            return idx, False

    return None, False


def _read_conclusion_value(row, con_col, search_range=2):
    """
    读取一行中的结论值。
    若定位列无效，则在右侧相邻列及行尾搜索结论关键词。
    """
    cells = get_row_cells_text_fast(row)
    if not cells:
        return ""

    conclusion_keywords = ['符合', 'Pass', '合格', '不合格', 'Fail', '不符合', '未检出', 'N.D.', 'ND']

    # 优先读取定位列
    if 0 <= con_col < len(cells):
        val = cells[con_col].strip()
        if val and val not in ['/', '-', '—', '结论', 'Conclusion', '结果', 'Result']:
            if any(k in val for k in conclusion_keywords) or len(val) <= 15:
                return val

    # 向右相邻列搜索
    start_col = min(con_col + 1, len(cells) - 1)
    for c in range(start_col, min(len(cells), start_col + search_range)):
        val = cells[c].strip()
        if val and val not in ['/', '-', '—']:
            if any(k in val for k in conclusion_keywords):
                return val

    # 兜底：行尾最后两列
    for c in range(max(0, len(cells) - 2), len(cells)):
        val = cells[c].strip()
        if val and val not in ['/', '-', '—', '结论', 'Conclusion', '结果', 'Result']:
            if any(k in val for k in conclusion_keywords):
                return val

    return ""


def extract_conclusion_from_result_tables_zh(doc):
    """从检测结果表结论列/行综合判断；支持合并单元格、跨行表头导致的列偏移。"""
    all_items = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row)
        if len(first_row) == 1:
            continue

        # 判断是否为结果表（支持跨行合并表头）
        if not _is_result_table(table):
            continue

        # 跳过说明/备注/注释表：表头含 说明/备注/注释/Note/Remark/Annotation 且不含检测项目列
        if any(k in first_row_text for k in ['说明', '备注', '注释', 'Note', 'Remark', 'Annotation']):
            has_item_col = any(k in first_row_text for k in ['检测项目', 'Test Item',
                                                              'Item', '测试项目'])
            if not has_item_col:
                continue

        # 定位结论列
        con_col, _ = _locate_conclusion_column(table, first_row)

        is_result_fallback = False
        if con_col is None:
            # 兜底：找“结果/Result”列
            for idx, cell in enumerate(first_row):
                if any(k in cell for k in ['结果', 'Result']) and \
                   not any(k in first_row_text for k in ['检测结果', 'Test Result',
                                                          '要求', 'Requirement']):
                    con_col = idx
                    is_result_fallback = True
                    break
        if con_col is None:
            continue

        item_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['检测项目', 'Test Item', 'Item', '测试项目']):
                item_col = idx
                break

        current_item = ""
        table_rows = list(table.rows)

        for row_idx, row in enumerate(table_rows[1:], start=1):
            cells = get_row_cells_text_fast(row)
            if not cells:
                continue

            first_cell = cells[0].strip() if len(cells) > 0 else ""
            is_conclusion_row = first_cell in ['结论', 'Conclusion']

            if item_col is not None and item_col < len(cells) and not is_conclusion_row:
                item_val = cells[item_col].strip()
                if item_val and item_val not in ['检测项目', 'Test Item', 'Item',
                                                  '测试项目', '结论', 'Conclusion']:
                    current_item = item_val

            # 读取结论值（支持列偏移兜底）
            val = _read_conclusion_value(row, con_col)
            if val:
                item_for = current_item if current_item not in ['结论', 'Conclusion'] else ""

                if is_result_fallback:
                    # 增强判断：逐值语义分析
                    resolved = _resolve_result_value(val, cells, con_col)
                    if resolved is None:
                        continue
                    if resolved[0] == 'pass':
                        all_items.append((item_for, '符合'))
                    elif resolved[0] == 'fail':
                        all_items.append((item_for, resolved[1]))
                    else:
                        all_items.append((item_for, '需人工查看'))
                else:
                    all_items.append((item_for, val))

            if is_conclusion_row:
                for idx, val in enumerate(cells[1:], start=1):
                    val = val.strip()
                    if not val or val in ['/', '-', '—']:
                        continue

                    target_item = ""
                    for prev_row in reversed(table_rows[1:row_idx]):
                        prev_cells = get_row_cells_text_fast(prev_row)
                        if idx < len(prev_cells):
                            prev_res = prev_cells[idx].strip()
                            prev_item = (prev_cells[item_col].strip()
                                         if item_col is not None and item_col < len(prev_cells)
                                         else "")
                            if prev_res and prev_item \
                                    and prev_item not in ['结论', 'Conclusion']:
                                target_item = prev_item
                                break

                    if not target_item:
                        target_item = current_item if current_item not in ['结论', 'Conclusion'] else ""

                    all_items.append((target_item, val))

    if all_items:
        return summarize_conclusion_items(all_items, all_fail_short=True)
    return ""


def extract_conclusion_from_paragraphs_zh(doc):
    """
    兜底：从段落文本中抓取检测结论。
    寻找包含"结论"、"符合"、"通过"、"合格"等关键词的句子。
    """
    paras = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    if not paras:
        return ""

    # 直接包含"检测结论"或"结论"标题的段落，其后续内容
    capture = False
    captured_lines = []
    for text in paras:
        if re.search(r'检测结论|测试结论|结论[：:]', text):
            # 去掉标题本身
            content = re.sub(r'.*?(检测结论|测试结论|结论)[：:\s]*', '', text, count=1).strip()
            if content and len(content) > 3:
                captured_lines.append(content)
            capture = True
            continue
        if capture:
            # 遇到常见终止词停止
            if any(k in text for k in ['检测结果', '检测仪器', '样品描述', '样品照片', '备注', '编制', '审核', '批准', 'End of Report']):
                break
            if len(text) < 3:
                continue
            captured_lines.append(text)

    if captured_lines:
        full = '\n'.join(captured_lines)
        cat = classify_text(full)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            return summarize_conclusion_items([("", v) for v in captured_lines], all_fail_short=True)
        # 无法分类但明确含结论关键词，保留原文
        return full

    # 没有明确结论标题时，找含结论语义的单句
    conclusion_keywords = ['符合', '通过', '合格', '满足', '达标', '未检出', '不符合', '不合格', '不通过']
    for text in paras:
        # 过滤掉太短的和明显不是结论的句子
        if len(text) < 5 or len(text) > 200:
            continue
        if any(k in text for k in conclusion_keywords):
            # 进一步过滤：必须含明确结论词，且不含检测项目/方法等无关词
            if any(k in text for k in ['检测项目', '检测方法', '检测仪器', '样品描述']):
                continue
            cat = classify_text(text)
            if cat == 'pass':
                return "符合"
            if cat == 'fail':
                return text
            # 文本含结论关键词但无法明确分类，若不含"需人工查看"等无效词则保留
            if '需人工查看' not in text:
                return text

    return ""


def extract_conclusion_zh(doc, filename):
    """按优先级提取检测结论"""
    con = extract_conclusion_from_filename_zh(filename)
    if con:
        return con

    con = extract_conclusion_from_independent_table_zh(doc)
    if con:
        return con

    con = extract_conclusion_from_mixed_table_zh(doc)
    if con:
        return con

    con = extract_conclusion_from_result_tables_zh(doc)
    if con:
        return con

    con = extract_conclusion_from_paragraphs_zh(doc)
    if con:
        return con

    return ""


# ==============================================
# 函数4：提取样品部件总数
# ==============================================
def extract_sample_count(doc):
    """
    提取样品部件总数（ds修改版）：
    0. 新增：优先扫描"样品描述"特征表（产品编号/样品序号/部件名称等），
       找到直接返回；有表但无部件拆分则返回1。
    1. 其次在"检测结果"模块内寻找：
       - 有序号/部件编号/部件号列时，取最大基号；同时用样品描述/部件描述去重行数做二次校验。
       - 无序号列时，只按样品描述/部件描述去重行数统计。
    2. 识别无表头部件表：首行首列是部件编号、第二列是描述。
    3. 识别邻苯类组合编号表（如 1+9+19+25+26+33）。
    4. 排除 CAS No. 被误当成序号列。
    5. 检测结果模块没有部件信息时，退到独立的"样品描述"模块，取最大序号。
    """
    max_serial = 0
    desc_count = 0
    found_result_part_info = False
    has_result_table_no_parts = False

    result_section_markers = ['检测结果', '测试结果']
    result_markers = ['结果', '限值', '方法检出限', '序号', 'Result', 'Limit', '结论', '符合', 'XRF', 'mg/kg']
    test_item_markers = ['铅', '镉', '汞', '铬', '溴', 'Pb', 'Cd', 'Hg', 'Cr', 'Br']
    method_table_markers = ['检测方法', '测试方法', '检测仪器', '测试仪器']
    serial_keywords = ['序号', '部件编号', '部件号', '产品编号', '样品编号',
                       '样品序号', '产品序号', '部件序号',
                       'No.', 'No', 'Part No', 'Sample No', '部件号']
    desc_keywords = ['样品描述', '部件描述', '描述', '部件名称', '样品名称',
                     'Part Description', 'Sample Description', 'Description']

    # 样品描述表特征关键词（含任一即优先识别为样品描述表）
    sample_desc_markers = ['样品描述', '样品组成', '样品序号', '产品序号', '部件序号', '产品编号', '部件名称',
                           'Sample Description', 'Sample composition', 'Part Description']
    # 样品描述表互斥关键词：含任一结果表特征的表格不按样品描述表处理
    sample_desc_exclude = ['检测项目', '限值', 'mg/kg', 'Result', 'Limit',
                           'CAS No', 'CAS No.', 'Test item', '检测仪器', '方法检出限', 'XRF']

    part_id_pattern = re.compile(r'^(\d+)(-\d+)?[A-Z]?\)?$')

    def _base_num(val):
        m = re.match(r'^(\d+)', str(val).strip())
        return int(m.group(1)) if m else None

    def _update_max(val):
        nonlocal max_serial
        n = _base_num(val)
        if n and n > max_serial:
            max_serial = n

    def _is_header_like_part_table(first_row):
        if len(first_row) < 2:
            return False
        h0 = first_row[0]
        h1 = first_row[1]
        has_no = any(k in h0 for k in ['No.', 'No', '序号', 'Part No', 'Sample No',
                                        '产品编号', '样品编号', '样品序号', '产品序号', '部件序号', '部件编号'])
        has_desc = any(k in h1 for k in ['Description', '描述', 'Part', 'Sample'])
        return has_no and has_desc

    # ================================================================
    # 阶段0（新增）：样品描述表优先提取
    # 含产品编号/样品序号等特征、且不含检测结果表特征的表格，
    # 优先按样品描述表提取序号/描述列，提取逻辑与结果表趋同。
    # 找到 → 直接返回；有表但无拆分 → 返回 1；未找到 → 继续后续逻辑。
    # ================================================================
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 跳过方法表
        if any(k in first_row_text for k in method_table_markers):
            continue

        has_sample_desc = any(k in first_row_text for k in sample_desc_markers)
        has_result_feature = any(k in first_row_text for k in sample_desc_exclude)

        if has_sample_desc and not has_result_feature:
            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                sample_max_serial = 0
                sample_desc_count = 0

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        m = re.match(r'^(\d+)', str(val).strip())
                        if m:
                            n = int(m.group(1))
                            if n > sample_max_serial:
                                sample_max_serial = n

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    sample_desc_count = len(descs)

                count = sample_max_serial if sample_max_serial > 0 else sample_desc_count
                if count > 2000:
                    return "样品部件总数异常请人工检查"
                if count == 0:
                    return 1
                return count

    # ================================================================
    # 阶段1（原有）：检测结果表/部件表提取
    # ================================================================
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 排除方法表
        if any(k in first_row_text for k in method_table_markers):
            continue

        # ===== A. 带表头的检测结果表 =====
        is_result_table = (any(k in first_row_text for k in result_section_markers) or
                           any(k in first_row_text for k in result_markers) or
                           any(k in first_row_text for k in test_item_markers))
        if is_result_table:
            # 排除 SVHC/REACH 化学物质清单表：表头含 CAS No. 且含 EC No. 或物质名称
            # 这类表的“序号”是化学物质编号（如 1~250），不是样品部件序号
            is_chemical_list = (
                'CAS' in first_row_text.upper() and
                ('EC' in first_row_text.upper() or
                 '物质名称' in first_row_text or
                 'Substance' in first_row_text)
            )
            if is_chemical_list:
                continue

            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                found_result_part_info = True

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        _update_max(val)

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    if len(descs) > desc_count:
                        desc_count = len(descs)
            else:
                # 有结果表但没有序号/描述列，通常是单一样品的结果表
                has_result_table_no_parts = True

        # ===== B. 无表头部件表（首行即数据）或 No./序号 头部件表 =====
        is_part_table = False
        start_row_idx = 0
        if len(first_row) >= 2:
            c0 = first_row[0].strip()
            c1 = first_row[1].strip() if len(first_row) > 1 else ""
            # 数据式：首列是编号，第二列是描述
            if (part_id_pattern.match(c0) and c1 and not c1.isdigit() and
                    not any(k in c1 for k in method_table_markers + result_section_markers + ['结论', 'Conclusion', '结果', 'Result', '方法', 'Method'])):
                is_part_table = True
            # 表头式：No./序号/Part No. + 描述
            elif _is_header_like_part_table(first_row):
                is_part_table = True
                start_row_idx = 1

        if is_part_table:
            numeric_rows = 0
            for row in table.rows[start_row_idx:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                if part_id_pattern.match(val):
                    _update_max(val)
                    numeric_rows += 1
            if numeric_rows > 0:
                found_result_part_info = True

        # ===== C. 邻苯类组合编号表：Test Items | Result(mg/kg) ... 第二行为组合编号 =====
        if (len(table.rows) >= 3 and
                any(k in first_row_text for k in ['Test Items', '测试项目']) and
                'Result' in first_row_text):
            sub_header = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[1])]
            if sub_header and ('Test Items' in sub_header[0] or '测试项目' in sub_header[0]):
                found_nums = False
                for cell in sub_header:
                    nums = re.findall(r'\b\d+\b', cell)
                    if nums:
                        found_nums = True
                        for n in nums:
                            _update_max(n)
                if found_nums:
                    found_result_part_info = True

    if found_result_part_info:
        count = max_serial if max_serial > 0 else desc_count
        if count > 2000:
            return "样品部件总数异常请人工检查"
        # 有结果表但序号/描述列均无有效值，视为单一样品的汇总结果
        if count == 0:
            return 1
        return count

    # 退到独立的"样品描述"模块
    count = extract_sample_count_from_sample_description_section(doc)
    if count > 0:
        if count > 2000:
            return "样品部件总数异常请人工检查"
        return count

    # 最终兜底：存在结果表但无部件拆分，视为 1 个样品
    if has_result_table_no_parts:
        return 1

    return "未检测到"


def extract_sample_count_from_sample_description_section(doc):
    """
    从独立的"样品描述"模块（段落标题）中提取最大序号。
    排除包含"照片/Photo"的段落。
    """
    body = doc.element.body
    in_section = False
    target_table = None
    section_paragraphs = []

    for child in body:
        tag = child.tag.split('}')[-1]
        text = ''.join(child.itertext()).strip()

        if tag == 'p':
            lower = text.lower()
            if '样品描述' in text or 'sample description' in lower:
                if '照片' in text or 'photo' in lower:
                    continue
                in_section = True
                continue
            if in_section and any(k in text for k in ['检测方法', '测试结果', '检测结果', '测试方法',
                                                        '检测项目', '测试项目', '要求', '结论']):
                break

        if in_section:
            if tag == 'p':
                section_paragraphs.append(text)
            elif tag == 'tbl' and target_table is None:
                target_table = child
                break

    max_serial = 0
    found = False

    def _try_match(s):
        nonlocal max_serial, found
        s = s.strip()
        for pattern in [r'^(\d+)\s*[:：.\)、]\s*', r'^(\d+)$']:
            m = re.match(pattern, s)
            if m:
                found = True
                seq = int(m.group(1))
                if seq > max_serial:
                    max_serial = seq
                return True
        return False

    if target_table is not None:
        for tc in target_table.iter():
            if tc.tag.endswith('tc'):
                cell_text = ''.join(tc.itertext()).strip()
                _try_match(cell_text)
    else:
        for text in section_paragraphs:
            for line in re.split(r'[\r\n]+', text):
                _try_match(line.strip())

    # 若段落标题方式没找到，再扫描全文档中表头为"样品描述"的表格
    if max_serial == 0:
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
            first_row_text = ' '.join(first_row)
            if '样品描述' not in first_row_text and 'sample description' not in first_row_text.lower():
                continue
            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                m = re.match(r'^(\d+)', val)
                if m:
                    seq = int(m.group(1))
                    if seq > max_serial:
                        max_serial = seq

    return max_serial
def is_standard_method(method):
    """判断是否为标准方法（包含标准前缀或年份）"""
    if not method:
        return False
    std_prefixes = ['GB/T', 'GB', 'IEC', 'ISO', 'ASTM', 'EPA', 'US EPA', 'QC/T', 'QB/T', 'SJ/T']
    if any(prefix in method for prefix in std_prefixes):
        return True
    if re.search(r'\d{4}', method):
        return True
    return False


def extract_test_methods_from_tables(doc):
    """从表格中提取检测方法/仪器"""
    methods = []
    seen = set()

    for table_idx, table in enumerate(doc.tables, 1):
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 识别"检测项目/检测方法/检测仪器"结构的表格（兼容更多同义词）
        if '检测项目' in first_row_text or '测试项目' in first_row_text or '项目' in first_row_text or \
           'Test Item' in first_row_text or 'Testing item' in first_row_text or \
           ('检测方法' in first_row_text or '测试方法' in first_row_text or '方法' in first_row_text or
            '检测仪器' in first_row_text or '测试仪器' in first_row_text or '仪器' in first_row_text):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到检测项目/方法/仪器表格")

            item_col = find_column_index(first_row, ["检测项目", "测试项目", "项目", "Test Item", "Testing item", "Test Items"])
            method_col = find_column_index(first_row, ["检测方法", "测试方法", "方法", "Test Method", "Testing Method", "Method"])
            instrument_col = find_column_index(first_row, ["检测仪器", "测试仪器", "仪器", "Test Instrument", "Instrument", "Equipment", "检测设备"])

            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) < 2:
                    continue

                item = cells[item_col] if item_col is not None and item_col < len(cells) else ""
                method = cells[method_col] if method_col is not None and method_col < len(cells) else ""
                instrument = cells[instrument_col] if instrument_col is not None and instrument_col < len(cells) else ""

                if not item or item == '<br />' or item.isdigit():
                    continue
                if 'Limit' in item or '△' in item or 'Note' in item or '注' in item:
                    continue
                if item in seen:
                    continue
                seen.add(item)

                methods.append({
                    "检测项目": item,
                    "检测方法": method,
                    "检测仪器": instrument
                })

        # 识别"仪器名称/型号/校准有效日期"结构的表格
        elif is_instrument_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到仪器明细表格")

            name_col = find_column_index(first_row, ["仪器名称", "仪器", "设备名称", "设备"])
            model_col = find_column_index(first_row, ["型号", "规格型号"])

            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) < 2:
                    continue

                instrument = cells[name_col] if name_col is not None and name_col < len(cells) else ""
                model = cells[model_col] if model_col is not None and model_col < len(cells) else ""

                if not instrument or instrument == '<br />':
                    continue

                dedup_key = f"{instrument}_{model}"
                if dedup_key in seen:
                    continue
                seen.add(dedup_key)

                full_instrument = instrument
                if model:
                    full_instrument += f" ({model})"

                methods.append({
                    "检测项目": "",
                    "检测方法": "",
                    "检测仪器": full_instrument
                })

    return methods


def extract_test_methods_from_paragraphs(doc):
    """从段落中提取检测方法（兜底）"""
    methods = []
    method_patterns = [
        r'参考\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'依据\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'采用\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'按照\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
    ]

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        for pattern in method_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                method = match.strip()
                if is_standard_method(method) and method not in [m["检测方法"] for m in methods]:
                    methods.append({
                        "检测项目": "",
                        "检测方法": method,
                        "检测仪器": ""
                    })

    return methods


def extract_test_items_from_result_tables(doc):
    """从检测结果表格中提取检测项目（兜底，用于填充检测项目列）"""
    items = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        if '检测项目' not in first_row_text:
            continue

        item_col = find_column_index(first_row, ["检测项目", "测试项目", "项目", "Test Item"])
        if item_col is None:
            continue

        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if item_col >= len(cells):
                continue

            item = cells[item_col].strip()
            if not item or item == '<br />' or item.isdigit():
                continue
            if 'Limit' in item or '△' in item or 'Note' in item or '注' in item:
                continue
            if item in seen:
                continue
            seen.add(item)
            items.append(item)

    return items


def extract_test_methods_long(doc):
    """长文档：从首页检测方法总结表格提取"""
    methods = []

    for table in doc.tables[:5]:
        if len(table.rows) >= 2 and len(table.rows[0].cells) == 2:
            first_cell = get_cell_text(table.rows[0].cells[0])
            if re.match(r'^(QC/T|GB/T|IEC|GB|ISO|ASTM|EPA|US EPA)\s*\d', first_cell, re.IGNORECASE):
                for row in table.rows:
                    cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                    if len(cells) == 2 and cells[0] and cells[1]:
                        full_method = f"{cells[0]} {cells[1]}"
                        methods.append({
                            "检测项目": "",
                            "检测方法": full_method,
                            "检测仪器": ""
                        })
                break

    return methods


def merge_test_methods_by_item(methods):
    """
    按检测项目合并方法/仪器：同一项目出现多次时，
    把多个方法、仪器用" / "合并到一条记录中。
    """
    merged = {}
    for m in methods:
        item = m.get("检测项目", "")
        if not item:
            continue
        if item not in merged:
            merged[item] = {"检测项目": item, "methods": [], "instruments": []}
        method = m.get("检测方法", "")
        instrument = m.get("检测仪器", "")
        if method and method not in merged[item]["methods"]:
            merged[item]["methods"].append(method)
        if instrument and instrument not in merged[item]["instruments"]:
            merged[item]["instruments"].append(instrument)

    result = []
    for item, data in merged.items():
        result.append({
            "检测项目": item,
            "检测方法": " / ".join(data["methods"]),
            "检测仪器": " / ".join(data["instruments"]),
        })
    return result


def extract_dates_from_paragraphs_zh(doc):
    """
    兜底：从段落文本中提取样品接收日期和检测期间。
    匹配常见中文日期表达。
    """
    recv_date = ""
    period = ""

    date_pattern = r'(\d{4}\s*[-年/]\s*\d{1,2}\s*[-月/]\s*\d{1,2})'
    period_pattern = r'(\d{4}\s*[-年/]\s*\d{1,2}\s*[-月/]\s*\d{1,2})\s*[-—~至]\s*(\d{4}\s*[-年/]\s*\d{1,2}\s*[-月/]\s*\d{1,2})'

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        # 优先匹配检测期间（两个日期之间用 - 至 连接）
        if not period and ('检测期间' in text or '测试周期' in text or '测试期间' in text):
            m = re.search(period_pattern, text)
            if m:
                period = f"{m.group(1).replace(' ', '')}至{m.group(2).replace(' ', '')}"
            else:
                # 单日期也先收下
                m = re.search(date_pattern, text)
                if m:
                    period = m.group(1).replace(' ', '')

        if not recv_date and ('样品接收日期' in text or '接收日期' in text or '收样日期' in text or '收到日期' in text):
            m = re.search(date_pattern, text)
            if m:
                recv_date = m.group(1).replace(' ', '')

        if recv_date and period:
            break

    return recv_date, period


# ==============================================
# 函数6：完整提取中文报告
# ==============================================
def extract_chinese_report(doc, filename):
    """中文报告完整提取，返回字段字典"""
    doc_type = detect_doc_type(doc)

    if VERBOSE:
        print(f"  📋 中文文档类型: {'长文档' if doc_type == 'long' else '短文档'}")

    # 1. 基础字段
    raw_fields = extract_basic_info(doc)

    result = {}
    for raw_key, raw_val in raw_fields.items():
        std_key = ZH_FIELD_MAPPING.get(raw_key)
        if std_key:
            result[std_key] = raw_val

    # 2. 检测要求/结论（优先使用专题方案2逻辑，空时回退旧逻辑）
    req = extract_requirement_zh(doc)
    con = extract_conclusion_zh(doc, filename)

    # 旧逻辑兜底
    if not req or not con:
        old_req, old_con = extract_test_requirement_and_conclusion(doc)
        if not req and old_req:
            req = old_req
        if not con and old_con:
            con = old_con

    # 兜底：仍未取到检测结论时，从结果表的“结论”列提取
    if not con:
        con = extract_conclusion_from_result_column_zh(doc)

    if req:
        result["检测要求"] = req
    if con:
        result["检测结论"] = con

    # 3. 报告编号从文件名取
    m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z]?\d*)', filename, re.IGNORECASE)
    if m:
        result["报告编号"] = m.group(1)

    # 3.5 日期兜底：若基础字段未取到样品接收日期/检测期间，从段落补
    if not result.get("样品接收日期") or not result.get("检测期间"):
        recv_date, period = extract_dates_from_paragraphs_zh(doc)
        if recv_date and not result.get("样品接收日期"):
            result["样品接收日期"] = recv_date
        if period and not result.get("检测期间"):
            result["检测期间"] = period

    # 4. 样品部件总数
    sample_count = extract_sample_count(doc)
    result["样品部件总数"] = str(sample_count)

    # 5. 检测方法：优先直接读取"检测项目/检测方法/检测仪器"表，无命中再按长文档方法总结兜底
    direct_methods = extract_test_methods_from_tables(doc)
    if direct_methods:
        test_methods = direct_methods
    elif doc_type == "long":
        test_methods = extract_test_methods_long(doc)
    else:
        test_methods = []

    # 统一用段落方法补全/兜底
    paragraph_methods = extract_test_methods_from_paragraphs(doc)
    if test_methods:
        # 如果表格方法缺少检测方法，用段落方法补全（按顺序填入）
        pm_iter = iter(paragraph_methods)
        for tm in test_methods:
            if not tm["检测方法"]:
                for pm in pm_iter:
                    if pm["检测方法"]:
                        tm["检测方法"] = pm["检测方法"]
                        break
        # 段落方法中尚未被使用的，追加为新条目
        used_methods = {tm["检测方法"] for tm in test_methods if tm["检测方法"]}
        for pm in paragraph_methods:
            if pm["检测方法"] and pm["检测方法"] not in used_methods:
                test_methods.append(pm)
                used_methods.add(pm["检测方法"])
    else:
        # 没有任何表格方法时，直接用段落方法兜底
        test_methods = paragraph_methods

    # 如果方法条目缺少检测项目，尝试从检测结果表补全
    if test_methods:
        test_items = extract_test_items_from_result_tables(doc)
        for i, item in enumerate(test_items):
            if i < len(test_methods) and not test_methods[i]["检测项目"]:
                test_methods[i]["检测项目"] = item

    # 兜底：如果仍没有任何方法条目，但文档段落中有标准方法，生成一个通用条目
    if not test_methods:
        std_methods = extract_test_methods_from_paragraphs(doc)
        if std_methods:
            test_methods = std_methods
    # 按检测项目合并重复方法/仪器（例如六价铬两种方法合并为一条）
    test_methods = merge_test_methods_by_item(test_methods)

    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]

    return result, test_methods


# ==============================================
# 主函数：处理单个文档
# ==============================================
def process_word_file(input_path, output_excel):
    """
    处理单个中文 Word 检测报告，提取信息并生成 Excel

    参数:
        input_path: Word 文档路径
        output_excel: 输出 Excel 路径
    """
    from docx import Document

    print("=" * 60)
    print(f"开始处理: {os.path.basename(input_path)}")
    print("=" * 60)

    # 1. 检查文件是否存在
    if not os.path.exists(normalize_path(input_path)):
        print(f"❌ 错误：文件不存在 - {input_path}")
        return False

    # 2. 修复文档（如果需要）
    try:
        actual_file = fix_docx_if_needed(input_path)
    except Exception as e:
        print(f"❌ 错误：无法打开文档 - {e}")
        print("   请检查文件是否为有效的 Word 文档(.docx格式)")
        return False

    # 3. 打开文档
    try:
        doc = Document(actual_file)
    except ImportError:
        print("❌ 错误：未安装 python-docx 库")
        print("   请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 错误：打开文档失败 - {e}")
        return False

    if VERBOSE:
        print(f"\n📄 文档信息:")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")

    # 4. 初始化结果字典
    result = {
        "文件名": os.path.basename(input_path),
        "报告编号": "",
        "申请商": "",
        "申请商地址": "",
        "制造商地址": "",
        "样品名称": "",
        "零件号": "",
        "参考零件号": "",
        "材质": "",
        "车型": "",
        "生产日期": "",
        "供应商代码": "",
        "测试类型": "",
        "样品编号": "",
        "样品接收日期": "",
        "检测期间": "",
        "检测要求": "",
        "检测结论": "",
        "报告日期": "",
        "样品部件总数": "",
        "备注": "",
    }

    # 添加检测方法列
    for i in range(1, MAX_TEST_METHODS + 1):
        result[f"检测项目{i}"] = ""
        result[f"检测方法{i}"] = ""
        result[f"检测仪器{i}"] = ""

    # 5. 提取
    extracted, test_methods = extract_chinese_report(doc, os.path.basename(input_path))
    result.update(extracted)

    # 6. 生成 Excel
    print(f"\n💾 生成 Excel 文件...")

    col_order = [
        "报告编号", "申请商", "申请商地址", "制造商地址",
        "样品名称", "零件号", "参考零件号", "材质",
        "车型", "生产日期", "供应商代码", "测试类型",
        "样品编号", "样品接收日期", "检测期间",
        "检测要求", "检测结论", "报告日期",
        "样品部件总数", "备注"
    ]

    actual_method_count = min(len(test_methods), MAX_TEST_METHODS)
    for i in range(1, actual_method_count + 1):
        col_order.append(f"检测项目{i}")
        col_order.append(f"检测方法{i}")
        col_order.append(f"检测仪器{i}")

    col_order.extend(["文件名"])

    df = pd.DataFrame([result])
    for col in col_order:
        if col not in df.columns:
            df[col] = ""
    df = df[col_order]

    try:
        df.to_excel(output_excel, index=False)
        print(f"✅ 成功保存: {output_excel}")
    except ImportError:
        print("❌ 错误：未安装 openpyxl 库")
        print("   请运行: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 错误：保存 Excel 失败 - {e}")
        return False

    # 统计信息
    non_empty = sum(1 for val in result.values() if val)
    print(f"\n📊 提取完成！共提取 {non_empty} 个非空字段")
    print("=" * 60)

    return True


# ==============================================
# 程序入口
# ==============================================
def main():
    print("\n" + "=" * 60)
    print("  中文检测报告信息提取工具 Plus 4")
    print("=" * 60)

    # 检查依赖
    try:
        import docx
        import pandas
    except ImportError as e:
        print(f"\n❌ 缺少依赖库: {e}")
        print("请运行以下命令安装依赖:")
        print("  pip install python-docx pandas openpyxl")
        sys.exit(1)

    # 处理文件
    success = process_word_file(INPUT_FILE, OUTPUT_EXCEL)

    if success:
        print("\n🎉 处理完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)


if __name__ == "__main__":
    main()
