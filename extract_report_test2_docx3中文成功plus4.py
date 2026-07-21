# -*- coding: utf-8 -*-
"""
中文检测报告信息提取工具 Plus 4
功能：从中文 RoHS/ELV/REACH 检测报告 Word 文档中提取关键信息，生成 Excel 表格
使用说明：
1. 修改下方【配置区】的文件路径和输出路径
2. 运行脚本即可生成 Excel 文件
3. 支持自动修复部分损坏的 docx 文件（NULL 关系、Bad CRC-32 等）
依赖库：python-docx, pandas, openpyxl
安装命令：pip install python-docx pandas openpyxl
"""
import os
import re
import sys
import zipfile
import tempfile
import html
import pandas as pd

# ==============================================
# 【配置区】请根据实际情况修改以下参数
# ==============================================
# 待处理的 Word 文档路径（支持 Windows 路径，如 r"D:\Kathy\PDF提取工具\test.docx"）
INPUT_FILE = r"D:\Kathy\PDF提取工具\ROHS-S19062103203001.docx"
# 输出 Excel 文件路径
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\检测报告提取信息中文docx4.xlsx"
# 最大检测方法数量（预留列数）
MAX_TEST_METHODS = 100
# 是否显示详细处理日志
VERBOSE = True


# ==============================================
# 【中文字段映射表】中文表头 → 标准含义
# ==============================================
ZH_FIELD_ALIASES = {
    # 申请商相关
    "申请商": ["申请商", "委托单位", "委托方", "客户", "客户名称"],
    "申请商地址": ["地址", "申请商地址", "委托单位地址", "客户地址"],

    # 样品相关
    "样品名称": ["样品名称", "产品名称", "样品种类", "样品描述", "产品描述", "零部件名称", "Part Name"],
    "零件号": ["零件号", "型号", "产品型号", "料号", "Part No", "Part number", "零件型号"],
    "参考零件号": ["参考零件号", "参考型号"],
    "材质": ["材质", "材料", "样品材质"],
    "车型": ["车型", "车辆型号", "车型/项目号", "项目号", "Vehicle model"],
    "生产日期": ["生产日期", "生产批号", "Production Date"],
    "供应商代码": ["供应商代码", "供应商", "Supplier"],
    "测试类型": ["测试类型", "检测类型", "测试项目"],

    # 汽车部新增关键字段
    "主机厂": ["主机厂", "汽车公司", "Automobile company", "Buyer", "买家"],
    "制造商": ["制造商", "OEM", "生产厂家", "工厂", "Manufacturer", "Factory"],

    # 样品编号相关
    "样品编号": ["样品编号", "样品号", "样本编号", "Sample No", "Sample Number", "Sample number"],

    # 日期相关
    "样品接收日期": ["样品接收日期", "接收日期", "样品收到日期", "收样日期"],
    "检测期间": ["检测期间", "测试周期", "检测周期", "测试期间"],
    "报告日期": ["报告日期", "日期", "签发日期", "发布日期"],

    # 人员相关
    "编制": ["编制", "编写", "编制人"],
    "审核": ["审核", "校核", "审核人"],
    "批准": ["批准", "签发", "批准人"],

    # 其他
    "备注": ["备注", "说明", "附注"],
}


def build_zh_field_mapping():
    """把别名表展开为 {别名: 标准字段名}"""
    mapping = {}
    for std_key, aliases in ZH_FIELD_ALIASES.items():
        for alias in aliases:
            mapping[alias] = std_key
    return mapping


ZH_FIELD_MAPPING = build_zh_field_mapping()


# ==============================================
# 工具函数：路径处理（解决 Windows 长路径问题）
# ==============================================
def normalize_path(file_path):
    """为 Windows 长路径添加 \\\\?\\\\ 前缀"""
    prefix = r"\\\\?\\"
    if sys.platform == 'win32' and not file_path.startswith(prefix):
        abs_path = os.path.abspath(file_path)
        if len(abs_path) > 230:  # 留余量
            return prefix + abs_path
    return file_path


# ==============================================
# 工具函数：修复有问题的 docx 文件
# ==============================================
def fix_docx_if_needed(file_path):
    """
    检查并修复有问题的 docx 文件
    支持：NULL 关系问题、Bad CRC-32 等 zip 损坏
    返回修复后的文件路径，如果不需要修复则返回原路径
    """
    file_path = normalize_path(file_path)

    def _try_open(path):
        from docx import Document
        try:
            Document(path)
            return True
        except Exception:
            return False

    if _try_open(file_path):
        return file_path

    # 需要修复
    if VERBOSE:
        print(f"  ⚠️  检测到文档损坏，正在自动修复...")

    temp_dir = tempfile.gettempdir()
    prefix = r"\\\\?\\"
    base_name = os.path.basename(file_path.replace(prefix, ''))
    fixed_path = os.path.join(temp_dir, f"fixed_{base_name}")

    try:
        with zipfile.ZipFile(file_path, 'r') as zin:
            with zipfile.ZipFile(fixed_path, 'w') as zout:
                for item in zin.infolist():
                    # 修复 NULL 关系
                    if item.filename == 'word/_rels/document.xml.rels':
                        try:
                            content = zin.read(item.filename).decode('utf-8')
                            content = re.sub(r'<Relationship[^>]*Target="[^"]*NULL[^"]*"[^>]*/>', '', content)
                            zout.writestr(item, content)
                        except zipfile.BadZipFile:
                            # 如果连关系文件都损坏，跳过（极少见）
                            continue
                        continue

                    # 尝试复制其他文件；遇到 Bad CRC 时，若是媒体文件则跳过
                    try:
                        data = zin.read(item.filename)
                        zout.writestr(item, data)
                    except zipfile.BadZipFile:
                        if item.filename.startswith('word/media/'):
                            if VERBOSE:
                                print(f"    跳过损坏的图片: {item.filename}")
                            continue
                        raise

        if _try_open(fixed_path):
            if VERBOSE:
                print(f"  ✅ 文档修复成功")
            return fixed_path
        else:
            raise Exception("修复后仍无法打开文档")
    except Exception as fix_error:
        if VERBOSE:
            print(f"  ❌ 文档修复失败: {fix_error}")
        raise


# ==============================================
# 工具函数：获取单元格完整文本（包括 SDT 内容控件）
# ==============================================
def get_cell_text(cell):
    """
    获取单元格的完整文本，包括 SDT（结构化文档标签/内容控件）中的内容
    解决 python-docx 无法读取内容控件文本的问题
    """
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    result = html.unescape(result)
    return result.strip()


def get_row_cells_text_fast(row):
    """
    获取一行中所有单元格的文本，包括被 SDT（结构化文档标签）包裹的单元格
    解决 python-docx 无法识别 SDT 包裹的 tc 单元格的问题
    """
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cell_text = html.unescape(cell_text)
        cells_text.append(cell_text.strip())
    return cells_text


# 兼容总提取代码的别名
get_row_cells_text = get_row_cells_text_fast


# ==============================================
# 工具函数：文本预处理
# ==============================================
def normalize_text(text):
    """清理文本：去首尾空、统一换行/空格"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def normalize_key(key):
    """清理字段名：去冒号、去首尾空、统一空格"""
    if not key:
        return ""
    key = normalize_text(key)
    key = key.rstrip('：:').strip()
    return key


def normalize_key_bilingual(key):
    """
    针对汽车部报告中“中文(English)”双语字段名做归一化。
    去掉英文括号及括号内容、尾部中英文冒号/空格，再交给 normalize_key。
    例如：
      '申请商(Applicant) :' -> '申请商'
      '车型/项目号(Model/Project No.):' -> '车型/项目号'
      'Part Name:' -> 'Part Name'
    """
    if not key:
        return ""
    # 先做一次基础清洗
    key = normalize_text(key)
    # 去掉英文括号及其中的内容（支持嵌套一层括号）
    key = re.sub(r'\s*\([A-Za-z0-9\s&/\.\-_,#()]+\)\s*', '', key)
    # 去掉尾部中英文冒号
    key = key.rstrip('：:').strip()
    return key


# ==============================================
# 函数1：判断文档类型（长文档/短文档）
# ==============================================
def detect_doc_type(doc):
    """
    判断中文文档类型：
    - 长文档：首页有检测方法总结表格（两列，方法编号+名称，如 QC/T、GB/T、IEC 开头）
    - 短文档：首页没有检测方法总结，检测方法在后面的详细表格中
    """
    # 方法1：通过表格数量判断（长文档表格很多）
    if len(doc.tables) > 30:
        return "long"

    # 方法2：检查前 5 个表格中是否有检测方法总结
    for table in doc.tables[:5]:
        if len(table.rows) >= 2 and len(table.rows[0].cells) == 2:
            first_cell = get_cell_text(table.rows[0].cells[0])
            if re.match(r'^(QC/T|GB/T|IEC|GB|ISO|ASTM|EPA|US EPA)\s*\d', first_cell, re.IGNORECASE):
                return "long"

    return "short"


# ==============================================
# 表格类型判断辅助函数
# ==============================================
def is_result_table(first_row):
    """判断是否为检测结果表格（应跳过）"""
    text = ' '.join(first_row)
    has_item = any(kw in text for kw in ["检测项目", "Test Item", "测试项目", "Test item(s)"])
    has_result = any(kw in text for kw in ["限值", "结果", "方法检出限", "Limit", "Result", "MDL"])
    return has_item and has_result


def is_sample_composition_table(table):
    """判断是否为样品组成/描述表格（第一列是数字序号）"""
    if len(table.rows) < 3:
        return False

    checked = 0
    digit_count = 0
    for row in table.rows[:5]:
        cells = get_row_cells_text_fast(row)
        if not cells:
            continue
        checked += 1
        first_cell = cells[0].strip()
        if re.match(r'^(\d+)(-\d+)?$', first_cell):
            digit_count += 1

    return checked >= 3 and digit_count >= 3


def is_instrument_table(first_row):
    """判断是否为仪器/设备表格"""
    text = ' '.join(first_row)
    instrument_keywords = ["仪器名称", "仪器", "设备名称", "设备", "Instrument", "Equipment"]
    return any(kw in text for kw in instrument_keywords)


def find_column_index(first_row, keywords):
    """根据关键词列表找到列索引"""
    for idx, cell in enumerate(first_row):
        for kw in keywords:
            if kw in cell:
                return idx
    return None


# ==============================================
# 函数2：提取基本信息
# ==============================================
def extract_basic_info(doc):
    """从两列、四列表格中提取中文报告基础字段"""
    fields = {}
    manufacturer_found = False

    for table_idx, table in enumerate(doc.tables[:20], 1):
        if len(table.rows) < 1:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' '.join(first_row)
        cell_count = len(first_row)

        # 跳过检测结果表格
        if is_result_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 检测结果表格，跳过")
            continue

        # 跳过样品组成/样品描述表格
        if is_sample_composition_table(table):
            if VERBOSE:
                print(f"  表格 {table_idx}: 样品组成表格，跳过")
            continue

        # 跳过仪器明细表（单独处理）
        if is_instrument_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 仪器表格，单独处理")
            continue

        # 处理两列表格（键值对）
        if cell_count == 2:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]

                # 处理只有 1 个单元格的行（可能是合并单元格的标签行）
                if len(cells) == 1:
                    key = normalize_key_bilingual(cells[0])
                    if key and len(key) < 60:
                        fields[key] = ""
                    continue

                if len(cells) == 2:
                    key = normalize_key_bilingual(cells[0])
                    val = cells[1].strip()

                    if not key or len(key) >= 60:
                        continue

                    # 跳过说明性文字
                    if '以下的检测样品' in key or '由客户提供' in key:
                        continue

                    # 特殊处理：遇到制造商/工厂后，下一个地址是制造商地址
                    if key in ["制造商", "工厂", "生产厂家"]:
                        manufacturer_found = True
                        if val and val != '/':
                            fields[key] = val
                        continue

                    # 特殊处理地址：区分申请商地址和制造商地址
                    if key == "地址":
                        if not manufacturer_found:
                            if val and val != '/':
                                fields["申请商地址"] = val
                        else:
                            if val and val != '/' and '制造商地址' not in fields:
                                fields["制造商地址"] = val
                        continue

                    # 忽略空值占位符
                    if val == '/':
                        val = ""

                    if val:
                        fields[key] = val

        # 处理四列表格（编制/审核/批准/日期等）
        elif cell_count == 4:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) == 4:
                    for i in range(0, 4, 2):
                        key = normalize_key_bilingual(cells[i])
                        val = cells[i + 1].strip()
                        if key and len(key) < 40:
                            if val and val != '/':
                                fields[key] = val

    return fields


# ==============================================
# 函数3：提取检测要求和检测结论（完整原文）
# ==============================================
def _split_requirement_conclusion(text):
    """
    对“检测要求及结论”混合文本进行拆分。
    把包含结论语义（符合/通过/合格/满足/达标）的最后一个分句作为检测结论，
    其余部分作为检测要求。
    """
    text = normalize_text(text)
    if not text:
        return "", ""

    # 按中文/英文句读拆分，保留分隔符
    parts = re.split(r'([。；，,!])', text)
    sentences = []
    cur = ""
    for p in parts:
        cur += p
        if p in "。；，,!":
            s = cur.strip()
            if s:
                sentences.append(s)
            cur = ""
    if cur.strip():
        sentences.append(cur.strip())

    conclusion_keywords = ['符合', '通过', '合格', '满足', '达标',
                           'comply', 'complies', 'pass', 'passed', 'conform']
    con_idx = None
    for i in range(len(sentences) - 1, -1, -1):
        low = sentences[i].lower()
        if any(k in low or k in sentences[i] for k in conclusion_keywords):
            con_idx = i
            break

    if con_idx is not None:
        requirement = "".join(sentences[:con_idx]).strip()
        conclusion = sentences[con_idx].strip()
        if not requirement:
            requirement = text
        return requirement, conclusion

    return text, ""


def extract_test_requirement_and_conclusion(doc):
    """提取检测要求和检测结论的完整原文"""
    test_requirement = ""
    test_conclusion = ""

    for table_idx, table in enumerate(doc.tables, 1):
        if test_requirement and test_conclusion:
            break

        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c) for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)

        # 情况1：2行2列表格（检测要求和结论在同一行）
        if (len(table.rows[0].cells) == 2 and
                '检测要求' in first_row_text and
                '结论' in first_row_text):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到检测要求/结论表格（2行2列）")
            if len(table.rows) >= 2:
                test_requirement = get_cell_text(table.rows[1].cells[0])
                test_conclusion = get_cell_text(table.rows[1].cells[1])
            break

        # 情况2/3：1列表格
        elif len(table.rows[0].cells) == 1:
            header = normalize_text(first_row[0])

            # 2a. 表头同时含“检测要求”和“结论”（如“检测要求及结论”）
            if '检测要求' in header and '结论' in header:
                if VERBOSE:
                    print(f"  表格 {table_idx}: 找到检测要求及结论混合表格")
                if len(table.rows) >= 2:
                    full = get_cell_text(table.rows[1].cells[0])
                    req, con = _split_requirement_conclusion(full)
                    if req and not test_requirement:
                        test_requirement = req
                    if con and not test_conclusion:
                        test_conclusion = con

            # 2b. 表头仅含“检测要求”，可能在后续行出现“检测结论:”标签
            elif '检测要求' in header:
                if not test_requirement and len(table.rows) >= 2:
                    test_requirement = get_cell_text(table.rows[1].cells[0])
                # 扫描后续标签行
                if not test_conclusion:
                    for r in range(1, len(table.rows)):
                        txt = normalize_text(get_cell_text(table.rows[r].cells[0]))
                        if '检测结论' in txt and r + 1 < len(table.rows):
                            test_conclusion = get_cell_text(table.rows[r + 1].cells[0])
                            if VERBOSE:
                                print(f"  表格 {table_idx}: 在单列表格后续行找到检测结论")
                            break

            # 2c. 表头含“检测结论”
            elif '结论' in header and '检测' in header:
                if not test_conclusion:
                    if VERBOSE:
                        print(f"  表格 {table_idx}: 找到检测结论表格（独立表格）")
                    if len(table.rows) >= 2:
                        test_conclusion = get_cell_text(table.rows[1].cells[0])

    return test_requirement, test_conclusion


def extract_conclusion_from_result_column_zh(doc):
    """
    兜底：从检测结果表格的“结论”列提取结论（如“符合”）。
    对多行结论去重后用“ / ”拼接。
    """
    conclusions = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        # 跳过“检测要求/结论”标签表
        if len(first_row) == 1 and ('检测要求' in first_row[0] or '检测结论' in first_row[0]):
            continue

        first_row_text = ' '.join(first_row)
        # 只考虑结果/结论相关表格
        if not any(k in first_row_text for k in ['结论', 'Conclusion', '符合', 'Pass', '结果', 'Result']):
            continue

        con_col = find_column_index(first_row, ['结论', 'Conclusion', '符合', 'Pass'])
        if con_col is None:
            continue

        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if con_col >= len(cells):
                continue
            val = cells[con_col].strip()
            if not val:
                continue
            # 跳过表头残留、检测项目名等
            if val in first_row:
                continue
            if val in ['结论', 'Conclusion', '结果', 'Result']:
                continue
            if val not in seen:
                seen.add(val)
                conclusions.append(val)

    # 兜底 2：结果表中直接有一行首列为“结论”，其右侧单元格为“符合/Pass”等
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        if not any(k in ' '.join(first_row) for k in ['结论', 'Conclusion', '符合', 'Pass', '结果', 'Result']):
            continue
        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if not cells:
                continue
            if cells[0].strip() in ['结论', 'Conclusion']:
                for cell in cells[1:]:
                    val = cell.strip()
                    if val and val not in ['结论', 'Conclusion', '/', '-', '—', '结果', 'Result']:
                        if val not in seen:
                            seen.add(val)
                            conclusions.append(val)

    return ' / '.join(conclusions) if conclusions else ""


# ==============================================
# 函数4：提取样品部件总数
# ==============================================
def extract_sample_count(doc):
    """
    提取样品部件总数（Plus 4 修复版）：
    1. 优先在"检测结果"模块内寻找：
       - 有序号/部件编号/部件号列时，取最大基号；同时用样品描述/部件描述去重行数做二次校验。
       - 无序号列时，只按样品描述/部件描述去重行数统计。
    2. 识别无表头部件表：首行首列是部件编号、第二列是描述。
    3. 识别邻苯类组合编号表（如 1+9+19+25+26+33）。
    4. 排除 CAS No. 被误当成序号列。
    5. 检测结果模块没有部件信息时，退到独立的"样品描述"模块，取最大序号。
    """
    max_serial = 0
    desc_count = 0
    found_result_part_info = False
    has_result_table_no_parts = False

    result_section_markers = ['检测结果', '测试结果']
    result_markers = ['结果', '限值', '方法检出限', '序号', 'Result', 'Limit', '结论', '符合', 'XRF', 'mg/kg']
    test_item_markers = ['铅', '镉', '汞', '铬', '溴', 'Pb', 'Cd', 'Hg', 'Cr', 'Br']
    method_table_markers = ['检测方法', '测试方法', '检测仪器', '测试仪器']
    serial_keywords = ['序号', '部件编号', '部件号', 'No.', 'No', 'Part No', 'Sample No', '部件号']
    desc_keywords = ['样品描述', '部件描述', '描述', 'Part Description', 'Sample Description']

    part_id_pattern = re.compile(r'^(\d+)(-\d+)?[A-Z]?\)?$')

    def _base_num(val):
        m = re.match(r'^(\d+)', str(val).strip())
        return int(m.group(1)) if m else None

    def _update_max(val):
        nonlocal max_serial
        n = _base_num(val)
        if n and n > max_serial:
            max_serial = n

    def _is_header_like_part_table(first_row):
        if len(first_row) < 2:
            return False
        h0 = first_row[0]
        h1 = first_row[1]
        has_no = any(k in h0 for k in ['No.', 'No', '序号', 'Part No', 'Sample No'])
        has_desc = any(k in h1 for k in ['Description', '描述', 'Part', 'Sample'])
        return has_no and has_desc

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 排除方法表
        if any(k in first_row_text for k in method_table_markers):
            continue

        # ===== A. 带表头的检测结果表 =====
        is_result_table = (any(k in first_row_text for k in result_section_markers) or
                           any(k in first_row_text for k in result_markers) or
                           any(k in first_row_text for k in test_item_markers))
        if is_result_table:
            # 排除 SVHC/REACH 化学物质清单表：表头含 CAS No. 且含 EC No. 或物质名称
            # 这类表的“序号”是化学物质编号（如 1~250），不是样品部件序号
            is_chemical_list = (
                'CAS' in first_row_text.upper() and
                ('EC' in first_row_text.upper() or
                 '物质名称' in first_row_text or
                 'Substance' in first_row_text)
            )
            if is_chemical_list:
                continue

            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                found_result_part_info = True

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        _update_max(val)

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    if len(descs) > desc_count:
                        desc_count = len(descs)
            else:
                # 有结果表但没有序号/描述列，通常是单一样品的结果表
                has_result_table_no_parts = True

        # ===== B. 无表头部件表（首行即数据）或 No./序号 头部件表 =====
        is_part_table = False
        start_row_idx = 0
        if len(first_row) >= 2:
            c0 = first_row[0].strip()
            c1 = first_row[1].strip() if len(first_row) > 1 else ""
            # 数据式：首列是编号，第二列是描述
            if (part_id_pattern.match(c0) and c1 and not c1.isdigit() and
                    not any(k in c1 for k in method_table_markers + result_section_markers + ['结论', 'Conclusion', '结果', 'Result', '方法', 'Method'])):
                is_part_table = True
            # 表头式：No./序号/Part No. + 描述
            elif _is_header_like_part_table(first_row):
                is_part_table = True
                start_row_idx = 1

        if is_part_table:
            numeric_rows = 0
            for row in table.rows[start_row_idx:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                if part_id_pattern.match(val):
                    _update_max(val)
                    numeric_rows += 1
            if numeric_rows > 0:
                found_result_part_info = True

        # ===== C. 邻苯类组合编号表：Test Items | Result(mg/kg) ... 第二行为组合编号 =====
        if (len(table.rows) >= 3 and
                any(k in first_row_text for k in ['Test Items', '测试项目']) and
                'Result' in first_row_text):
            sub_header = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[1])]
            if sub_header and ('Test Items' in sub_header[0] or '测试项目' in sub_header[0]):
                found_nums = False
                for cell in sub_header:
                    nums = re.findall(r'\b\d+\b', cell)
                    if nums:
                        found_nums = True
                        for n in nums:
                            _update_max(n)
                if found_nums:
                    found_result_part_info = True

    if found_result_part_info:
        count = max_serial if max_serial > 0 else desc_count
        if count > 2000:
            return "样品部件总数异常请人工检查"
        # 有结果表但序号/描述列均无有效值，视为单一样品的汇总结果
        if count == 0:
            return 1
        return count

    # 退到独立的"样品描述"模块
    count = extract_sample_count_from_sample_description_section(doc)
    if count > 0:
        if count > 2000:
            return "样品部件总数异常请人工检查"
        return count

    # 最终兜底：存在结果表但无部件拆分，视为 1 个样品
    if has_result_table_no_parts:
        return 1

    return "未检测到"


def extract_sample_count_from_sample_description_section(doc):
    """
    从独立的"样品描述"模块（段落标题）中提取最大序号。
    排除包含"照片/Photo"的段落。
    """
    body = doc.element.body
    in_section = False
    target_table = None
    section_paragraphs = []

    for child in body:
        tag = child.tag.split('}')[-1]
        text = ''.join(child.itertext()).strip()

        if tag == 'p':
            lower = text.lower()
            if '样品描述' in text or 'sample description' in lower:
                if '照片' in text or 'photo' in lower:
                    continue
                in_section = True
                continue
            if in_section and any(k in text for k in ['检测方法', '测试结果', '检测结果', '测试方法',
                                                        '检测项目', '测试项目', '要求', '结论']):
                break

        if in_section:
            if tag == 'p':
                section_paragraphs.append(text)
            elif tag == 'tbl' and target_table is None:
                target_table = child
                break

    max_serial = 0
    found = False

    def _try_match(s):
        nonlocal max_serial, found
        s = s.strip()
        for pattern in [r'^(\d+)\s*[:：.\)、]\s*', r'^(\d+)$']:
            m = re.match(pattern, s)
            if m:
                found = True
                seq = int(m.group(1))
                if seq > max_serial:
                    max_serial = seq
                return True
        return False

    if target_table is not None:
        for tc in target_table.iter():
            if tc.tag.endswith('tc'):
                cell_text = ''.join(tc.itertext()).strip()
                _try_match(cell_text)
    else:
        for text in section_paragraphs:
            for line in re.split(r'[\r\n]+', text):
                _try_match(line.strip())

    # 若段落标题方式没找到，再扫描全文档中表头为"样品描述"的表格
    if max_serial == 0:
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
            first_row_text = ' '.join(first_row)
            if '样品描述' not in first_row_text and 'sample description' not in first_row_text.lower():
                continue
            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                m = re.match(r'^(\d+)', val)
                if m:
                    seq = int(m.group(1))
                    if seq > max_serial:
                        max_serial = seq

    return max_serial
def is_standard_method(method):
    """判断是否为标准方法（包含标准前缀或年份）"""
    if not method:
        return False
    std_prefixes = ['GB/T', 'GB', 'IEC', 'ISO', 'ASTM', 'EPA', 'US EPA', 'QC/T', 'QB/T', 'SJ/T']
    if any(prefix in method for prefix in std_prefixes):
        return True
    if re.search(r'\d{4}', method):
        return True
    return False


def extract_test_methods_from_tables(doc):
    """从表格中提取检测方法/仪器"""
    methods = []
    seen = set()

    for table_idx, table in enumerate(doc.tables, 1):
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 识别"检测项目/检测方法/检测仪器"结构的表格
        if '检测项目' in first_row_text and ('检测方法' in first_row_text or '检测仪器' in first_row_text):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到检测项目/方法/仪器表格")

            item_col = find_column_index(first_row, ["检测项目", "测试项目", "项目"])
            method_col = find_column_index(first_row, ["检测方法", "测试方法", "方法"])
            instrument_col = find_column_index(first_row, ["检测仪器", "测试仪器", "仪器"])

            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) < 2:
                    continue

                item = cells[item_col] if item_col is not None and item_col < len(cells) else ""
                method = cells[method_col] if method_col is not None and method_col < len(cells) else ""
                instrument = cells[instrument_col] if instrument_col is not None and instrument_col < len(cells) else ""

                if not item or item == '<br />' or item.isdigit():
                    continue
                if 'Limit' in item or '△' in item or 'Note' in item or '注' in item:
                    continue
                if item in seen:
                    continue
                seen.add(item)

                methods.append({
                    "检测项目": item,
                    "检测方法": method,
                    "检测仪器": instrument
                })

        # 识别"仪器名称/型号/校准有效日期"结构的表格
        elif is_instrument_table(first_row):
            if VERBOSE:
                print(f"  表格 {table_idx}: 找到仪器明细表格")

            name_col = find_column_index(first_row, ["仪器名称", "仪器", "设备名称", "设备"])
            model_col = find_column_index(first_row, ["型号", "规格型号"])

            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) < 2:
                    continue

                instrument = cells[name_col] if name_col is not None and name_col < len(cells) else ""
                model = cells[model_col] if model_col is not None and model_col < len(cells) else ""

                if not instrument or instrument == '<br />':
                    continue

                dedup_key = f"{instrument}_{model}"
                if dedup_key in seen:
                    continue
                seen.add(dedup_key)

                full_instrument = instrument
                if model:
                    full_instrument += f" ({model})"

                methods.append({
                    "检测项目": "",
                    "检测方法": "",
                    "检测仪器": full_instrument
                })

    return methods


def extract_test_methods_from_paragraphs(doc):
    """从段落中提取检测方法（兜底）"""
    methods = []
    method_patterns = [
        r'参考\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'依据\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'采用\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
        r'按照\s*([A-Za-z0-9\s\-/:]+(?:\s*:\s*\d{4})?)',
    ]

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        for pattern in method_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                method = match.strip()
                if is_standard_method(method) and method not in [m["检测方法"] for m in methods]:
                    methods.append({
                        "检测项目": "",
                        "检测方法": method,
                        "检测仪器": ""
                    })

    return methods


def extract_test_items_from_result_tables(doc):
    """从检测结果表格中提取检测项目（兜底，用于填充检测项目列）"""
    items = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        if '检测项目' not in first_row_text:
            continue

        item_col = find_column_index(first_row, ["检测项目", "测试项目", "项目", "Test Item"])
        if item_col is None:
            continue

        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if item_col >= len(cells):
                continue

            item = cells[item_col].strip()
            if not item or item == '<br />' or item.isdigit():
                continue
            if 'Limit' in item or '△' in item or 'Note' in item or '注' in item:
                continue
            if item in seen:
                continue
            seen.add(item)
            items.append(item)

    return items


def extract_test_methods_long(doc):
    """长文档：从首页检测方法总结表格提取"""
    methods = []

    for table in doc.tables[:5]:
        if len(table.rows) >= 2 and len(table.rows[0].cells) == 2:
            first_cell = get_cell_text(table.rows[0].cells[0])
            if re.match(r'^(QC/T|GB/T|IEC|GB|ISO|ASTM|EPA|US EPA)\s*\d', first_cell, re.IGNORECASE):
                for row in table.rows:
                    cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                    if len(cells) == 2 and cells[0] and cells[1]:
                        full_method = f"{cells[0]} {cells[1]}"
                        methods.append({
                            "检测项目": "",
                            "检测方法": full_method,
                            "检测仪器": ""
                        })
                break

    return methods


def merge_test_methods_by_item(methods):
    """
    按检测项目合并方法/仪器：同一项目出现多次时，
    把多个方法、仪器用" / "合并到一条记录中。
    """
    merged = {}
    for m in methods:
        item = m.get("检测项目", "")
        if not item:
            continue
        if item not in merged:
            merged[item] = {"检测项目": item, "methods": [], "instruments": []}
        method = m.get("检测方法", "")
        instrument = m.get("检测仪器", "")
        if method and method not in merged[item]["methods"]:
            merged[item]["methods"].append(method)
        if instrument and instrument not in merged[item]["instruments"]:
            merged[item]["instruments"].append(instrument)

    result = []
    for item, data in merged.items():
        result.append({
            "检测项目": item,
            "检测方法": " / ".join(data["methods"]),
            "检测仪器": " / ".join(data["instruments"]),
        })
    return result


# ==============================================
# 函数6：完整提取中文报告
# ==============================================
def extract_chinese_report(doc, filename):
    """中文报告完整提取，返回字段字典"""
    doc_type = detect_doc_type(doc)

    if VERBOSE:
        print(f"  📋 中文文档类型: {'长文档' if doc_type == 'long' else '短文档'}")

    # 1. 基础字段
    raw_fields = extract_basic_info(doc)

    result = {}
    for raw_key, raw_val in raw_fields.items():
        std_key = ZH_FIELD_MAPPING.get(raw_key)
        if std_key:
            result[std_key] = raw_val

    # 2. 检测要求/结论
    test_req, test_con = extract_test_requirement_and_conclusion(doc)
    if test_req:
        result["检测要求"] = test_req
    if test_con:
        result["检测结论"] = test_con

    # 2.1 兜底：仍未取到检测结论时，从结果表的“结论”列提取
    if not result.get("检测结论", "").strip():
        fallback_con = extract_conclusion_from_result_column_zh(doc)
        if fallback_con:
            result["检测结论"] = fallback_con

    # 3. 报告编号从文件名取
    m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z]?\d*)', filename, re.IGNORECASE)
    if m:
        result["报告编号"] = m.group(1)

    # 4. 样品部件总数
    sample_count = extract_sample_count(doc)
    result["样品部件总数"] = str(sample_count)

    # 5. 检测方法：优先直接读取"检测项目/检测方法/检测仪器"表，无命中再按长文档方法总结兜底
    direct_methods = extract_test_methods_from_tables(doc)
    if direct_methods:
        test_methods = direct_methods
    elif doc_type == "long":
        test_methods = extract_test_methods_long(doc)
    else:
        test_methods = []
        paragraph_methods = extract_test_methods_from_paragraphs(doc)

        # 如果表格方法缺少检测方法，用段落方法补全
        for tm in test_methods:
            if not tm["检测方法"]:
                for pm in paragraph_methods:
                    if pm["检测方法"]:
                        tm["检测方法"] = pm["检测方法"]
                        break

        # 如果段落方法没有被使用，添加为新条目
        used_methods = {tm["检测方法"] for tm in test_methods}
        for pm in paragraph_methods:
            if pm["检测方法"] and pm["检测方法"] not in used_methods:
                test_methods.append(pm)

    # 如果方法条目缺少检测项目，尝试从检测结果表补全
    if test_methods:
        test_items = extract_test_items_from_result_tables(doc)
        for i, item in enumerate(test_items):
            if i < len(test_methods) and not test_methods[i]["检测项目"]:
                test_methods[i]["检测项目"] = item

    # 按检测项目合并重复方法/仪器（例如六价铬两种方法合并为一条）
    test_methods = merge_test_methods_by_item(test_methods)

    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]

    return result, test_methods


# ==============================================
# 主函数：处理单个文档
# ==============================================
def process_word_file(input_path, output_excel):
    """
    处理单个中文 Word 检测报告，提取信息并生成 Excel

    参数:
        input_path: Word 文档路径
        output_excel: 输出 Excel 路径
    """
    from docx import Document

    print("=" * 60)
    print(f"开始处理: {os.path.basename(input_path)}")
    print("=" * 60)

    # 1. 检查文件是否存在
    if not os.path.exists(normalize_path(input_path)):
        print(f"❌ 错误：文件不存在 - {input_path}")
        return False

    # 2. 修复文档（如果需要）
    try:
        actual_file = fix_docx_if_needed(input_path)
    except Exception as e:
        print(f"❌ 错误：无法打开文档 - {e}")
        print("   请检查文件是否为有效的 Word 文档(.docx格式)")
        return False

    # 3. 打开文档
    try:
        doc = Document(actual_file)
    except ImportError:
        print("❌ 错误：未安装 python-docx 库")
        print("   请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 错误：打开文档失败 - {e}")
        return False

    if VERBOSE:
        print(f"\n📄 文档信息:")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")

    # 4. 初始化结果字典
    result = {
        "文件名": os.path.basename(input_path),
        "报告编号": "",
        "申请商": "",
        "申请商地址": "",
        "制造商地址": "",
        "样品名称": "",
        "零件号": "",
        "参考零件号": "",
        "材质": "",
        "车型": "",
        "生产日期": "",
        "供应商代码": "",
        "测试类型": "",
        "样品编号": "",
        "样品接收日期": "",
        "检测期间": "",
        "检测要求": "",
        "检测结论": "",
        "报告日期": "",
        "样品部件总数": "",
        "备注": "",
    }

    # 添加检测方法列
    for i in range(1, MAX_TEST_METHODS + 1):
        result[f"检测项目{i}"] = ""
        result[f"检测方法{i}"] = ""
        result[f"检测仪器{i}"] = ""

    # 5. 提取
    extracted, test_methods = extract_chinese_report(doc, os.path.basename(input_path))
    result.update(extracted)

    # 6. 生成 Excel
    print(f"\n💾 生成 Excel 文件...")

    col_order = [
        "报告编号", "申请商", "申请商地址", "制造商地址",
        "样品名称", "零件号", "参考零件号", "材质",
        "车型", "生产日期", "供应商代码", "测试类型",
        "样品编号", "样品接收日期", "检测期间",
        "检测要求", "检测结论", "报告日期",
        "样品部件总数", "备注"
    ]

    actual_method_count = min(len(test_methods), MAX_TEST_METHODS)
    for i in range(1, actual_method_count + 1):
        col_order.append(f"检测项目{i}")
        col_order.append(f"检测方法{i}")
        col_order.append(f"检测仪器{i}")

    col_order.extend(["文件名"])

    df = pd.DataFrame([result])
    for col in col_order:
        if col not in df.columns:
            df[col] = ""
    df = df[col_order]

    try:
        df.to_excel(output_excel, index=False)
        print(f"✅ 成功保存: {output_excel}")
    except ImportError:
        print("❌ 错误：未安装 openpyxl 库")
        print("   请运行: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 错误：保存 Excel 失败 - {e}")
        return False

    # 统计信息
    non_empty = sum(1 for val in result.values() if val)
    print(f"\n📊 提取完成！共提取 {non_empty} 个非空字段")
    print("=" * 60)

    return True


# ==============================================
# 程序入口
# ==============================================
def main():
    print("\n" + "=" * 60)
    print("  中文检测报告信息提取工具 Plus 4")
    print("=" * 60)

    # 检查依赖
    try:
        import docx
        import pandas
    except ImportError as e:
        print(f"\n❌ 缺少依赖库: {e}")
        print("请运行以下命令安装依赖:")
        print("  pip install python-docx pandas openpyxl")
        sys.exit(1)

    # 处理文件
    success = process_word_file(INPUT_FILE, OUTPUT_EXCEL)

    if success:
        print("\n🎉 处理完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)


if __name__ == "__main__":
    main()
