# -*- coding: utf-8 -*-
"""
检测要求和检测结论字段专属提取方案
功能：仅重新提取 Word 检测报告中的"检测要求"和"检测结论"两个字段，
      输出新 Excel，通过"文件路径"列与原汇总结果一一对应，便于复制回填。
"""

# ============================================================
# 【配置区 - 改这里就行】
# ============================================================

# 输入：现有的汇总结果 Excel（必须包含"文件路径"列）
INPUT_EXCEL = r"D:\Kathy\PDF提取工具\中英文汇总提取结果_20260722_003948全提取2筛选版.xlsx"

# 输出：新的专属提取结果 Excel
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\检测要求和检测结论字段专属提取结果3.xlsx"

# 测试限制：0=跑全部；N>0=只跑前 N 条（用于快速调测）
TEST_LIMIT = 0

# 读取的 sheet 名
SHEET_NAME = "提取结果汇总"

# 是否显示进度条
SHOW_PROGRESS = True

# ============================================================
# 【依赖导入】
# ============================================================
import os
import re
import html
import pandas as pd
from docx import Document
from tqdm import tqdm


# ============================================================
# 【工具函数】
# ============================================================
def get_cell_text(cell):
    """读取单元格完整文本（含 SDT 内容控件）"""
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    return html.unescape(result).strip()


def get_row_cells_text_fast(row):
    """读取一行所有单元格文本"""
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cells_text.append(html.unescape(cell_text).strip())
    return cells_text


def normalize_text(text):
    """文本清洗"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def contains_pass(text):
    """是否包含符合语义"""
    if not text:
        return False
    t = text.lower()
    return any(k in t for k in ['符合', '合格', '通过', 'pass', 'comply', 'complies',
                                 'conform', 'conforms', '阴性', '未检出'])


def contains_fail(text):
    """是否包含不符合语义"""
    if not text:
        return False
    t = text.lower()
    return any(k in t for k in ['不符合', '不合格', '不通过', 'fail', 'failed', 'failing'])


def is_nd(text):
    """是否为 N.D.（未检出）"""
    if not text:
        return False
    return text.strip().lower() in ['nd', 'n.d.', 'n.d', '未检出']


def is_numeric_result(val):
    """是否为具体数值结果（按用户要求视为不符合）"""
    if not val or val in ['/', '-', '—', 'P', 'X', '']:
        return False
    if is_nd(val) or contains_pass(val) or contains_fail(val):
        return False
    return bool(re.search(r'\d', val))


def classify_text(text):
    """对整段文本判断：pass / fail / partial / unknown"""
    if not text:
        return 'unknown'
    has_fail = contains_fail(text)
    has_pass = contains_pass(text)
    if has_fail and not has_pass:
        return 'fail'
    if has_fail and has_pass:
        return 'partial'
    if has_pass and not has_fail:
        return 'pass'
    return 'unknown'


def format_fail_details(items):
    """格式化不符合项列表"""
    seen = set()
    details = []
    for item, val in items:
        key = f"{item}|{val}"
        if key in seen:
            continue
        seen.add(key)
        if item:
            if val in ['不符合', 'Fail', 'failed', '不合格', '不通过']:
                details.append(item)
            else:
                details.append(f"{item}({val})")
        else:
            details.append(val)
    return details


def summarize_conclusion_items(items, all_fail_short=False):
    """
    综合多项结论：
    - items: [(项目, 值), ...]
    - all_fail_short=True 时，全部不符合只输出"不符合"
    """
    if not items:
        return ""

    pass_items = [(i, v) for i, v in items
                  if (contains_pass(v) and not contains_fail(v)) or is_nd(v)]
    fail_items = [(i, v) for i, v in items
                  if contains_fail(v) or is_numeric_result(v)]

    if fail_items and not pass_items:
        if all_fail_short:
            return "不符合"
        details = format_fail_details(fail_items)
        if details:
            return "不符合，不符合的内容有：" + "、".join(details)
        return "不符合"

    if fail_items:
        details = format_fail_details(fail_items)
        if details:
            return "部分项不符合，不符合的内容有：" + "、".join(details)
        return "部分项不符合"

    if pass_items:
        return "符合"

    # 只有未知值，保留原文
    return " / ".join(set(v for _, v in items))


# ============================================================
# 【检测要求提取】
# ============================================================
def extract_requirement(doc):
    """提取检测要求"""
    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row)

        # 单列表格
        if len(first_row) == 1:
            header = first_row[0]
            if '检测要求' in header or 'Test Requirement' in header:
                lines = []
                for row in table.rows[1:]:
                    cell_text = get_cell_text(row.cells[0]).strip()
                    if cell_text:
                        lines.append(cell_text)
                return '\n'.join(lines)

        # 2行2列混合表
        if len(first_row) == 2:
            has_req = ('检测要求' in first_row_text or 'Test Requirement' in first_row_text)
            has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text)
            if has_req and has_con and len(table.rows) >= 2:
                return get_cell_text(table.rows[1].cells[0]).strip()

    return ""


# ============================================================
# 【检测结论提取】
# ============================================================
def extract_conclusion_from_filename(filename):
    """从文件名判断结论"""
    name = os.path.splitext(filename)[0]
    if '不合格' in name or '不符合' in name:
        return '不符合'
    if '合格' in name or '符合' in name:
        return '符合'
    return ""


def extract_conclusion_from_independent_table(doc):
    """从独立检测结论表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 1:
            continue

        header = first_row[0]
        if '检测结论' not in header and 'Test Conclusion' not in header:
            continue

        lines = []
        for row in table.rows[1:]:
            val = get_cell_text(row.cells[0]).strip()
            if val:
                lines.append(val)

        full_text = '\n'.join(lines)
        if not full_text:
            return ""

        cat = classify_text(full_text)
        if cat == 'pass':
            return "符合"
        if cat == 'fail':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)
        if cat == 'partial':
            items = [("", v) for v in lines]
            return summarize_conclusion_items(items, all_fail_short=True)

        # 无法判断，保留原文
        return full_text

    return ""


def extract_conclusion_from_mixed_table(doc):
    """从2行2列检测要求/结论混合表格提取"""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = get_row_cells_text_fast(table.rows[0])
        if len(first_row) != 2:
            continue

        first_row_text = ' | '.join(first_row)
        has_req = ('检测要求' in first_row_text or 'Test Requirement' in first_row_text)
        has_con = ('结论' in first_row_text or 'Conclusion' in first_row_text)
        if not (has_req and has_con):
            continue

        con_col = 1
        if '结论' in first_row[0] or 'Conclusion' in first_row[0]:
            con_col = 0

        items = []
        for row in table.rows[1:]:
            cells = get_row_cells_text_fast(row)
            if con_col < len(cells):
                val = cells[con_col].strip()
                if val:
                    items.append(("", val))

        if items:
            summary = summarize_conclusion_items(items, all_fail_short=True)
            if summary:
                return summary
            return "\n".join([v for _, v in items])

    return ""


def extract_conclusion_from_result_tables(doc):
    """从检测结果表结论列/行综合判断"""
    all_items = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' | '.join(first_row)
        if len(first_row) == 1:
            continue

        is_result = any(k in first_row_text for k in ['结果', 'Result', '结论',
                                                       'Conclusion', '限值', 'Limit'])
        if not is_result:
            continue

        # 结论列
        con_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['结论', 'Conclusion']):
                con_col = idx
                break

        # 检测项目列
        item_col = None
        for idx, cell in enumerate(first_row):
            if any(k in cell for k in ['检测项目', 'Test Item', 'Item', '测试项目']):
                item_col = idx
                break

        current_item = ""
        table_rows = list(table.rows)

        for row_idx, row in enumerate(table_rows[1:], start=1):
            cells = get_row_cells_text_fast(row)
            if not cells:
                continue

            first_cell = cells[0].strip() if len(cells) > 0 else ""
            is_conclusion_row = first_cell in ['结论', 'Conclusion']

            # 更新当前项目（处理合并单元格），结论行不更新
            if item_col is not None and item_col < len(cells) and not is_conclusion_row:
                item_val = cells[item_col].strip()
                if item_val and item_val not in ['检测项目', 'Test Item', 'Item',
                                                  '测试项目', '结论', 'Conclusion']:
                    current_item = item_val

            # 结论列
            if con_col is not None and con_col < len(cells):
                val = cells[con_col].strip()
                if val and val not in ['/', '-', '—']:
                    item_for = current_item if current_item not in ['结论', 'Conclusion'] else ""
                    all_items.append((item_for, val))

            # 结论行
            if is_conclusion_row:
                for idx, val in enumerate(cells[1:], start=1):
                    val = val.strip()
                    if not val or val in ['/', '-', '—']:
                        continue

                    target_item = ""
                    # 向上扫描该列，找结果列有具体数字的项目
                    for prev_row in reversed(table_rows[1:row_idx]):
                        prev_cells = get_row_cells_text_fast(prev_row)
                        if idx < len(prev_cells):
                            prev_res = prev_cells[idx].strip()
                            prev_item = (prev_cells[item_col].strip()
                                         if item_col is not None and item_col < len(prev_cells)
                                         else "")
                            if is_numeric_result(prev_res) and prev_item \
                                    and prev_item not in ['结论', 'Conclusion']:
                                target_item = prev_item
                                break

                    if not target_item:
                        target_item = current_item if current_item not in ['结论', 'Conclusion'] else ""

                    all_items.append((target_item, val))

    if all_items:
        return summarize_conclusion_items(all_items, all_fail_short=True)
    return ""


def extract_conclusion(doc, filename):
    """按优先级提取检测结论"""
    # 1. 文件名
    con = extract_conclusion_from_filename(filename)
    if con:
        return con

    # 2. 独立检测结论表格
    con = extract_conclusion_from_independent_table(doc)
    if con:
        return con

    # 3. 2行2列混合表格
    con = extract_conclusion_from_mixed_table(doc)
    if con:
        return con

    # 4. 从检测结果表综合判断
    con = extract_conclusion_from_result_tables(doc)
    if con:
        return con

    return ""


# ============================================================
# 【主程序】
# ============================================================
def main():
    print("=" * 70)
    print("检测要求和检测结论字段专属提取方案")
    print(f"输入：{INPUT_EXCEL}")
    print(f"输出：{OUTPUT_EXCEL}")
    if TEST_LIMIT > 0:
        print(f"[测试模式] 只处理前 {TEST_LIMIT} 条")
    print("=" * 70)

    # 1. 读取汇总 Excel
    print("\n[1/4] 读取汇总 Excel...")
    df = pd.read_excel(INPUT_EXCEL, sheet_name=SHEET_NAME)
    total = len(df)
    print(f"      共 {total} 行")

    if '文件路径' not in df.columns:
        print("[错误] Excel 中不存在'文件路径'列")
        return

    file_paths = df['文件路径'].dropna().tolist()
    if TEST_LIMIT > 0:
        file_paths = file_paths[:TEST_LIMIT]

    # 2. 批量处理
    print("\n[2/4] 开始提取检测要求和检测结论...")
    results = []

    iterator = tqdm(file_paths, desc="提取进度", unit="份") if SHOW_PROGRESS else file_paths
    for fp in iterator:
        try:
            doc = Document(fp)
            req = extract_requirement(doc)
            con = extract_conclusion(doc, os.path.basename(fp))
            results.append({
                "文件路径": fp,
                "检测要求（新）": req,
                "检测结论（新）": con,
            })
        except Exception as e:
            results.append({
                "文件路径": fp,
                "检测要求（新）": "",
                "检测结论（新）": f"提取失败: {str(e)[:80]}",
            })

    # 3. 保存结果
    print("\n[3/4] 保存结果 Excel...")
    df_out = pd.DataFrame(results)
    df_out = df_out[["文件路径", "检测要求（新）", "检测结论（新）"]]
    df_out.to_excel(OUTPUT_EXCEL, index=False)

    # 4. 统计
    print("\n[4/4] 统计信息：")
    print(f"      处理文件数：{len(results)}")
    print(f"      检测要求非空：{df_out['检测要求（新）'].astype(str).str.strip().ne('').sum()}")
    print(f"      检测结论非空：{df_out['检测结论（新）'].astype(str).str.strip().ne('').sum()}")
    print(f"      输出文件：{OUTPUT_EXCEL}")
    print("\n[完成] 可直接按'文件路径'列 VLOOKUP 或复制填回原汇总表。")


if __name__ == "__main__":
    main()
