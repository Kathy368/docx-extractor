# -*- coding: utf-8 -*-
"""
英文检测报告信息提取工具 4
功能：从英文 RoHS/ELV/REACH 检测报告 Word 文档中提取关键信息，生成 Excel 表格

改进点：
1. 支持自动修复部分损坏的 docx 文件（NULL 关系、Bad CRC-32 等）
2. 新增无序号标注时的样品部件总数统计（按 Test item(s) 行数）
3. 新增图片型英文报告的 OCR 识别兜底方案

OCR 依赖（可选）：
- 首选：Tesseract OCR + pytesseract + Pillow
  安装：pip install pytesseract Pillow
  并下载安装 Tesseract：https://github.com/UB-Mannheim/tesseract/wiki
- 备选：easyocr 或 paddleocr

基础依赖库：python-docx, pandas, openpyxl
安装命令：pip install python-docx pandas openpyxl
"""
import os
import re
import sys
import zipfile
import shutil
import tempfile
import html
import pandas as pd

# ==============================================
# 【配置区】请根据实际情况修改以下参数
# ==============================================
# 待处理的 Word 文档路径（支持 Windows 路径，如 r"D:\Kathy\PDF提取工具\test.docx"）
INPUT_FILE = r"G:\中心实验室\中心实验室-新(2023.07.31)\002 报告\报告组\常规组报告 2017.11月份开始\待审核报告 新\非汽车事业部\A 傲川\A 傲川 2023.4.4\S23040400301C-硅元素.docx"
# 输出 Excel 文件路径
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\英测2提取结果.xlsx"
# 最大检测方法数量（预留列数）
MAX_TEST_METHODS = 100
# 是否显示详细处理日志
VERBOSE = True


# ==============================================
# 【英文字段映射表】英文表头 → 中文含义
# ==============================================
EN_FIELD_MAPPING = {
    # 申请商相关
    "Applicant": "申请商",
    "Client": "申请商",
    "Company": "申请商",
    "Address": "申请商地址",

    # 样品相关
    "Product name": "样品名称",
    "Product Name": "样品名称",
    "Sample name": "样品名称",
    "Sample Name": "样品名称",
    "Sample Description": "样品名称",
    "Model": "零件号",
    "Part No.": "零件号",
    "Part Number": "零件号",
    "Model No.": "零件号",
    "Test model": "零件号",
    "Client Ref. Info.": "客户参考信息",
    "Client Reference": "客户参考信息",
    "Material": "材质",
    "Production Date": "生产日期",
    "Manufacturer": "制造商",
    "Manufacturer & Factory": "制造商",
    "Factory": "制造商",
    "Trade mark": "商标",
    "Trade Mark": "商标",

    # 样品编号相关
    "Sample No.": "样品编号",
    "Sample Number": "样品编号",
    "Sample ID": "样品编号",

    # 日期相关
    "Sample Received Date": "样品接收日期",
    "Date of Receipt": "样品接收日期",
    "Date Received": "样品接收日期",
    "Testing Period": "检测期间",
    "Test Period": "检测期间",
    "Test Date": "检测期间",
    "Date": "报告日期",
    "Report Date": "报告日期",
    "Issue Date": "报告日期",

    # 人员相关
    "Compiled by": "编制",
    "Prepared by": "编制",
    "Reviewed by": "审核",
    "Checked by": "审核",
    "Approved by": "批准",

    # 其他
    "Note": "说明",
    "Notes": "说明",
    "Remark": "备注",
    "Remarks": "备注",
}


# ==============================================
# 工具函数：路径处理（解决 Windows 长路径问题）
# ==============================================
def normalize_path(file_path):
    """为 Windows 长路径添加 \\\\?\\\\ 前缀"""
    prefix = r"\\\\?\\"
    if sys.platform == 'win32' and not file_path.startswith(prefix):
        abs_path = os.path.abspath(file_path)
        if len(abs_path) > 230:
            return prefix + abs_path
    return file_path


# ==============================================
# 工具函数：修复有问题的 docx 文件
# ==============================================
def fix_docx_if_needed(file_path):
    """
    检查并修复有问题的 docx 文件
    支持：NULL 关系问题、Bad CRC-32 等 zip 损坏
    返回修复后的文件路径，如果不需要修复则返回原路径
    """
    file_path = normalize_path(file_path)

    def _try_open(path):
        from docx import Document
        try:
            Document(path)
            return True
        except Exception:
            return False

    if _try_open(file_path):
        return file_path

    if VERBOSE:
        print(f"  ⚠️  检测到文档损坏，正在自动修复...")

    temp_dir = tempfile.gettempdir()
    prefix = r"\\\\?\\"
    base_name = os.path.basename(file_path.replace(prefix, ''))
    fixed_path = os.path.join(temp_dir, f"fixed_{base_name}")

    try:
        with zipfile.ZipFile(file_path, 'r') as zin:
            with zipfile.ZipFile(fixed_path, 'w') as zout:
                for item in zin.infolist():
                    if item.filename == 'word/_rels/document.xml.rels':
                        try:
                            content = zin.read(item.filename).decode('utf-8')
                            content = re.sub(r'<Relationship[^>]*Target="[^"]*NULL[^"]*"[^>]*/>', '', content)
                            zout.writestr(item, content)
                        except zipfile.BadZipFile:
                            continue
                        continue

                    try:
                        data = zin.read(item.filename)
                        zout.writestr(item, data)
                    except zipfile.BadZipFile:
                        if item.filename.startswith('word/media/'):
                            if VERBOSE:
                                print(f"    跳过损坏的图片: {item.filename}")
                            continue
                        raise

        if _try_open(fixed_path):
            if VERBOSE:
                print(f"  ✅ 文档修复成功")
            return fixed_path
        else:
            raise Exception("修复后仍无法打开文档")
    except Exception as fix_error:
        if VERBOSE:
            print(f"  ❌ 文档修复失败: {fix_error}")
        raise


# ==============================================
# 工具函数：获取单元格完整文本（包括 SDT 内容控件）
# ==============================================
def get_cell_text(cell):
    """
    获取单元格的完整文本，包括 SDT（结构化文档标签/内容控件）中的内容
    """
    xml = cell._tc.xml
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    result = ''.join(texts)
    result = html.unescape(result)
    return result.strip()


def get_row_cells_text_fast(row):
    """
    获取一行中所有单元格的文本，包括被 SDT（结构化文档标签）包裹的单元格
    """
    xml = row._tr.xml
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        cell_text = html.unescape(cell_text)
        cells_text.append(cell_text.strip())
    return cells_text


# 兼容总提取代码的别名
get_row_cells_text = get_row_cells_text_fast


def normalize_text(text):
    """清理文本：去首尾空、统一换行/空格"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'[\r\n]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def normalize_key(key):
    """清理字段名：去冒号、去首尾空、统一空格"""
    if not key:
        return ""
    key = normalize_text(key)
    key = key.rstrip('：:').strip()
    return key


# ==============================================
# 函数1：提取基本信息
# ==============================================
def find_column_index(first_row, keywords):
    """根据关键词列表找到列索引"""
    for idx, cell in enumerate(first_row):
        for kw in keywords:
            if kw in cell:
                return idx
    return None


def extract_basic_info(doc):
    """从两列表格中提取英文报告基础字段"""
    fields = {}
    manufacturer_address_found = False

    for table_idx, table in enumerate(doc.tables[:20]):
        if len(table.rows) < 1:
            continue

        first_row = get_row_cells_text_fast(table.rows[0])
        first_row_text = ' '.join(first_row)
        first_cell_count = len(first_row)

        # 跳过检测结果表格
        if 'Part No.' in first_row_text or 'Test Items' in first_row_text or 'Test item(s)' in first_row_text:
            continue

        # 跳过检测方法表格
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue

        # 跳过 XRF 限值表格
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue

        # 跳过检测要求/结论表格（单独处理）
        if 'Test Requirement' in first_row_text or 'Test Conclusion' in first_row_text:
            continue

        # 处理两列表格（键值对）
        if first_cell_count == 2:
            for row_idx, row in enumerate(table.rows):
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]

                # 处理只有 1 个单元格的行
                if len(cells) == 1:
                    key = normalize_key(cells[0])
                    if key and len(key) < 60:
                        fields[key] = ""
                    continue

                if len(cells) == 2:
                    key = normalize_key(cells[0])
                    val = cells[1].strip()
                    key = re.sub(r'\s+', ' ', key).strip()

                    if key and len(key) < 60:
                        # 跳过说明性文字
                        if 'The following sample' in key or 'was/were submitted' in key:
                            continue

                        # 特殊处理：遇到 Manufacturer/Factory 后，下一个 Address 是制造商地址
                        if key in ['Manufacturer', 'Manufacturer & Factory', 'Factory']:
                            manufacturer_address_found = True
                            if val:
                                fields[key] = val
                            continue

                        # 特殊处理 Address
                        if key == 'Address':
                            if not manufacturer_address_found:
                                if val:
                                    fields[key] = val
                            else:
                                if val and '制造商地址' not in fields:
                                    fields['制造商地址'] = val
                            continue

                        if val:
                            fields[key] = val

        # 处理四列表格（编制/审核/批准/日期等）
        elif first_cell_count == 4:
            for row in table.rows:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if len(cells) == 4:
                    for i in range(0, 4, 2):
                        key = normalize_key(cells[i])
                        val = cells[i+1].strip()
                        if key and len(key) < 40:
                            if val:
                                fields[key] = val

    return fields


# ==============================================
# 函数2：提取检测要求和检测结论（完整原文）
# ==============================================
def extract_test_requirement_and_conclusion(doc):
    """提取 Test Requirement 和 Test Conclusion 的完整原文（优先从首页提取）"""
    test_requirement = ""
    test_conclusion = ""

    for table_idx, table in enumerate(doc.tables):
        if test_requirement and test_conclusion:
            break

        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c) for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)

        # 情况1：2行2列表格（Test Requirement 和 Conclusion 在同一行）
        if (len(table.rows[0].cells) == 2 and
            'Test Requirement' in first_row_text and
            'Conclusion' in first_row_text):
            if VERBOSE:
                print(f"  找到检测要求/结论表格（2行2列格式，首页）")
            if len(table.rows) >= 2:
                test_requirement = get_cell_text(table.rows[1].cells[0])
                test_conclusion = get_cell_text(table.rows[1].cells[1])
            break

        # 情况2：2行1列表格（Test Requirement 单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Requirement' in first_row_text:
            if not test_requirement:
                if VERBOSE:
                    print(f"  找到检测要求表格（独立表格）")
                if len(table.rows) >= 2:
                    test_requirement = get_cell_text(table.rows[1].cells[0])

        # 情况3：2行1列表格（Test Conclusion 单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Conclusion' in first_row_text:
            if not test_conclusion:
                if VERBOSE:
                    print(f"  找到检测结论表格（独立表格）")
                if len(table.rows) >= 2:
                    test_conclusion = get_cell_text(table.rows[1].cells[0])

    return test_requirement, test_conclusion


# ==============================================
# 函数3：提取样品部件总数
# ==============================================
def extract_sample_count(doc):
    """
    提取样品部件总数（成功4 修复版）：
    1. 优先在 Test Result(s) 模块内寻找：
       - 有 Part No / Sample No. 列时，取最大基号；同时用 Part Description / Sample Description 去重行数做二次校验。
       - 无序号列时，只按 Part Description / Sample Description 去重行数统计。
    2. 识别无表头部件表（首行首列为数字编号，第二列为描述）。
    3. 识别邻苯类组合编号表（如 1+9+19+25+26+33）。
    4. 排除 CAS No. 被误当成序号列。
    5. Test Result(s) 模块没有部件信息时，退到独立的 Sample Description 模块，取最大序号。
    """
    max_serial = 0
    desc_count = 0
    found_result_part_info = False

    result_section_markers = ['Test Result(s)', 'Test Result', 'Test Results', 'Result(s),mg/kg', '结果,mg/kg']
    result_markers = ['Result', 'Limit', 'Conclusion', 'mg/kg']
    test_item_keywords = {'pb', 'cd', 'hg', 'cr', 'br', 'lead', 'cadmium', 'mercury', 'chromium', 'brominated'}
    serial_keywords = ['Part No', 'Sample No', 'No.']
    desc_keywords = ['Part Description', 'Sample Description', 'Description']

    part_id_pattern = re.compile(r'^(\d+)(-\d+)?[A-Z]?\)?$')

    def _base_num(val):
        m = re.match(r'^(\d+)', str(val).strip())
        return int(m.group(1)) if m else None

    def _update_max(val):
        nonlocal max_serial
        n = _base_num(val)
        if n and n > max_serial:
            max_serial = n

    def _is_header_like_part_table(first_row):
        if len(first_row) < 2:
            return False
        h0 = first_row[0]
        h1 = first_row[1]
        has_no = any(k in h0 for k in ['No.', 'No', '序号', 'Part No', 'Sample No'])
        has_desc = any(k in h1 for k in ['Description', '描述', 'Part', 'Sample'])
        return has_no and has_desc

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)
        first_row_lower = first_row_text.lower()

        # 排除方法/限值表
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue

        # ===== A. 带表头的 Test Result(s) 表 =====
        is_result_table = (any(k in first_row_text for k in result_section_markers) or
                           any(k in first_row_text for k in result_markers) or
                           any(kw in first_row_lower for kw in test_item_keywords))
        if is_result_table:
            serial_col = find_column_index(first_row, serial_keywords)
            desc_col = find_column_index(first_row, desc_keywords)

            # 排除 CAS No. 被当成序号列
            if serial_col is not None:
                header = first_row[serial_col]
                if 'CAS' in header.upper():
                    serial_col = None

            if serial_col is not None or desc_col is not None:
                found_result_part_info = True

                if serial_col is not None:
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if serial_col >= len(cells):
                            continue
                        val = cells[serial_col].strip()
                        _update_max(val)

                if desc_col is not None:
                    descs = set()
                    for row in table.rows[1:]:
                        cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                        if desc_col >= len(cells):
                            continue
                        d = cells[desc_col].strip()
                        if d and d not in desc_keywords and not d.isdigit():
                            descs.add(d)
                    if len(descs) > desc_count:
                        desc_count = len(descs)

        # ===== B. 无表头部件表 或 No./序号 头部件表 =====
        is_part_table = False
        start_row_idx = 0
        if len(first_row) >= 2:
            c0 = first_row[0].strip()
            c1 = first_row[1].strip() if len(first_row) > 1 else ""
            if (part_id_pattern.match(c0) and c1 and not c1.isdigit() and
                    not any(k in c1.lower() for k in ['conclusion', 'result', 'method', 'limit', 'note'])):
                is_part_table = True
            elif _is_header_like_part_table(first_row):
                is_part_table = True
                start_row_idx = 1

        if is_part_table:
            numeric_rows = 0
            for row in table.rows[start_row_idx:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                if part_id_pattern.match(val):
                    _update_max(val)
                    numeric_rows += 1
            if numeric_rows > 0:
                found_result_part_info = True

        # ===== C. 邻苯类组合编号表 =====
        if (len(table.rows) >= 3 and
                any(k in first_row_text for k in ['Test Items', '测试项目']) and
                'Result' in first_row_text):
            sub_header = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[1])]
            if sub_header and ('Test Items' in sub_header[0] or '测试项目' in sub_header[0]):
                found_nums = False
                for cell in sub_header:
                    nums = re.findall(r'\b\d+\b', cell)
                    if nums:
                        found_nums = True
                        for n in nums:
                            _update_max(n)
                if found_nums:
                    found_result_part_info = True

    if found_result_part_info:
        return max_serial if max_serial > 0 else desc_count

    # 退到独立的 Sample Description 模块
    return extract_sample_count_from_sample_description_section(doc)


def extract_sample_count_from_sample_description_section(doc):
    """
    从独立的 Sample Description 模块（段落标题）中提取最大序号。
    排除包含 Photo / 照片的段落。
    """
    body = doc.element.body
    in_section = False
    target_table = None
    section_paragraphs = []

    for child in body:
        tag = child.tag.split('}')[-1]
        text = ''.join(child.itertext()).strip()

        if tag == 'p':
            lower = text.lower()
            if 'sample description' in lower:
                if 'photo' in lower or '照片' in text:
                    continue
                in_section = True
                continue
            if in_section and any(k in lower for k in ['test method', 'test result', 'test item', 'conclusion', 'requirement']):
                break

        if in_section:
            if tag == 'p':
                section_paragraphs.append(text)
            elif tag == 'tbl' and target_table is None:
                target_table = child
                break

    max_serial = 0
    found = False

    def _try_match(s):
        nonlocal max_serial, found
        s = s.strip()
        for pattern in [r'^(\d+)\s*[:：.\)、]\s*', r'^(\d+)$']:
            m = re.match(pattern, s)
            if m:
                found = True
                seq = int(m.group(1))
                if seq > max_serial:
                    max_serial = seq
                return True
        return False

    if target_table is not None:
        for tc in target_table.iter():
            if tc.tag.endswith('tc'):
                cell_text = ''.join(tc.itertext()).strip()
                _try_match(cell_text)
    else:
        for text in section_paragraphs:
            for line in re.split(r'[\r\n]+', text):
                _try_match(line.strip())

    # 若段落标题方式没找到，再扫描全文档中表头为 Sample Description 的表格
    if max_serial == 0:
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
            first_row_text = ' '.join(first_row)
            if 'Sample Description' not in first_row_text and 'sample description' not in first_row_text.lower():
                continue
            for row in table.rows[1:]:
                cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
                if not cells:
                    continue
                val = cells[0].strip()
                m = re.match(r'^(\d+)', val)
                if m:
                    seq = int(m.group(1))
                    if seq > max_serial:
                        max_serial = seq

    return max_serial


def extract_sample_count_from_test_items(doc):
    """
    兜底：当样品没有按 Part No. 拆分时，根据结果表的 Result(s) 列数推断样品数。
    例如 Test item(s) | Limit | Result(s) 只有一列结果 → 1 个样品。
    """
    max_count = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 只关心结果表
        if not ('Result' in first_row_text or 'result' in first_row_text.lower()):
            continue
        # 已带 Part/Sample 编号的表由 extract_sample_count 处理，这里跳过
        if any(k in first_row_text for k in ['Part No', 'Sample No', 'No.']):
            continue

        # 若表头是数字编号（如 1 | 2 | 3），取最大编号
        nums = []
        for cell in first_row:
            cell = cell.strip()
            if re.match(r'^\d+$', cell):
                nums.append(int(cell))
        if nums:
            max_count = max(max_count, max(nums))
            continue

        # 若有 Test item(s) + Result(s) 列，则视为 1 个样品
        has_test_item = any('Test item' in c or 'item' in c.lower() for c in first_row)
        has_result = any('Result' in c for c in first_row)
        if has_test_item and has_result:
            max_count = max(max_count, 1)

    return max_count


def extract_test_methods(doc):
    """从 Chemical Test 章节的检测方法表格中提取"""
    methods = []
    seen_items = set()

    for table_idx, table in enumerate(doc.tables[:30]):
        if len(table.rows) < 3:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        # 检测方法表格特征：有 Test item / Testing item 和 Test method 列
        if (('Test item' not in first_row_text and 'Testing item' not in first_row_text) or
                'Test method' not in first_row_text):
            continue

        if VERBOSE:
            print(f"  找到检测方法表格（{len(table.rows)}行）")

        for row_idx, row in enumerate(table.rows[1:]):
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if len(cells) < 3:
                continue

            item = cells[0].strip()
            method = cells[1].strip() if len(cells) > 1 else ""
            instrument = cells[2].strip() if len(cells) > 2 else ""

            # 跳过说明行和空行
            if not item or not method:
                continue
            if 'Limit' in item or '△' in item or 'Note' in item:
                continue

            # 按检测项目名称去重
            if item in seen_items:
                continue
            seen_items.add(item)

            methods.append({
                "检测项目": item,
                "检测方法": method,
                "检测仪器": instrument
            })

    return methods


# ==============================================
# 函数5：判断是否为图片型文档
# ==============================================
def is_image_based_document(doc):
    """
    判断文档是否为图片型（正文几乎没有可提取文本，如 CTT...EN_sign.docx）
    """
    paras = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    total_text = ' '.join(paras)

    # 非空段落少于 3 个，且总文本长度小于 50，认为是图片型
    if len(paras) < 3 and len(total_text) < 50:
        return True

    # 检查前 5 个表格是否都几乎无文本
    empty_tables = 0
    for table in doc.tables[:5]:
        has_text = False
        for row in table.rows[:3]:
            for cell in row.cells:
                if get_cell_text(cell):
                    has_text = True
                    break
            if has_text:
                break
        if not has_text:
            empty_tables += 1

    if empty_tables >= 3 and len(total_text) < 100:
        return True

    return False


# ==============================================
# 函数6：图片型文档 OCR 识别
# ==============================================
def extract_images_from_docx(docx_path):
    """
    从 docx 文件中提取所有内嵌图片到临时目录
    返回图片文件路径列表
    """
    image_paths = []
    temp_dir = tempfile.mkdtemp(prefix="docx_ocr_")

    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            for item in zin.infolist():
                if item.filename.startswith('word/media/'):
                    try:
                        data = zin.read(item.filename)
                        ext = os.path.splitext(item.filename)[1].lower()
                        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                            img_name = f"img_{len(image_paths):04d}{ext}"
                            img_path = os.path.join(temp_dir, img_name)
                            with open(img_path, 'wb') as f:
                                f.write(data)
                            image_paths.append(img_path)
                    except Exception:
                        continue
    except Exception:
        pass

    return image_paths, temp_dir


def ocr_image(image_path):
    """
    对单张图片进行 OCR 识别
    优先使用 pytesseract；未安装则尝试 easyocr/paddleocr
    对低分辨率/小字号图片会自动放大，提升识别率
    """
    from PIL import Image, ImageEnhance

    try:
        img = Image.open(image_path)
    except Exception:
        return ""

    # 跳过极小图片（通常是图标、装饰线）
    if img.width < 80 or img.height < 20:
        return ""

    # 转换并增强：统一转灰度，低分辨率时放大
    img = img.convert('L')
    scale = 1.0
    if img.height < 80:
        scale = max(scale, 200.0 / img.height)
    if img.width < 300:
        scale = max(scale, 600.0 / img.width)
    if scale > 1.0:
        new_size = (int(img.width * scale), int(img.height * scale))
        img = img.resize(new_size, Image.LANCZOS)

    # 适度增强对比度
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.5)

    # 尝试 pytesseract
    try:
        import pytesseract
        # 兼容 Windows 上未将 Tesseract 加入 PATH 的情况
        _tess_path = r'C:\Users\szbc124\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
        if os.path.exists(_tess_path):
            pytesseract.pytesseract.tesseract_cmd = _tess_path

        # 优先只用英文包；若未安装中文包，eng+chi_sim 会报错
        for lang in ('eng', 'eng+chi_sim'):
            try:
                text = pytesseract.image_to_string(img, lang=lang)
                if text and text.strip():
                    return text
            except Exception:
                continue
    except ImportError:
        pass
    except Exception:
        pass

    # 尝试 easyocr
    try:
        import easyocr
        reader = easyocr.Reader(['en', 'ch_sim'], gpu=False)
        results = reader.readtext(image_path)
        text = '\n'.join([r[1] for r in results])
        return text
    except ImportError:
        pass
    except Exception:
        pass

    # 尝试 paddleocr
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        result = ocr.ocr(image_path, cls=True)
        texts = []
        if result and result[0]:
            for line in result[0]:
                if line:
                    texts.append(line[1][0])
        return '\n'.join(texts)
    except ImportError:
        pass
    except Exception:
        pass

    return ""


def extract_text_via_ocr(docx_path):
    """
    对图片型 docx 的所有内嵌图片进行 OCR，合并识别结果
    按文档顺序处理，主要保留可能是文字行的图片
    """
    from PIL import Image

    image_paths, temp_dir = extract_images_from_docx(docx_path)
    if not image_paths:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return ""

    # 图片已按文档出现顺序命名（img_0000, img_0001...），保持该顺序
    all_texts = []
    total_len = 0
    max_images = 400       # 文字行型 docx 可能拆成很多小图，放宽数量
    text_threshold = 8000  # 累计识别到约 8000 字符后停止，避免过长

    for idx, img_path in enumerate(image_paths[:max_images]):
        # 优先处理整行文字/段落/表格图片；跳过图标、小碎片、超大签名页
        size = os.path.getsize(img_path)
        if size < 512 or size > 2 * 1024 * 1024:
            continue
        try:
            with Image.open(img_path) as img:
                # 文字行/段落/表格宽度通常 ≥200；过低则多为图标
                if img.width < 200 or img.height < 15 or img.height > 2000:
                    continue
        except Exception:
            continue

        try:
            text = ocr_image(img_path)
            text = text.strip() if text else ""
            if text:
                all_texts.append(text)
                total_len += len(text)
                if total_len >= text_threshold:
                    break
        except Exception:
            continue

    shutil.rmtree(temp_dir, ignore_errors=True)
    return '\n'.join(all_texts)


def extract_fields_from_ocr_text(text):
    """
    从 OCR 识别出的英文文本中用正则提取关键字段
    返回 {英文字段名: 值} 的字典
    """
    fields = {}
    if not text:
        return fields

    # 标准化：合并多余空格、统一冒号
    text = re.sub(r'\s+', ' ', text)

    # 标签映射：英文标签 → 字段名
    ocr_patterns = [
        (r'Report\s*No\.?\s*[:#]\s*([A-Z]{1,5}\d{6,15}[A-Z0-9]*)', 'Report No.'),
        (r'Applicant\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Client|Company|Product|Sample|Date|Test)|$)', 'Applicant'),
        (r'Client\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Company|Product|Sample|Date|Test)|$)', 'Client'),
        (r'Company\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Product|Sample|Date|Test)|$)', 'Company'),
        (r'Address\s*[:：]\s*([^\n]+?)(?=\s+(?:Product|Sample|Model|Date|Test|Manufacturer)|$)', 'Address'),
        (r'Product\s*Name\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Sample|Date|Test)|$)', 'Product name'),
        (r'Sample\s*Name\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Date|Test)|$)', 'Sample name'),
        (r'Sample\s*Description\s*[:：]\s*([^\n]+?)(?=\s+(?:Model|Part|Date|Test)|$)', 'Sample Description'),
        (r'Model\s*[:：]\s*([^\n]+?)(?=\s+(?:Part|Sample|Date|Test)|$)', 'Model'),
        (r'Part\s*No\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Sample|Date|Test|Model)|$)', 'Part No.'),
        (r'Part\s*Number\s*[:：]\s*([^\n]+?)(?=\s+(?:Sample|Date|Test|Model)|$)', 'Part Number'),
        (r'Sample\s*No\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Received)|$)', 'Sample No.'),
        (r'Sample\s*Received\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Period)|$)', 'Sample Received Date'),
        (r'Date\s*of\s*Receipt\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Period)|$)', 'Date of Receipt'),
        (r'Testing\s*Period\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Report)|$)', 'Testing Period'),
        (r'Test\s*Period\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Report)|$)', 'Test Period'),
        (r'Report\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Issue)|$)', 'Report Date'),
        (r'Issue\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Report)|$)', 'Issue Date'),
        (r'Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Test|Report|No)|$)', 'Date'),
        (r'Manufacturer\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Factory|Date|Test)|$)', 'Manufacturer'),
        (r'Factory\s*[:：]\s*([^\n]+?)(?=\s+(?:Address|Date|Test)|$)', 'Factory'),
        (r'Material\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test|Production)|$)', 'Material'),
        (r'Production\s*Date\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Production Date'),
        (r'Trade\s*Mark\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Trade Mark'),
        (r'Client\s*Ref\.?\s*Info\.?\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Test)|$)', 'Client Ref. Info.'),
        (r'Compiled\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Reviewed|Checked|Approved)|$)', 'Compiled by'),
        (r'Prepared\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Reviewed|Checked|Approved)|$)', 'Prepared by'),
        (r'Reviewed\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Checked|Approved|Date)|$)', 'Reviewed by'),
        (r'Checked\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Approved|Date)|$)', 'Checked by'),
        (r'Approved\s*by\s*[:：]\s*([^\n]+?)(?=\s+(?:Date|Report)|$)', 'Approved by'),
    ]

    for pattern, field_name in ocr_patterns:
        if field_name in fields and fields[field_name]:
            continue
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            # 清理末尾可能截断的短语
            value = re.sub(r'\s+(?:Test|Date|Report|No|Sample)\s*$', '', value, flags=re.IGNORECASE)
            if value:
                fields[field_name] = value

    # 检测要求和结论：尝试从 Test Requirement / Test Conclusion 段落提取
    req_match = re.search(
        r'Test\s*Requirement\s*[:：]?\s*([^\n]*(?:\n(?![A-Z][a-zA-Z\s]{2,20}[:：]).*)*)',
        text, re.IGNORECASE
    )
    if req_match:
        fields['Test Requirement'] = req_match.group(1).strip()

    con_match = re.search(
        r'Test\s*Conclusion\s*[:：]?\s*([^\n]*(?:\n(?![A-Z][a-zA-Z\s]{2,20}[:：]).*)*)',
        text, re.IGNORECASE
    )
    if con_match:
        fields['Test Conclusion'] = con_match.group(1).strip()

    return fields


# ==============================================
# 函数7：完整提取英文报告
# ==============================================
def extract_english_report(doc, filename, file_path=None):
    """英文报告完整提取，返回字段字典和检测方法列表"""

    # 图片型文档：使用 OCR 识别
    if is_image_based_document(doc):
        if VERBOSE:
            print(f"  ⚠️  检测到图片型文档，尝试 OCR 识别...")

        result = {}
        ocr_text = ""

        if file_path:
            try:
                ocr_text = extract_text_via_ocr(file_path)
                if ocr_text:
                    ocr_fields = extract_fields_from_ocr_text(ocr_text)
                    for raw_key, raw_val in ocr_fields.items():
                        std_key = EN_FIELD_MAPPING.get(raw_key)
                        if std_key and raw_val:
                            result[std_key] = raw_val
                    # 保留 OCR 原文作为备注，便于人工核对
                    if '备注' not in result or not result['备注']:
                        result['备注'] = f"OCR识别内容：{ocr_text[:500]}"
                else:
                    result['备注'] = "图片型文档，OCR 未识别到有效文本（请检查是否已安装 Tesseract/pytesseract）"
            except Exception as e:
                result['备注'] = f"图片型文档，OCR 识别失败: {str(e)[:120]}"
        else:
            result['备注'] = "图片型文档，缺少文件路径无法 OCR"

        # 若 OCR 未识别出报告编号，再尝试从文件名补
        if not result.get("报告编号"):
            m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z0-9]*)', filename, re.IGNORECASE)
            if m:
                result["报告编号"] = m.group(1)

        return result, []

    # 1. 基础字段
    raw_fields = extract_basic_info(doc)

    result = {}
    for raw_key, raw_val in raw_fields.items():
        std_key = EN_FIELD_MAPPING.get(raw_key)
        if std_key and std_key in result and raw_val:
            # 已存在则不覆盖（取首次出现的值）
            pass
        elif std_key and raw_val:
            result[std_key] = raw_val

    # 单独处理制造商地址
    if '制造商地址' in raw_fields and raw_fields['制造商地址']:
        result['制造商地址'] = raw_fields['制造商地址']

    # 2. 检测要求/结论
    test_req, test_con = extract_test_requirement_and_conclusion(doc)
    if test_req:
        result["检测要求"] = test_req
    if test_con:
        result["检测结论"] = test_con

    # 3. 报告编号从文件名补
    m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z]?\d*)', filename, re.IGNORECASE)
    if m and not result.get("报告编号"):
        result["报告编号"] = m.group(1)

    # 4. 样品部件总数
    sample_count = extract_sample_count(doc)
    if sample_count == 0:
        sample_count = extract_sample_count_from_test_items(doc)
    result["样品部件总数"] = str(sample_count)

    # 5. 检测方法
    test_methods = extract_test_methods(doc)
    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]

    return result, test_methods


# ==============================================
# 主函数：处理单个文档
# ==============================================
def process_word_file(input_path, output_excel):
    """
    处理单个英文 Word 检测报告，提取信息并生成 Excel

    参数:
        input_path: Word 文档路径
        output_excel: 输出 Excel 路径
    """
    from docx import Document

    print("=" * 60)
    print(f"开始处理: {os.path.basename(input_path)}")
    print("=" * 60)

    # 1. 检查文件是否存在
    if not os.path.exists(normalize_path(input_path)):
        print(f"❌ 错误：文件不存在 - {input_path}")
        return False

    # 2. 修复文档（如果需要）
    try:
        actual_file = fix_docx_if_needed(input_path)
    except Exception as e:
        print(f"❌ 错误：无法打开文档 - {e}")
        print("   请检查文件是否为有效的 Word 文档(.docx格式)")
        return False

    # 3. 打开文档
    try:
        doc = Document(actual_file)
    except ImportError:
        print("❌ 错误：未安装 python-docx 库")
        print("   请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 错误：打开文档失败 - {e}")
        return False

    if VERBOSE:
        print(f"\n📄 文档信息:")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")

    # 4. 初始化结果字典
    result = {
        "文件名": os.path.basename(input_path),
        "报告编号": "",
        "申请商": "",
        "申请商地址": "",
        "制造商地址": "",
        "样品名称": "",
        "零件号": "",
        "商标": "",
        "客户参考信息": "",
        "材质": "",
        "生产日期": "",
        "制造商": "",
        "样品编号": "",
        "样品接收日期": "",
        "检测期间": "",
        "检测要求": "",
        "检测结论": "",
        "报告日期": "",
        "样品部件总数": "",
        "编制": "",
        "审核": "",
        "批准": "",
        "备注": "",
    }

    # 5. 提取
    extracted, test_methods = extract_english_report(doc, os.path.basename(input_path), input_path)
    result.update(extracted)

    # 6. 生成 Excel
    print(f"\n💾 生成 Excel 文件...")

    col_order = [
        "报告编号", "申请商", "申请商地址", "制造商地址",
        "样品名称", "零件号", "商标",
        "客户参考信息", "材质", "生产日期", "制造商", "样品编号",
        "样品接收日期", "检测期间", "检测要求", "检测结论",
        "报告日期", "样品部件总数",
        "编制", "审核", "批准"
    ]

    actual_method_count = min(len(test_methods), MAX_TEST_METHODS)
    for i in range(1, actual_method_count + 1):
        col_order.append(f"检测项目{i}")
        col_order.append(f"检测方法{i}")
        col_order.append(f"检测仪器{i}")

    col_order.extend(["备注", "文件名"])

    df = pd.DataFrame([result])
    for col in col_order:
        if col not in df.columns:
            df[col] = ""
    df = df[col_order]

    try:
        df.to_excel(output_excel, index=False)
        print(f"✅ 成功保存: {output_excel}")
    except ImportError:
        print("❌ 错误：未安装 openpyxl 库")
        print("   请运行: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 错误：保存 Excel 失败 - {e}")
        return False

    # 统计信息
    non_empty = sum(1 for val in result.values() if val)
    print(f"\n📊 提取完成！共提取 {non_empty} 个非空字段")
    print("=" * 60)

    return True


# ==============================================
# 程序入口
# ==============================================
def main():
    print("\n" + "=" * 60)
    print("  英文检测报告信息提取工具 4")
    print("=" * 60)

    # 检查依赖
    try:
        import docx
        import pandas
    except ImportError as e:
        print(f"\n❌ 缺少依赖库: {e}")
        print("请运行以下命令安装依赖:")
        print("  pip install python-docx pandas openpyxl")
        sys.exit(1)

    # 处理文件
    success = process_word_file(INPUT_FILE, OUTPUT_EXCEL)

    if success:
        print("\n🎉 处理完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)


if __name__ == "__main__":
    main()
