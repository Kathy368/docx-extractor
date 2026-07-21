# -*- coding: utf-8 -*-
"""
检测报告批量提取工具 v6.0（总提取代码6）
功能：从嵌套文件夹中批量提取中英文 Word 检测报告信息，输出三份 Excel
- 中文报告提取结果_<时间戳>.xlsx
- 英文报告提取结果_<时间戳>.xlsx
- 中英文汇总提取结果_<时间戳>.xlsx

改进点：
1. 中文模块升级为 extract_report_test2_docx3中文成功plus4.py，基础信息入口改为 extract_basic_info
2. 英文模块升级为 extract_report_test2_docx2英文成功4.py，复用其 extract_english_report 以支持特殊格式、无序号样品统计与图片 OCR 兜底
3. 文件名过滤增强：新增证书/更新/Supplement Sample Notice/申请单/通知书/通知单等排除规则
4. 保留报告编号正则判定、白名单关键词、Word 临时文件过滤逻辑
"""

import os
import re
import sys
import time
import zipfile
import tempfile
import html
import pandas as pd
from tqdm import tqdm

# 兼容 Windows GBK 终端：强制 stdout/stderr 使用 utf-8，避免打印表情/特殊符号时编码错误
try:
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
except Exception:
    pass

# ==============================================
# ========== 【配置填写区 - 改这里就行】 ==========
# ==============================================

# ====== 1. 路径配置 ======
# 报告根目录路径
ROOT_FOLDER = r"D:\Kathy\PDF提取工具\测试报告2"

# 输出文件夹
OUTPUT_FOLDER = r"D:\Kathy\PDF提取工具\测试报告2"


# ====== 2. 【测试开关 - 非常重要】 ======
# 测试模式1：只扫描前N个一级公司文件夹（0=全部扫描）
TEST_FOLDER_LIMIT = 0

# 测试模式2：只提取前N份报告（0=全部提取）
TEST_REPORT_LIMIT = 0


# ====== 3. 其他配置 ======
# 最大检测方法/仪器数量
MAX_TEST_METHODS = 100

# 调试模式：打印每份报告的提取详情（False=只显示进度条）
DEBUG_MODE = False


# ==============================================
# 导入中文/英文提取模块
# ==============================================
import extract_report_test2_docx3中文成功plus4 as _zh_module
import extract_report_test2_docx2英文成功4 as _en_module

# 同步导入模块的 VERBOSE 开关，避免批量时产生过多日志
_zh_module.VERBOSE = DEBUG_MODE
_en_module.VERBOSE = DEBUG_MODE

from extract_report_test2_docx3中文成功plus4 import (
    fix_docx_if_needed,
    get_cell_text,
    normalize_text,
    normalize_key,
)


def get_row_cells_text_fast(row):
    """
    统一使用 XML 读取一行单元格的完整文本。
    避免 python-docx 的 cell.text 在 SDT（内容控件）下只读到 '~' 或空字符串。
    """
    return [get_cell_text(c) for c in row.cells]


# 兼容旧别名
get_row_cells_text = get_row_cells_text_fast


# ==============================================
# 一、报告文件名判定规则
# ==============================================

# 报告编号正则：1-5 个字母 + 6-15 位数字 + 可选 1 个字母 + 可选数字
REPORT_NO_PATTERN = re.compile(
    r'(?:^|[^A-Za-z0-9])([A-Z]{1,5}\d{6,15}[A-Z]?\d*)',
    re.IGNORECASE
)

# 非报告关键词黑名单：文件名中出现这些，直接排除
NON_REPORT_KEYWORDS = [
    # 证书/更新/补充通知类
    '证书', '更新', 'Supplement Sample Notice',
    # 数据反馈/结果通知类
    '数据反馈', '结果通知', '反馈',
    # 申请表/委托单类
    '申请表', '申请单', '委托单', '测试申请', '检测申请', '申请',
    # 通知书/通知单类
    '通知书', '通知单',
    # 签名页/补样类
    'EN_sign', '_sign', '补样', '补样清单',
    # 参考/模板/工作指导书类
    '参考', '模板', '范本', 'WI-', 'Work Instruction',
    # 流程/签名状态类
    '待放签名', '已审核', '待审核', '待签名',
    # 其他非报告
    '照片', '图片', '截图', '邮件', '合同', '报价单', '发票',
]

# 报告关键词白名单：文件名中出现这些，更倾向于判定为报告
REPORT_KEYWORDS = [
    '报告', '检测报告', '测试报告', '试验报告',
    'Report', 'TEST REPORT', 'Test Report', 'Testing Report',
    'Certificate', 'Cert',
    'RoHS', 'ROHS', 'SVHC', 'REACH', 'ELV', '卤素', 'HF',
]


def is_report_by_filename(filename):
    """
    根据文件名判断是否是报告文件
    规则：
    1. 跳过 Word 临时文件
    2. 命中黑名单关键词 → 不是报告
    3. 包含报告编号正则，或命中白名单关键词 → 是报告
    """
    # 跳过 Word 临时文件
    if filename.startswith('~$'):
        return False

    name_no_ext = os.path.splitext(filename)[0]
    name_lower = name_no_ext.lower()

    # 黑名单优先
    for kw in NON_REPORT_KEYWORDS:
        if kw.lower() in name_lower:
            return False

    # 白名单或报告编号
    if REPORT_NO_PATTERN.search(name_no_ext):
        return True

    for kw in REPORT_KEYWORDS:
        if kw.lower() in name_lower:
            return True

    return False


def extract_conclusion_from_filename(filename):
    """
    从文件名中提取检测结论关键词。
    若文件名包含 "符合" / "不符合" / "合格" / "不合格"，直接返回对应关键词。
    优先匹配更长的否定词，避免 "符合" 误命中 "不符合"。
    """
    name_no_ext = os.path.splitext(filename)[0]
    # 按长度降序排列：优先 "不合格" > "不符合" > "合格" > "符合"
    conclusion_keywords = ["不合格", "不符合", "合格", "符合"]
    for kw in conclusion_keywords:
        if kw in name_no_ext:
            return kw
    return ""


# 扫描时跳过的文件夹黑名单
SKIP_FOLDERS = {'.venv', 'venv', '__pycache__', '.git', '.idea', 'node_modules', '报告图片', '检测结果图片'}


# ==============================================
# 二、穷尽扫描（保留实时计数进度）
# ==============================================

def scan_all_folders(root_folder):
    """递归扫描所有文件夹，识别 docx 报告文件"""
    report_files = []
    no_report_projects = []

    try:
        items = os.listdir(root_folder)
    except Exception as e:
        print(f"[错误] 读取根目录失败：{e}")
        return [], []

    # 收集一级项目文件夹
    project_folders = sorted([
        item for item in items
        if os.path.isdir(os.path.join(root_folder, item)) and item not in SKIP_FOLDERS
    ])

    # 同时收集根目录下的 docx 文件（作为一个虚拟项目）
    root_docx = [
        os.path.join(root_folder, f)
        for f in items
        if f.lower().endswith('.docx') and not f.startswith('~$')
    ]

    # 测试模式：只取前N个一级文件夹
    if TEST_FOLDER_LIMIT > 0 and len(project_folders) > TEST_FOLDER_LIMIT:
        project_folders = project_folders[:TEST_FOLDER_LIMIT]
        print(f"\n[警告]  【文件夹测试模式】只扫描前 {TEST_FOLDER_LIMIT} 个一级公司文件夹")

    total_dirs = 0
    total_files = 0
    total_docx = 0
    start_time = time.time()

    print(f"\n[文件夹] 开始扫描，共 {len(project_folders)} 个一级项目文件夹...")

    # 先处理根目录下的文件
    if root_docx:
        root_reports = [f for f in root_docx if is_report_by_filename(os.path.basename(f))]
        total_files += len(root_docx)
        total_docx += len(root_docx)
        if root_reports:
            report_files.extend(root_reports)
            print(f"   [根目录] 发现 {len(root_docx)} 个 DOCX，其中 {len(root_reports)} 份报告")
        else:
            no_report_projects.append({
                "项目/文件夹名称": "【根目录】",
                "完整路径": root_folder,
                "文件总数": len(root_docx),
                "DOCX数量": len(root_docx),
                "DOCX文件名示例": "\n".join([os.path.basename(fp) for fp in root_docx[:20]]),
                "情况说明": "根目录下未匹配报告规则"
            })

    for project_idx, project in enumerate(project_folders, 1):
        project_path = os.path.join(root_folder, project)
        if not os.path.isdir(project_path):
            continue

        project_reports = []
        all_docx = []
        project_files = 0
        project_dirs = 0

        for dirpath, dirnames, filenames in os.walk(project_path):
            # 跳过黑名单子文件夹
            dirnames[:] = [d for d in dirnames if d not in SKIP_FOLDERS]

            project_dirs += 1
            project_files += len(filenames)

            for f in filenames:
                if f.startswith('~$'):
                    continue
                f_lower = f.lower()
                f_path = os.path.join(dirpath, f)

                # 只处理 docx
                if f_lower.endswith('.docx'):
                    all_docx.append(f_path)
                    if is_report_by_filename(f):
                        project_reports.append(f_path)

        total_dirs += project_dirs
        total_files += len(set(all_docx))
        total_docx += len(all_docx)

        elapsed = time.time() - start_time
        print(f"   [{project_idx}/{len(project_folders)}] {project[:35]:35} "
              f"| 文件夹:{total_dirs} | 文件:{total_files} | DOCX:{total_docx} | 报告:{len(report_files) + len(project_reports)} | {elapsed:.1f}s")

        if project_reports:
            report_files.extend(project_reports)
        elif all_docx:
            no_report_projects.append({
                "项目/文件夹名称": project,
                "完整路径": project_path,
                "文件总数": project_files,
                "DOCX数量": len(all_docx),
                "DOCX文件名示例": "\n".join([os.path.basename(fp) for fp in all_docx[:20]]),
                "情况说明": "未匹配报告编号格式，需人工确认"
            })
        else:
            no_report_projects.append({
                "项目/文件夹名称": project,
                "完整路径": project_path,
                "文件总数": project_files,
                "DOCX数量": 0,
                "DOCX文件名示例": "无",
                "情况说明": "无DOCX文件"
            })

    total_time = time.time() - start_time
    print(f"\n[OK] 扫描完成！用时 {total_time:.1f} 秒")
    print(f"   总文件夹：{total_dirs} | 总文件：{total_files} | DOCX：{total_docx}")
    print(f"   识别报告：{len(report_files)} 份 | 待确认项目：{len(no_report_projects)} 个")

    return report_files, no_report_projects


# ==============================================
# 三、通用工具函数
# ==============================================

def detect_language(doc):
    """
    判断文档是中文还是英文报告
    优先依据前 10 个非空段落：
    - 首段或前 3 段出现 "检测报告" → 中文
    - 首段或前 3 段出现 "TEST REPORT" → 英文
    - 兜底：前 10 段中 CJK 字符占比 > 20% → 中文，否则英文
    """
    paras = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    first_ten = paras[:10]
    first_ten_text = '\n'.join(first_ten)

    # 强规则：关键词匹配
    for p in first_ten[:3]:
        if '检测报告' in p:
            return "zh"
        if p.upper() == 'TEST REPORT' or 'TEST REPORT' in p.upper():
            return "en"

    # 兜底：CJK 字符占比
    total_chars = len(first_ten_text)
    chinese_chars = len(re.findall(r'[一-鿿]', first_ten_text))

    if total_chars == 0:
        return "zh"

    ratio = chinese_chars / total_chars
    return "zh" if ratio > 0.2 else "en"


# ==============================================
# 四、中文报告提取模块
# ==============================================

from extract_report_test2_docx3中文成功plus4 import (
    detect_doc_type as detect_doc_type_zh,
    extract_basic_info as extract_basic_info_zh,
    extract_test_requirement_and_conclusion as extract_test_req_conclusion_zh,
    extract_conclusion_from_result_column_zh,
    extract_sample_count as extract_sample_count_zh,
    extract_test_methods_from_tables,
    extract_test_methods_from_paragraphs,
    extract_test_methods_long,
    extract_test_items_from_result_tables,
    find_column_index,
    ZH_FIELD_MAPPING,
)


def extract_zh_test_methods_direct(doc):
    """
    中文检测三字段直接对应：仅读取表头同时包含"检测项目"与"检测方法/检测仪器"的表格，
    按行直接映射为 检测项目、检测方法、检测仪器。
    同一检测项目出现多次时，将多个方法/仪器用" / "合并到一条；缺失字段保持空值。
    """
    methods_map = {}

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [normalize_text(c) for c in get_row_cells_text_fast(table.rows[0])]
        first_row_text = ' '.join(first_row)

        if not ('检测项目' in first_row_text and ('检测方法' in first_row_text or '检测仪器' in first_row_text)):
            continue

        item_col = find_column_index(first_row, ['检测项目', '测试项目', '项目'])
        method_col = find_column_index(first_row, ['检测方法', '测试方法', '方法'])
        instrument_col = find_column_index(first_row, ['检测仪器', '测试仪器', '仪器'])

        for row in table.rows[1:]:
            cells = [normalize_text(c) for c in get_row_cells_text_fast(row)]
            if len(cells) < 2:
                continue

            item = cells[item_col] if item_col is not None and item_col < len(cells) else ""
            method = cells[method_col] if method_col is not None and method_col < len(cells) else ""
            instrument = cells[instrument_col] if instrument_col is not None and instrument_col < len(cells) else ""

            if not item or item == '<br />' or item.isdigit():
                continue
            if any(k in item for k in ['Limit', '△', 'Note', '注', '限值', '方法检出限', 'MDL', '序号']):
                continue

            if item not in methods_map:
                methods_map[item] = {
                    "检测项目": item,
                    "methods": [],
                    "instruments": [],
                }
            if method and method not in methods_map[item]["methods"]:
                methods_map[item]["methods"].append(method)
            if instrument and instrument not in methods_map[item]["instruments"]:
                methods_map[item]["instruments"].append(instrument)

    methods = []
    for item, data in methods_map.items():
        methods.append({
            "检测项目": item,
            "检测方法": " / ".join(data["methods"]),
            "检测仪器": " / ".join(data["instruments"]),
        })

    return methods


def extract_chemical_test_methods_en(doc):
    """
    从 Chemical Test 部分表格提取检测项目/检测方法/检测仪器。
    表格特征：表头同时包含 Test item / Testing item、Test method、Test instrument。
    兼容 Word 中分 run 书写导致 "Testinstrument" 这类无空格表头。
    """
    methods = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c).strip() for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row).lower()

        has_item = bool(re.search(r'test\s*item|testing\s*item', first_row_text))
        has_method = bool(re.search(r'test\s*method', first_row_text))
        has_instrument = bool(re.search(r'test\s*instrument', first_row_text))

        if has_item and has_method and has_instrument:
            item_col = method_col = instrument_col = None
            for idx, cell in enumerate(first_row):
                cell_lower = cell.lower()
                if re.search(r'test\s*item|testing\s*item', cell_lower) and item_col is None:
                    item_col = idx
                elif re.search(r'test\s*method', cell_lower) and method_col is None:
                    method_col = idx
                elif re.search(r'test\s*instrument', cell_lower) and instrument_col is None:
                    instrument_col = idx

            for row in table.rows[1:]:
                cells = [get_cell_text(c).strip() for c in row.cells]
                if item_col is None or item_col >= len(cells):
                    continue

                item = cells[item_col]
                method = cells[method_col] if method_col is not None and method_col < len(cells) else ""
                instrument = cells[instrument_col] if instrument_col is not None and instrument_col < len(cells) else ""

                if not item or re.search(r'test\s*item', item.lower()):
                    continue
                if 'limit' in item.lower() and 'mdl' in item.lower():
                    continue

                key = f"{item}|{method}|{instrument}"
                if key in seen:
                    continue
                seen.add(key)

                methods.append({
                    "检测项目": item,
                    "检测方法": method,
                    "检测仪器": instrument,
                })

    return methods


def extract_conclusion_from_test_results_en(doc):
    """
    兜底：从 Test Result(s) 表格的 Conclusion 列/行提取结论（如 Pass）。
    不删除、不改变原有 Test Conclusion 提取逻辑，仅作为补充。
    """
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c).strip() for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row).lower()

        # 识别 Test Result(s) 表格
        if 'part no.' in first_row_text and 'part description' in first_row_text:
            conclusion_col = None
            for idx, cell in enumerate(first_row):
                if 'conclusion' in cell.lower():
                    conclusion_col = idx
                    break

            if conclusion_col is not None:
                for row in table.rows[1:]:
                    cells = [get_cell_text(c).strip() for c in row.cells]
                    if conclusion_col < len(cells):
                        val = cells[conclusion_col]
                        if val and 'pass' in val.lower():
                            return val

        # 某些报告有独立的 Conclusion 行（如 Test Items / Test Result / Conclusion）
        if 'conclusion' in first_row_text:
            for row in table.rows[1:]:
                cells = [get_cell_text(c).strip() for c in row.cells]
                for cell in cells:
                    if 'pass' in cell.lower():
                        return cell

    return ""


def extract_chinese_report(doc, filename):
    """中文报告完整提取"""
    doc_type = detect_doc_type_zh(doc)

    if DEBUG_MODE:
        print(f"  [文档] 中文文档类型: {'长文档' if doc_type == 'long' else '短文档'}")

    # 1. 基础字段（使用新版模块的 extract_basic_info 入口）
    raw_fields = extract_basic_info_zh(doc)

    result = {}
    for raw_key, raw_val in raw_fields.items():
        std_key = ZH_FIELD_MAPPING.get(raw_key)
        if std_key:
            result[std_key] = raw_val

    # 2. 检测要求/结论（文件名含结论词时优先采用）
    filename_conclusion = extract_conclusion_from_filename(filename)
    if filename_conclusion:
        result["检测结论"] = filename_conclusion
    else:
        test_req, test_con = extract_test_req_conclusion_zh(doc)
        if test_req:
            result["检测要求"] = test_req
        if test_con:
            result["检测结论"] = test_con

        # 2.1 兜底：仍未取到检测结论时，从结果表的“结论”列提取
        if not result.get("检测结论", "").strip():
            fallback_con = extract_conclusion_from_result_column_zh(doc)
            if fallback_con:
                result["检测结论"] = fallback_con

    # 3. 报告编号从文件名补
    m = REPORT_NO_PATTERN.search(filename)
    if m:
        result["报告编号"] = m.group(1)

    # 4. 样品部件总数：按"检测结果 / Test Result(s)"模块优先，无部件信息时退到独立"样品描述"模块取最大序号
    sample_count = extract_sample_count_zh(doc)
    result["样品部件总数"] = str(sample_count)

    # 5. 检测方法/仪器：任务 B 要求直接对应"检测项目/检测方法/检测仪器"三列，缺失保持空值
    direct_methods = extract_zh_test_methods_direct(doc)
    if direct_methods:
        test_methods = direct_methods
    elif doc_type == "long":
        # 长文档若无三列表格，保留首页方法总结作为兜底
        test_methods = extract_test_methods_long(doc)
    else:
        test_methods = []

    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]

    result["_doc_type"] = doc_type
    return result


def extract_instruments_only_en(doc):
    """
    兜底：从仅含 Test instrument / Instrument model 的表格中提取检测仪器。
    适用于没有 Test item / Test method 列但单独列出仪器的报告（如 PFBS 报告）。
    """
    instruments = []
    seen = set()

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_row = [get_cell_text(c).strip() for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row).lower()

        # 只处理含 instrument 信息但不含完整方法表的表格
        if not ('test instrument' in first_row_text or 'instrument model' in first_row_text):
            continue
        if 'test item' in first_row_text and 'test method' in first_row_text:
            continue

        name_col = model_col = None
        for idx, cell in enumerate(first_row):
            cl = cell.lower()
            if 'test instrument' in cl and name_col is None:
                name_col = idx
            elif 'instrument model' in cl and model_col is None:
                model_col = idx
            elif 'model' in cl and model_col is None:
                model_col = idx

        for row in table.rows[1:]:
            cells = [get_cell_text(c).strip() for c in row.cells]
            name = cells[name_col] if name_col is not None and name_col < len(cells) else ""
            model = cells[model_col] if model_col is not None and model_col < len(cells) else ""

            if not name or name.lower() in ['test instrument', 'instrument']:
                continue

            full = name
            if model:
                full += f" ({model})"

            if full in seen:
                continue
            seen.add(full)
            instruments.append({
                "检测项目": "",
                "检测方法": "",
                "检测仪器": full,
            })

    return instruments


# ==============================================
# 五、英文报告提取模块
# ==============================================

from extract_report_test2_docx2英文成功4 import (
    extract_english_report as extract_english_report_module,
)


def extract_english_report(doc, filename, file_path=None):
    """
    英文报告完整提取
    直接复用英文成功4模块的 extract_english_report，以获得：
    - 标准格式表格提取
    - 无序号样品统计兜底
    - 图片型报告的 OCR 识别兜底
    返回前补充/修正以下字段：
    - 样品部件总数：直接使用模块结果（已包含 Test Result(s) / Sample Description 多层兜底）
    - 检测项目/方法/仪器：优先使用按表头映射的 Chemical Test 表格覆盖原结果；无方法时尝试仅仪器表
    - 检测结论：若模块未提取到，尝试从 Test Result(s) 的 Conclusion 列/行读取；仍无则最终写 Pass
    """
    result, _ = extract_english_report_module(doc, filename, file_path)

    # 1. 样品部件总数：直接使用英文成功4模块的结果（已包含 Test Result(s) / Sample Description 多层兜底）
    if not result.get("样品部件总数"):
        result["样品部件总数"] = "0"

    # 2. 检测项目/方法/仪器：优先使用 Chemical Test 表格（按表头映射列，修复 BOSER 这类 CAS No. 错位）
    chemical_methods = extract_chemical_test_methods_en(doc)
    if chemical_methods:
        # 清除模块可能已写入的错误/旧方法字段
        for i in range(1, MAX_TEST_METHODS + 1):
            result.pop(f"检测项目{i}", None)
            result.pop(f"检测方法{i}", None)
            result.pop(f"检测仪器{i}", None)
        for i, method in enumerate(chemical_methods[:MAX_TEST_METHODS], 1):
            result[f"检测项目{i}"] = method["检测项目"]
            result[f"检测方法{i}"] = method["检测方法"]
            result[f"检测仪器{i}"] = method["检测仪器"]
    else:
        # 若完全没有仪器，再尝试仅仪器表（如 S24011507201E）
        has_instrument = any(
            result.get(f"检测仪器{i}", "").strip()
            for i in range(1, MAX_TEST_METHODS + 1)
        )
        if not has_instrument:
            instrument_only = extract_instruments_only_en(doc)
            for i, method in enumerate(instrument_only[:MAX_TEST_METHODS], 1):
                result[f"检测项目{i}"] = method["检测项目"]
                result[f"检测方法{i}"] = method["检测方法"]
                result[f"检测仪器{i}"] = method["检测仪器"]

    # 3. 检测结论：文件名含结论词时优先采用
    filename_conclusion = extract_conclusion_from_filename(filename)
    if filename_conclusion:
        result["检测结论"] = filename_conclusion
    else:
        # 模块未提取到时，从 Test Result(s) Conclusion 兜底；仍无则最终写 Pass
        if not result.get("检测结论", "").strip():
            conclusion = extract_conclusion_from_test_results_en(doc)
            if conclusion:
                result["检测结论"] = conclusion
        if not result.get("检测结论", "").strip():
            result["检测结论"] = "Pass"

    return result


# ==============================================
# 六、单份报告统一处理入口
# ==============================================

def process_single_report(file_path, idx, total):
    """处理单份报告，返回结果字典和语言类型"""
    filename = os.path.basename(file_path)

    result = {
        "文件名": filename,
        "文件路径": file_path,
        "提取状态": "成功",
    }

    try:
        # 1. 修复并打开文档
        actual_file = fix_docx_if_needed(file_path)
        from docx import Document
        doc = Document(actual_file)

        # 2. 判断语言
        lang = detect_language(doc)

        # 3. 分别提取
        if lang == "zh":
            extracted = extract_chinese_report(doc, filename)
            result.update(extracted)
        else:
            extracted = extract_english_report(doc, filename, file_path)
            result.update(extracted)

        return result, lang

    except Exception as e:
        result["提取状态"] = f"失败: {str(e)[:80]}"
        return result, "unknown"


# ==============================================
# 七、表头构建
# ==============================================

def build_header_list():
    """构建三个 Excel 的表头列表"""

    # 通用基础字段
    base_common = [
        "报告编号", "申请商", "申请商地址", "样品名称", "零件号",
        "样品编号", "材质", "样品接收日期", "检测期间", "检测要求",
        "检测结论", "报告日期", "样品部件总数",
    ]

    # 中文特有字段
    zh_only = [
        "参考零件号", "车型", "生产日期", "供应商代码", "测试类型", "主机厂",
    ]

    # 英文特有字段
    en_only = [
        "制造商", "制造商地址", "客户参考信息", "商标", "编制", "审核", "批准",
    ]

    # 动态检测方法字段
    method_fields = []
    for i in range(1, MAX_TEST_METHODS + 1):
        method_fields.extend([
            f"检测项目{i}", f"检测方法{i}", f"检测仪器{i}"
        ])

    # 辅助字段
    aux_fields = ["文件名", "文件路径", "提取状态"]

    # 中文报告表头
    zh_headers = base_common + zh_only + method_fields + aux_fields

    # 英文报告表头
    en_headers = base_common + en_only + method_fields + aux_fields

    # 汇总版表头（并集）
    merged_base = [
        "报告编号", "申请商", "申请商地址", "制造商", "制造商地址",
        "样品名称", "零件号", "参考零件号", "客户参考信息",
        "材质", "车型", "生产日期", "商标",
        "供应商代码", "测试类型", "主机厂",
        "样品编号", "样品接收日期", "检测期间",
        "检测要求", "检测结论", "报告日期",
        "样品部件总数",
        "编制", "审核", "批准",
    ]
    merged_headers = merged_base + method_fields + aux_fields

    return zh_headers, en_headers, merged_headers


# ==============================================
# 八、主程序
# ==============================================

def main():
    print("=" * 70)
    print("[搜索] 检测报告批量提取工具 v6.0（总提取代码6）")
    print(f"根目录：{ROOT_FOLDER}")
    print(f"输出到：{OUTPUT_FOLDER}")
    print(f"检测方法列数：{MAX_TEST_METHODS} 组")

    # 测试模式提示
    if TEST_FOLDER_LIMIT > 0:
        print(f"[警告]  文件夹测试模式：只扫描前 {TEST_FOLDER_LIMIT} 个一级文件夹")
    if TEST_REPORT_LIMIT > 0:
        print(f"[警告]  报告测试模式：只提取前 {TEST_REPORT_LIMIT} 份报告")
    if TEST_FOLDER_LIMIT == 0 and TEST_REPORT_LIMIT == 0:
        print(f"[OK] 正式模式：全部提取")

    print("=" * 70)

    if not os.path.exists(ROOT_FOLDER):
        print(f"[错误] 错误：根目录不存在 - {ROOT_FOLDER}")
        return

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    # 阶段1：扫描
    print("\n" + "=" * 70)
    print("【阶段 1/4】扫描文件夹")
    print("=" * 70)

    report_files, no_report_projects = scan_all_folders(ROOT_FOLDER)

    if not report_files:
        print("\n[错误] 没有识别到任何 DOCX 报告文件！")
        return

    # 测试模式：只取前N份报告
    if TEST_REPORT_LIMIT > 0 and len(report_files) > TEST_REPORT_LIMIT:
        report_files = report_files[:TEST_REPORT_LIMIT]
        print(f"\n[警告]  【报告测试模式】只提取前 {TEST_REPORT_LIMIT} 份报告")

    # 阶段2-3：提取
    print("\n" + "=" * 70)
    print(f"【阶段 2-3/4】提取信息（共 {len(report_files)} 份）")
    print("=" * 70)

    zh_results = []
    en_results = []
    error_count = 0
    start_phase2 = time.time()

    for i, file_path in enumerate(tqdm(report_files, desc="提取进度", unit="份"), 1):
        if DEBUG_MODE:
            print(f"\n[{i}/{len(report_files)}] {os.path.basename(file_path)}")

        result, lang = process_single_report(file_path, i, len(report_files))
        result["序号"] = i

        if lang == "zh":
            zh_results.append(result)
        elif lang == "en":
            en_results.append(result)
        else:
            zh_results.append(result)

        if "失败" in str(result.get("提取状态", "")):
            error_count += 1
            if DEBUG_MODE:
                print(f"  [错误] 失败: {result['提取状态']}")

    phase2_time = time.time() - start_phase2

    # 阶段4：输出三份Excel
    print("\n" + "=" * 70)
    print("【阶段 4/4】保存三份 Excel")
    print("=" * 70)

    timestamp = time.strftime("%Y%m%d_%H%M%S")
    zh_headers, en_headers, merged_headers = build_header_list()

    # 4.1 中文报告Excel
    zh_excel = os.path.join(OUTPUT_FOLDER, f"中文报告提取结果_{timestamp}.xlsx")
    if zh_results:
        df_zh = pd.DataFrame(zh_results)
        for col in zh_headers:
            if col not in df_zh.columns:
                df_zh[col] = ""
        df_zh = df_zh[zh_headers]

        # 删除尾部全空的方法列
        for i in range(MAX_TEST_METHODS, 0, -1):
            cols = [f"检测项目{i}", f"检测方法{i}", f"检测仪器{i}"]
            if all(c in df_zh.columns for c in cols):
                if all(df_zh[c].astype(str).str.strip().eq('').all() for c in cols):
                    df_zh = df_zh.drop(columns=cols)
                else:
                    break

        df_zh.to_excel(zh_excel, index=False)
        print(f"  [OK] 中文报告：{len(zh_results)} 份 → {os.path.basename(zh_excel)}")
    else:
        print("  [警告]  无中文报告")

    # 4.2 英文报告Excel
    en_excel = os.path.join(OUTPUT_FOLDER, f"英文报告提取结果_{timestamp}.xlsx")
    if en_results:
        df_en = pd.DataFrame(en_results)
        for col in en_headers:
            if col not in df_en.columns:
                df_en[col] = ""
        df_en = df_en[en_headers]

        for i in range(MAX_TEST_METHODS, 0, -1):
            cols = [f"检测项目{i}", f"检测方法{i}", f"检测仪器{i}"]
            if all(c in df_en.columns for c in cols):
                if all(df_en[c].astype(str).str.strip().eq('').all() for c in cols):
                    df_en = df_en.drop(columns=cols)
                else:
                    break

        df_en.to_excel(en_excel, index=False)
        print(f"  [OK] 英文报告：{len(en_results)} 份 → {os.path.basename(en_excel)}")
    else:
        print("  [警告]  无英文报告")

    # 4.3 汇总版Excel
    merged_excel = os.path.join(OUTPUT_FOLDER, f"中英文汇总提取结果_{timestamp}.xlsx")
    all_results = zh_results + en_results
    if all_results:
        df_merged = pd.DataFrame(all_results)
        for col in merged_headers:
            if col not in df_merged.columns:
                df_merged[col] = ""
        df_merged = df_merged[merged_headers]

        for i in range(MAX_TEST_METHODS, 0, -1):
            cols = [f"检测项目{i}", f"检测方法{i}", f"检测仪器{i}"]
            if all(c in df_merged.columns for c in cols):
                if all(df_merged[c].astype(str).str.strip().eq('').all() for c in cols):
                    df_merged = df_merged.drop(columns=cols)
                else:
                    break

        # 待人工确认项目 sheet
        df_no_report = pd.DataFrame(no_report_projects) if no_report_projects else pd.DataFrame(
            columns=["项目/文件夹名称", "完整路径", "文件总数", "DOCX数量",
                     "DOCX文件名示例", "情况说明"])

        # 异常清单
        df_errors = pd.DataFrame([r for r in all_results if "失败" in str(r.get("提取状态", ""))])

        with pd.ExcelWriter(merged_excel, engine='openpyxl') as writer:
            df_merged.to_excel(writer, sheet_name='提取结果汇总', index=False)
            df_no_report.to_excel(writer, sheet_name='待人工确认项目', index=False)
            if not df_errors.empty:
                df_errors.to_excel(writer, sheet_name='提取异常清单', index=False)

        print(f"  [OK] 汇总版：{len(all_results)} 份 → {os.path.basename(merged_excel)}")
        print(f"     含 3 个 Sheet：提取结果汇总 / 待人工确认项目 / 提取异常清单")

    # 汇总统计
    success = len(all_results) - error_count

    print(f"\n[完成] 全部完成！")
    print(f"   中文报告：{len(zh_results)} 份")
    print(f"   英文报告：{len(en_results)} 份")
    print(f"   成功：{success} 份 | 失败：{error_count} 份")
    print(f"   提取耗时：{phase2_time:.1f} 秒")
    print(f"\n[目录] 输出文件目录：{OUTPUT_FOLDER}")
    print(f"[文件夹] 每份报告的完整路径都在 Excel 的「文件路径」列里")


if __name__ == "__main__":
    main()
