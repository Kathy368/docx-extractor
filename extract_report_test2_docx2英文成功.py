# -*- coding: utf-8 -*-
"""
英文检测报告信息提取工具
功能：从英文RoHS/ELV检测报告Word文档中提取关键信息，生成Excel表格
使用说明：
1. 修改下方【配置区】的文件路径和输出路径
2. 运行脚本即可生成Excel文件
3. 支持自动修复部分损坏的docx文件（如NULL关系问题）
依赖库：python-docx, pandas, openpyxl
安装命令：pip install python-docx pandas openpyxl
"""
import os
import re
import sys
import zipfile
import shutil
import tempfile
import pandas as pd
# ==============================================
# 【配置区】请根据实际情况修改以下参数
# ==============================================
# 待处理的Word文档路径（支持Windows路径，如 r"D:\Kathy\PDF提取工具\test.docx"）
INPUT_FILE = r"D:\Kathy\PDF提取工具\ROHS-S19062103203001.docx"
# 输出Excel文件路径
OUTPUT_EXCEL = r"D:\Kathy\PDF提取工具\英测4提取结果.xlsx"
# 最大检测方法数量（预留列数）
MAX_TEST_METHODS = 100
# 是否显示详细处理日志
VERBOSE = True
# ==============================================
# 【英文字段映射表】英文表头 → 中文含义
# ==============================================
EN_FIELD_MAPPING = {
    # 申请商相关
    "Applicant": "申请商",
    "Client": "申请商",
    "Company": "申请商",
    "Address": "申请商地址",
    
    # 样品相关
    "Product name": "样品名称",
    "Product Name": "样品名称",
    "Sample name": "样品名称",
    "Sample Name": "样品名称",
    "Sample Description": "样品名称",
    "Model": "零件号",
    "Part No.": "零件号",
    "Part Number": "零件号",
    "Model No.": "零件号",
    "Client Ref. Info.": "客户参考信息",
    "Client Reference": "客户参考信息",
    "Material": "材质",
    "Production Date": "生产日期",
    "Manufacturer": "制造商",
    "Manufacturer & Factory": "制造商",
    "Factory": "制造商",
    "Trade mark": "商标",
    
    # 样品编号相关
    "Sample No.": "样品编号",
    "Sample Number": "样品编号",
    "Sample ID": "样品编号",
    
    # 日期相关
    "Sample Received Date": "样品接收日期",
    "Date of Receipt": "样品接收日期",
    "Date Received": "样品接收日期",
    "Testing Period": "检测期间",
    "Test Period": "检测期间",
    "Test Date": "检测期间",
    "Date": "报告日期",
    "Report Date": "报告日期",
    "Issue Date": "报告日期",
    
    # 人员相关
    "Compiled by": "编制",
    "Prepared by": "编制",
    "Reviewed by": "审核",
    "Checked by": "审核",
    "Approved by": "批准",
    
    # 其他
    "Note": "说明",
    "Notes": "说明",
    "Remark": "备注",
    "Remarks": "备注",
}
# ==============================================
# 工具函数：修复有问题的docx文件
# ==============================================
def fix_docx_if_needed(file_path):
    """
    检查并修复有问题的docx文件（如NULL关系问题）
    返回修复后的文件路径，如果不需要修复则返回原路径
    """
    try:
        # 先尝试直接打开
        from docx import Document
        Document(file_path)
        return file_path
    except KeyError as e:
        if "NULL" in str(e):
            if VERBOSE:
                print(f"  ⚠️  检测到文档有NULL关系问题，正在自动修复...")
            
            # 创建临时修复文件
            temp_dir = tempfile.gettempdir()
            base_name = os.path.basename(file_path)
            fixed_path = os.path.join(temp_dir, f"fixed_{base_name}")
            
            try:
                with zipfile.ZipFile(file_path, 'r') as zin:
                    with zipfile.ZipFile(fixed_path, 'w') as zout:
                        for item in zin.infolist():
                            if item.filename == 'word/_rels/document.xml.rels':
                                # 修复NULL关系（支持NULL、../NULL等多种形式）
                                content = zin.read(item.filename).decode('utf-8')
                                content = re.sub(r'<Relationship[^>]*Target="[^"]*NULL[^"]*"[^>]*/>', '', content)
                                zout.writestr(item, content)
                            else:
                                zout.writestr(item, zin.read(item.filename))
                
                if VERBOSE:
                    print(f"  ✅ 文档修复成功")
                return fixed_path
            except Exception as fix_error:
                if VERBOSE:
                    print(f"  ❌ 文档修复失败: {fix_error}")
                raise
        else:
            raise
    except Exception:
        # 其他错误直接抛出
        raise
# ==============================================
# 工具函数：获取单元格完整文本（包括SDT内容控件）
# ==============================================
def get_cell_text(cell):
    """
    获取单元格的完整文本，包括SDT（结构化文档标签/内容控件）中的内容
    解决python-docx无法读取内容控件文本的问题
    """
    import re
    import html
    xml = cell._tc.xml
    # 提取所有w:t标签中的文本
    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    # 还原HTML转义字符（如&amp; → &）
    result = ''.join(texts)
    result = html.unescape(result)
    return result.strip()

def get_row_cells_text(row):
    """
    获取一行中所有单元格的文本，包括被SDT（结构化文档标签）包裹的单元格
    解决python-docx无法识别SDT包裹的tc单元格的问题
    """
    import re
    import html
    xml = row._tr.xml
    # 提取所有w:tc元素（包括在w:sdtContent里面的）
    tcs = re.findall(r'<w:tc\b.*?</w:tc>', xml, re.DOTALL)
    cells_text = []
    for tc_xml in tcs:
        # 提取每个tc中的所有w:t文本
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', tc_xml)
        cell_text = ''.join(texts)
        # 还原HTML转义字符（如&amp; → &）
        cell_text = html.unescape(cell_text)
        cells_text.append(cell_text.strip())
    return cells_text

# ==============================================
# 函数1：提取基本信息
# ==============================================
def extract_basic_info(doc):
    """从两列表格中提取基本信息字段"""
    fields = {}
    manufacturer_address_found = False  # 标记是否已经遇到制造商，用于区分多个Address
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 1:
            continue
        
        # 使用get_row_cells_text获取完整的单元格列表（包括SDT包裹的单元格）
        first_row = get_row_cells_text(table.rows[0])
        first_row_text = ' '.join(first_row)
        first_cell_count = len(first_row)
        
        # 跳过检测结果表格
        if 'Part No.' in first_row_text or 'Test Items' in first_row_text:
            continue
        
        # 跳过检测方法表格
        if 'Test item' in first_row_text and 'Test method' in first_row_text:
            continue
        
        # 跳过XRF限值表格
        if 'Limit of IEC' in first_row_text or 'XRF screening' in first_row_text:
            continue
        
        # 跳过检测要求/结论表格（单独处理）
        if 'Test Requirement' in first_row_text or 'Test Conclusion' in first_row_text:
            continue
        
        # 处理两列表格（键值对）
        if first_cell_count == 2:
            for row_idx, row in enumerate(table.rows):
                # 使用get_row_cells_text获取完整单元格文本（包括SDT包裹的）
                cells = [c.replace('\n', ' ') for c in get_row_cells_text(row)]
                
                # 处理只有1个单元格的行（可能是合并单元格的标签行）
                if len(cells) == 1:
                    key = cells[0].rstrip('：:').strip()
                    if key and len(key) < 60:
                        # 只有标签没有值，先记录下来
                        fields[key] = ""
                    continue
                
                if len(cells) == 2:
                    key = cells[0].rstrip('：:').strip()
                    val = cells[1].strip()
                    
                    # 清理key中的多余空格（处理换行导致的多空格问题）
                    key = re.sub(r'\s+', ' ', key).strip()
                    
                    if key and len(key) < 60:
                        # 跳过说明性文字
                        if 'The following sample' in key or 'was/were submitted' in key:
                            continue
                        
                        # 特殊处理：遇到Manufacturer/Factory后，下一个Address是制造商地址
                        if key in ['Manufacturer', 'Manufacturer & Factory', 'Factory']:
                            manufacturer_address_found = True
                            if val:
                                fields[key] = val
                            continue
                        
                        # 特殊处理Address：区分申请商地址和制造商地址
                        if key == 'Address':
                            if not manufacturer_address_found:
                                # 第一个Address是申请商地址
                                if val:
                                    fields[key] = val
                            else:
                                # 制造商后面的Address是制造商地址
                                if val and '制造商地址' not in fields:
                                    fields['制造商地址'] = val
                            continue
                        
                        if val:
                            fields[key] = val
        
        # 处理四列表格（编制/审核/批准/日期等）
        elif first_cell_count == 4:
            for row in table.rows:
                # 使用get_row_cells_text获取完整单元格文本（包括SDT包裹的）
                cells = [c.replace('\n', ' ') for c in get_row_cells_text(row)]
                if len(cells) == 4:
                    for i in range(0, 4, 2):
                        key = cells[i].rstrip('：:').strip()
                        val = cells[i+1].strip()
                        if key and len(key) < 40:
                            if val:
                                fields[key] = val
    
    return fields
# ==============================================
# 函数2：提取检测要求和检测结论（完整原文）
# ==============================================
def extract_test_requirement_and_conclusion(doc):
    """提取Test Requirement和Test Conclusion的完整原文（优先从首页提取）"""
    test_requirement = ""
    test_conclusion = ""
    
    for table_idx, table in enumerate(doc.tables):
        # 如果两个都找到了，就不用继续找了（优先前面的表格/首页）
        if test_requirement and test_conclusion:
            break
            
        if len(table.rows) < 2:
            continue
        
        # 使用get_cell_text获取完整文本（包括SDT内容控件）
        first_row = [get_cell_text(c) for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)
        
        # 情况1：2行2列表格（Test Requirement和Conclusion在同一行）- 最常见的首页格式
        if (len(table.rows[0].cells) == 2 and 
            'Test Requirement' in first_row_text and 
            'Conclusion' in first_row_text):
            
            if VERBOSE:
                print(f"  找到检测要求/结论表格（2行2列格式，首页）")
            
            if len(table.rows) >= 2:
                test_requirement = get_cell_text(table.rows[1].cells[0])
                test_conclusion = get_cell_text(table.rows[1].cells[1])
            break  # 首页找到就直接返回
        
        # 情况2：2行1列表格（Test Requirement单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Requirement' in first_row_text:
            if not test_requirement:
                if VERBOSE:
                    print(f"  找到检测要求表格（独立表格）")
                
                if len(table.rows) >= 2:
                    test_requirement = get_cell_text(table.rows[1].cells[0])
        
        # 情况3：2行1列表格（Test Conclusion单独表格）
        elif len(table.rows[0].cells) == 1 and 'Test Conclusion' in first_row_text:
            if not test_conclusion:
                if VERBOSE:
                    print(f"  找到检测结论表格（独立表格）")
                
                if len(table.rows) >= 2:
                    test_conclusion = get_cell_text(table.rows[1].cells[0])
    
    return test_requirement, test_conclusion
# ==============================================
# 函数3：提取样品部件总数（从Test Result表格）
# ==============================================
def extract_sample_count(doc):
    """从所有检测结果表格中提取最大的部件序号"""
    max_seq = 0
    table_count = 0
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        
        table_max = 0
        is_result_table = False
        
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if cells:
                first_cell = cells[0]
                
                # 检查是否有Part No.表头
                if 'Part No.' in first_cell:
                    is_result_table = True
                    continue
                
                # 检查第一列是否是纯数字（部件序号）
                m = re.match(r'^(\d+)$', first_cell)
                if m:
                    seq = int(m.group(1))
                    if seq > table_max:
                        table_max = seq
        
        if table_max > 0:
            table_count += 1
            if table_max > max_seq:
                max_seq = table_max
    
    if VERBOSE:
        print(f"  找到 {table_count} 个检测结果表格，最大部件序号: {max_seq}")
    
    return max_seq
# ==============================================
# 函数4：提取检测方法（从Chemical Test章节）
# ==============================================
def extract_test_methods(doc):
    """从Chemical Test章节的检测方法表格中提取"""
    methods = []
    seen_items = set()
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 3:
            continue
        
        first_row = [c.text.strip() for c in table.rows[0].cells]
        first_row_text = ' '.join(first_row)
        
        # 检测方法表格特征：有Test item和Test method列
        if 'Test item' not in first_row_text or 'Test method' not in first_row_text:
            continue
        
        if VERBOSE:
            print(f"  找到检测方法表格（{len(table.rows)}行）")
        
        for row_idx, row in enumerate(table.rows[1:]):
            cells = [c.text.strip().replace('\n', '') for c in row.cells]
            if len(cells) < 3:
                continue
            
            item = cells[0].strip()
            method = cells[1].strip() if len(cells) > 1 else ""
            instrument = cells[2].strip() if len(cells) > 2 else ""
            
            # 跳过说明行和空行
            if not item or not method:
                continue
            if 'Limit' in item or '△' in item or 'Note' in item:
                continue
            
            # 按检测项目名称去重
            if item in seen_items:
                continue
            seen_items.add(item)
            
            methods.append({
                "检测项目": item,
                "检测方法": method,
                "检测仪器": instrument
            })
    
    if VERBOSE:
        print(f"  共提取到 {len(methods)} 项检测方法")
    
    return methods
# ==============================================
# 主函数：处理单个文档
# ==============================================
def process_word_file(input_path, output_excel):
    """
    处理单个Word检测报告，提取信息并生成Excel
    
    参数:
        input_path: Word文档路径
        output_excel: 输出Excel路径
    """
    from docx import Document
    
    print("=" * 60)
    print(f"开始处理: {os.path.basename(input_path)}")
    print("=" * 60)
    
    # 1. 检查文件是否存在
    if not os.path.exists(input_path):
        print(f"❌ 错误：文件不存在 - {input_path}")
        return False
    
    # 2. 修复文档（如果需要）
    try:
        actual_file = fix_docx_if_needed(input_path)
    except Exception as e:
        print(f"❌ 错误：无法打开文档 - {e}")
        print("   请检查文件是否为有效的Word文档(.docx格式)")
        return False
    
    # 3. 打开文档
    try:
        doc = Document(actual_file)
    except ImportError:
        print("❌ 错误：未安装 python-docx 库")
        print("   请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 错误：打开文档失败 - {e}")
        return False
    
    if VERBOSE:
        print(f"\n📄 文档信息:")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")
    
    # 4. 初始化结果字典
    result = {
        "文件名": os.path.basename(input_path),
        "报告编号": "",
        "申请商": "",
        "申请商地址": "",
        "制造商地址": "",
        "样品名称": "",
        "零件号": "",
        "商标": "",
        "客户参考信息": "",
        "材质": "",
        "生产日期": "",
        "制造商": "",
        "样品编号": "",
        "样品接收日期": "",
        "检测期间": "",
        "检测要求": "",
        "检测结论": "",
        "报告日期": "",
        "样品部件总数": "",
        "编制": "",
        "审核": "",
        "批准": "",
        "备注": "",
    }
    
    # 添加检测方法列
    for i in range(1, MAX_TEST_METHODS + 1):
        result[f"检测项目{i}"] = ""
        result[f"检测方法{i}"] = ""
        result[f"检测仪器{i}"] = ""
    
    # 5. 提取基本信息
    if VERBOSE:
        print(f"\n📋 【1/4】提取基本信息...")
    
    raw_fields = extract_basic_info(doc)
    
    # 映射英文字段到中文（支持模糊匹配，忽略大小写和空格差异）
    def normalize_field(s):
        """标准化字段名：小写、去空格、去冒号"""
        return re.sub(r'[\s:：]', '', s).lower()
    
    # 先构建标准化的映射表
    normalized_mapping = {}
    for en_key, cn_key in EN_FIELD_MAPPING.items():
        normalized_mapping[normalize_field(en_key)] = cn_key
    
    # 匹配字段
    for raw_key, raw_val in raw_fields.items():
        norm_key = normalize_field(raw_key)
        if norm_key in normalized_mapping:
            cn_key = normalized_mapping[norm_key]
            if cn_key in result and raw_val:
                result[cn_key] = raw_val
    
    # 单独处理制造商地址
    if '制造商地址' in raw_fields and '制造商地址' in result:
        result['制造商地址'] = raw_fields['制造商地址']
    
    # 从文件名提取报告编号
    file_name = os.path.basename(input_path)
    m = re.search(r'([A-Z]{1,5}\d{6,15}[A-Z]?\d*)', file_name, re.IGNORECASE)
    if m and not result["报告编号"]:
        result["报告编号"] = m.group(1)
    
    if VERBOSE:
        basic_count = sum(1 for k, v in result.items() 
                         if v and k not in ["文件名"] and not k.startswith("检测"))
        print(f"  ✅ 提取到 {basic_count} 个基本信息字段")
    
    # 6. 提取检测要求和检测结论
    if VERBOSE:
        print(f"\n📋 【2/4】提取检测要求和结论...")
    
    test_req, test_con = extract_test_requirement_and_conclusion(doc)
    result["检测要求"] = test_req
    result["检测结论"] = test_con
    
    # 7. 提取样品部件总数
    if VERBOSE:
        print(f"\n📋 【3/4】统计样品部件总数...")
    
    sample_count = extract_sample_count(doc)
    result["样品部件总数"] = str(sample_count)
    
    # 8. 提取检测方法
    if VERBOSE:
        print(f"\n📋 【4/4】提取检测方法...")
    
    test_methods = extract_test_methods(doc)
    
    for i, method in enumerate(test_methods[:MAX_TEST_METHODS], 1):
        result[f"检测项目{i}"] = method["检测项目"]
        result[f"检测方法{i}"] = method["检测方法"]
        result[f"检测仪器{i}"] = method["检测仪器"]
    
    # 9. 生成Excel
    print(f"\n💾 生成Excel文件...")
    
    # 构建列顺序
    col_order = [
        "报告编号", "申请商", "申请商地址", "制造商地址",
        "样品名称", "零件号", "商标",
        "客户参考信息", "材质", "生产日期", "制造商", "样品编号",
        "样品接收日期", "检测期间", "检测要求", "检测结论",
        "报告日期", "样品部件总数",
        "编制", "审核", "批准"
    ]
    
    # 添加检测方法列
    actual_method_count = min(len(test_methods), MAX_TEST_METHODS)
    for i in range(1, actual_method_count + 1):
        col_order.append(f"检测项目{i}")
        col_order.append(f"检测方法{i}")
        col_order.append(f"检测仪器{i}")
    
    col_order.extend(["备注", "文件名"])
    
    # 创建DataFrame
    df = pd.DataFrame([result])
    
    # 确保所有列都存在
    for col in col_order:
        if col not in df.columns:
            df[col] = ""
    
    df = df[col_order]
    
    # 保存Excel
    try:
        df.to_excel(output_excel, index=False)
        print(f"✅ 成功保存: {output_excel}")
    except ImportError:
        print("❌ 错误：未安装 openpyxl 库")
        print("   请运行: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 错误：保存Excel失败 - {e}")
        return False
    
    # 统计信息
    non_empty = sum(1 for val in result.values() if val)
    print(f"\n📊 提取完成！共提取 {non_empty} 个非空字段")
    print("=" * 60)
    
    return True
# ==============================================
# 程序入口
# ==============================================
def main():
    print("\n" + "=" * 60)
    print("  英文检测报告信息提取工具")
    print("=" * 60)
    
    # 检查依赖
    try:
        import docx
        import pandas
    except ImportError as e:
        print(f"\n❌ 缺少依赖库: {e}")
        print("请运行以下命令安装依赖:")
        print("  pip install python-docx pandas openpyxl")
        sys.exit(1)
    
    # 处理文件
    success = process_word_file(INPUT_FILE, OUTPUT_EXCEL)
    
    if success:
        print("\n🎉 处理完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)
if __name__ == "__main__":
    main()
